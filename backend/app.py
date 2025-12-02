from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from datetime import datetime, timedelta
import jwt
import sqlite3
import hashlib
import uuid
import os
import io
from dotenv import load_dotenv

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter

# Load environment variables
load_dotenv()  # Try .env first
load_dotenv('.env.example')  # Fallback to .env.example

# Try to import Groq (optional dependency)
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False
    print("⚠️  Groq not installed. AI features will use rule-based fallback.")

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

SECRET_KEY = "supersecret123"
DATABASE = "users.db"
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")

# Initialize Groq client if available
groq_client = None
if GROQ_AVAILABLE and GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print("✅ Groq AI client initialized")
    except Exception as e:
        print(f"⚠️  Groq initialization failed: {e}")


def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok"}), 200

@app.route('/api/v1/auth/register', methods=['POST', 'OPTIONS'])
def register():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        full_name = data.get('full_name', '')
        org_name = data.get('organization_name', '')
        
        if not email or not password:
            return jsonify({"detail": "Email and password required"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check if user exists
        c.execute("SELECT * FROM users WHERE email = ?", (email,))
        if c.fetchone():
            conn.close()
            return jsonify({"detail": "Email already registered"}), 400
        
        # Create user
        user_id = str(uuid.uuid4())
        hashed_pw = hash_password(password)
        created_at = datetime.utcnow().isoformat()
        default_tier = 'free'
        default_project_limit = 3
        
        c.execute("""INSERT INTO users (id, email, password, full_name, organization_name, created_at, tier, tier_expires_at, project_limit)
                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                  (user_id, email, hashed_pw, full_name, org_name, created_at, default_tier, None, default_project_limit))
        conn.commit()
        conn.close()
        
        # Create token
        token = jwt.encode({
            'user_id': user_id,
            'email': email,
            'exp': datetime.utcnow() + timedelta(hours=24)
        }, SECRET_KEY, algorithm='HS256')
        
        return jsonify({
            "access_token": token,
            "token_type": "bearer",
            "user": {
                "id": user_id,
                "email": email,
                "full_name": full_name,
                "organization_name": org_name,
                "tier": default_tier,
                "tier_expires_at": None,
                "project_limit": default_project_limit
            }
        }), 201
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/auth/login', methods=['POST', 'OPTIONS'])
def login():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = request.get_json()
            email = data.get('email')
            password = data.get('password')
        else:
            email = request.form.get('username') or request.form.get('email')
            password = request.form.get('password')
        
        if not email or not password:
            return jsonify({"detail": "Email and password required"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("SELECT id, email, password, full_name, organization_name, created_at, tier, tier_expires_at, project_limit FROM users WHERE email = ?", (email,))
        user = c.fetchone()
        conn.close()
        
        if not user or user[2] != hash_password(password):
            return jsonify({"detail": "Invalid credentials"}), 401
        
        # Create token
        token = jwt.encode({
            'user_id': user[0],
            'email': user[1],
            'exp': datetime.utcnow() + timedelta(hours=24)
        }, SECRET_KEY, algorithm='HS256')
        
        return jsonify({
            "access_token": token,
            "token_type": "bearer",
            "user": {
                "id": user[0],
                "email": user[1],
                "full_name": user[3],
                "organization_name": user[4],
                "tier": user[6] or 'free',
                "tier_expires_at": user[7],
                "project_limit": user[8] or 3
            }
        }), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/auth/me', methods=['GET', 'OPTIONS'])
def get_me():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("SELECT id, email, password, full_name, organization_name, created_at, tier, tier_expires_at, project_limit FROM users WHERE id = ?", (payload['user_id'],))
        user = c.fetchone()
        
        # Get project count for this user
        c.execute("SELECT COUNT(*) FROM projects WHERE user_id = ?", (payload['user_id'],))
        project_count = c.fetchone()[0]
        conn.close()
        
        if not user:
            return jsonify({"detail": "User not found"}), 404
        
        tier = user[6] or 'free'
        project_limit = user[8] or 3
        
        # Set limits based on tier
        tier_limits = {
            'free': {'projects': 3, 'cbam_export': False, 'brsr_export': False, 'scenario_compare': False, 'ai_advisor': False, 'verification': False},
            'pro': {'projects': -1, 'cbam_export': True, 'brsr_export': True, 'scenario_compare': True, 'ai_advisor': True, 'verification': False},
            'enterprise': {'projects': -1, 'cbam_export': True, 'brsr_export': True, 'scenario_compare': True, 'ai_advisor': True, 'verification': True}
        }
        
        limits = tier_limits.get(tier, tier_limits['free'])
        
        return jsonify({
            "id": user[0],
            "email": user[1],
            "full_name": user[3],
            "organization_name": user[4],
            "tier": tier,
            "tier_expires_at": user[7],
            "project_limit": limits['projects'],
            "project_count": project_count,
            "features": limits
        }), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": "Invalid token"}), 401

@app.route('/api/v1/auth/profile', methods=['PUT', 'OPTIONS'])
def update_profile():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.json
        full_name = data.get('full_name', '').strip()
        organization_name = data.get('organization_name', '').strip()
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("""UPDATE users SET full_name = ?, organization_name = ? WHERE id = ?""",
                  (full_name or None, organization_name or None, payload['user_id']))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Profile updated successfully"}), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error updating profile: {e}")
        return jsonify({"detail": "Failed to update profile"}), 500

@app.route('/api/v1/auth/change-password', methods=['PUT', 'OPTIONS'])
def change_password():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.json
        current_password = data.get('current_password', '')
        new_password = data.get('new_password', '')
        
        if not current_password or not new_password:
            return jsonify({"detail": "Both current and new password are required"}), 400
        
        if len(new_password) < 6:
            return jsonify({"detail": "New password must be at least 6 characters"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("SELECT password FROM users WHERE id = ?", (payload['user_id'],))
        user = c.fetchone()
        
        if not user:
            conn.close()
            return jsonify({"detail": "User not found"}), 404
        
        if not bcrypt.checkpw(current_password.encode('utf-8'), user[0].encode('utf-8')):
            conn.close()
            return jsonify({"detail": "Current password is incorrect"}), 400
        
        new_hashed = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        c.execute("UPDATE users SET password = ? WHERE id = ?", (new_hashed, payload['user_id']))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Password changed successfully"}), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error changing password: {e}")
        return jsonify({"detail": "Failed to change password"}), 500

@app.route('/api/v1/auth/stats', methods=['GET', 'OPTIONS'])
def get_account_stats():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project count
        c.execute("SELECT COUNT(*) FROM projects WHERE user_id = ?", (payload['user_id'],))
        project_count = c.fetchone()[0]
        
        # Get total GWP across all projects
        c.execute("SELECT COALESCE(SUM(gwp_total), 0) FROM projects WHERE user_id = ?", (payload['user_id'],))
        total_gwp = c.fetchone()[0] or 0
        
        # Get analyses run (projects with status 'calculated' or 'verified')
        c.execute("SELECT COUNT(*) FROM projects WHERE user_id = ? AND status IN ('calculated', 'verified')", (payload['user_id'],))
        analyses_run = c.fetchone()[0]
        
        # Get reports generated (we can track this based on projects that have been exported)
        # For now, use calculated projects as a proxy
        reports_generated = analyses_run
        
        conn.close()
        
        return jsonify({
            "project_count": project_count,
            "total_gwp": round(total_gwp, 2),
            "analyses_run": analyses_run,
            "reports_generated": reports_generated
        }), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error getting account stats: {e}")
        return jsonify({"detail": "Failed to get account stats"}), 500

@app.route('/api/v1/auth/upgrade', methods=['POST', 'OPTIONS'])
def upgrade_to_pro():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.json
        target_tier = data.get('tier', 'pro')
        
        if target_tier not in ['pro', 'enterprise']:
            return jsonify({"detail": "Invalid tier"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Update user tier to pro (without payment for now)
        c.execute("""UPDATE users SET tier = ?, project_limit = ? WHERE id = ?""",
                  (target_tier, -1, payload['user_id']))  # -1 means unlimited
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": f"Successfully upgraded to {target_tier.title()} plan!", "tier": target_tier}), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error upgrading account: {e}")
        return jsonify({"detail": "Failed to upgrade account"}), 500

@app.route('/api/v1/auth/account', methods=['DELETE', 'OPTIONS'])
def delete_account():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Delete all user's materials first (due to foreign key)
        c.execute("""DELETE FROM materials WHERE project_id IN 
                     (SELECT id FROM projects WHERE user_id = ?)""", (payload['user_id'],))
        
        # Delete all user's projects
        c.execute("DELETE FROM projects WHERE user_id = ?", (payload['user_id'],))
        
        # Delete the user
        c.execute("DELETE FROM users WHERE id = ?", (payload['user_id'],))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Account deleted successfully"}), 200
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token expired"}), 401
    except Exception as e:
        print(f"Error deleting account: {e}")
        return jsonify({"detail": "Failed to delete account"}), 500


# ==================== TEAM MANAGEMENT API ====================

def get_user_from_token():
    """Helper function to extract user from JWT token"""
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        return None, jsonify({"detail": "Not authenticated"}), 401
    
    try:
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        return payload['user_id'], None, None
    except jwt.ExpiredSignatureError:
        return None, jsonify({"detail": "Token expired"}), 401
    except Exception:
        return None, jsonify({"detail": "Invalid token"}), 401


@app.route('/api/v1/teams', methods=['GET', 'OPTIONS'])
def list_teams():
    """List all teams the user is a member of"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get teams where user is owner or member
        c.execute("""
            SELECT DISTINCT t.id, t.name, t.description, t.owner_id, t.created_at,
                   (SELECT COUNT(*) FROM team_members WHERE team_id = t.id) as member_count,
                   CASE WHEN t.owner_id = ? THEN 'owner' 
                        ELSE (SELECT role FROM team_members WHERE team_id = t.id AND user_id = ?) 
                   END as user_role
            FROM teams t
            LEFT JOIN team_members tm ON t.id = tm.team_id
            WHERE t.owner_id = ? OR tm.user_id = ?
            ORDER BY t.created_at DESC
        """, (user_id, user_id, user_id, user_id))
        
        teams = []
        for row in c.fetchall():
            teams.append({
                "id": row[0],
                "name": row[1],
                "description": row[2],
                "owner_id": row[3],
                "created_at": row[4],
                "member_count": row[5] or 1,
                "user_role": row[6] or 'owner'
            })
        
        conn.close()
        return jsonify(teams), 200
        
    except Exception as e:
        print(f"Error listing teams: {e}")
        return jsonify({"detail": "Failed to list teams"}), 500


@app.route('/api/v1/teams', methods=['POST', 'OPTIONS'])
def create_team():
    """Create a new team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        description = data.get('description', '').strip()
        
        if not name:
            return jsonify({"detail": "Team name is required"}), 400
        
        team_id = str(uuid.uuid4())
        created_at = datetime.utcnow().isoformat()
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        c.execute("""INSERT INTO teams (id, name, description, owner_id, created_at)
                     VALUES (?, ?, ?, ?, ?)""",
                  (team_id, name, description, user_id, created_at))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "id": team_id,
            "name": name,
            "description": description,
            "owner_id": user_id,
            "created_at": created_at,
            "message": "Team created successfully"
        }), 201
        
    except Exception as e:
        print(f"Error creating team: {e}")
        return jsonify({"detail": "Failed to create team"}), 500


@app.route('/api/v1/teams/<team_id>', methods=['GET', 'OPTIONS'])
def get_team(team_id):
    """Get team details with members"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get team
        c.execute("SELECT * FROM teams WHERE id = ?", (team_id,))
        team = c.fetchone()
        
        if not team:
            conn.close()
            return jsonify({"detail": "Team not found"}), 404
        
        # Check if user has access
        c.execute("""SELECT 1 FROM teams WHERE id = ? AND owner_id = ?
                     UNION
                     SELECT 1 FROM team_members WHERE team_id = ? AND user_id = ?""",
                  (team_id, user_id, team_id, user_id))
        
        if not c.fetchone():
            conn.close()
            return jsonify({"detail": "Access denied"}), 403
        
        # Get members
        c.execute("""
            SELECT u.id, u.email, u.full_name, tm.role, tm.joined_at
            FROM team_members tm
            JOIN users u ON tm.user_id = u.id
            WHERE tm.team_id = ?
        """, (team_id,))
        
        members = []
        for row in c.fetchall():
            members.append({
                "id": row[0],
                "email": row[1],
                "full_name": row[2],
                "role": row[3],
                "joined_at": row[4]
            })
        
        # Get owner info
        c.execute("SELECT id, email, full_name FROM users WHERE id = ?", (team[3],))
        owner = c.fetchone()
        
        # Get pending invites
        c.execute("""SELECT id, email, status, created_at FROM team_invites 
                     WHERE team_id = ? AND status = 'pending'""", (team_id,))
        invites = []
        for row in c.fetchall():
            invites.append({
                "id": row[0],
                "email": row[1],
                "status": row[2],
                "created_at": row[3]
            })
        
        # Get shared projects
        c.execute("""
            SELECT p.id, p.name, p.status, pc.permission
            FROM project_collaborators pc
            JOIN projects p ON pc.project_id = p.id
            WHERE pc.team_id = ?
        """, (team_id,))
        
        projects = []
        for row in c.fetchall():
            projects.append({
                "id": row[0],
                "name": row[1],
                "status": row[2],
                "permission": row[3]
            })
        
        conn.close()
        
        return jsonify({
            "id": team[0],
            "name": team[1],
            "description": team[2],
            "owner": {
                "id": owner[0],
                "email": owner[1],
                "full_name": owner[2]
            } if owner else None,
            "created_at": team[4],
            "members": members,
            "pending_invites": invites,
            "shared_projects": projects,
            "is_owner": team[3] == user_id
        }), 200
        
    except Exception as e:
        print(f"Error getting team: {e}")
        return jsonify({"detail": "Failed to get team"}), 500


@app.route('/api/v1/teams/<team_id>', methods=['PUT', 'OPTIONS'])
def update_team(team_id):
    """Update team details"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check ownership
        c.execute("SELECT owner_id FROM teams WHERE id = ?", (team_id,))
        team = c.fetchone()
        
        if not team:
            conn.close()
            return jsonify({"detail": "Team not found"}), 404
        
        if team[0] != user_id:
            conn.close()
            return jsonify({"detail": "Only team owner can update team"}), 403
        
        data = request.get_json()
        name = data.get('name', '').strip()
        description = data.get('description', '').strip()
        
        if not name:
            conn.close()
            return jsonify({"detail": "Team name is required"}), 400
        
        c.execute("UPDATE teams SET name = ?, description = ? WHERE id = ?",
                  (name, description, team_id))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Team updated successfully"}), 200
        
    except Exception as e:
        print(f"Error updating team: {e}")
        return jsonify({"detail": "Failed to update team"}), 500


@app.route('/api/v1/teams/<team_id>', methods=['DELETE', 'OPTIONS'])
def delete_team(team_id):
    """Delete a team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check ownership
        c.execute("SELECT owner_id FROM teams WHERE id = ?", (team_id,))
        team = c.fetchone()
        
        if not team:
            conn.close()
            return jsonify({"detail": "Team not found"}), 404
        
        if team[0] != user_id:
            conn.close()
            return jsonify({"detail": "Only team owner can delete team"}), 403
        
        # Delete related data
        c.execute("DELETE FROM project_collaborators WHERE team_id = ?", (team_id,))
        c.execute("DELETE FROM team_invites WHERE team_id = ?", (team_id,))
        c.execute("DELETE FROM team_members WHERE team_id = ?", (team_id,))
        c.execute("DELETE FROM teams WHERE id = ?", (team_id,))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Team deleted successfully"}), 200
        
    except Exception as e:
        print(f"Error deleting team: {e}")
        return jsonify({"detail": "Failed to delete team"}), 500


@app.route('/api/v1/teams/<team_id>/invite', methods=['POST', 'OPTIONS'])
def invite_member(team_id):
    """Invite a member to the team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        data = request.get_json()
        email = data.get('email', '').strip().lower()
        
        if not email:
            return jsonify({"detail": "Email is required"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check if user has invite permissions (owner or admin)
        c.execute("""SELECT 1 FROM teams WHERE id = ? AND owner_id = ?
                     UNION
                     SELECT 1 FROM team_members WHERE team_id = ? AND user_id = ? AND role = 'admin'""",
                  (team_id, user_id, team_id, user_id))
        
        if not c.fetchone():
            conn.close()
            return jsonify({"detail": "You don't have permission to invite members"}), 403
        
        # Check if already a member
        c.execute("""SELECT u.id FROM users u 
                     JOIN team_members tm ON u.id = tm.user_id
                     WHERE u.email = ? AND tm.team_id = ?""", (email, team_id))
        if c.fetchone():
            conn.close()
            return jsonify({"detail": "User is already a team member"}), 400
        
        # Check for existing pending invite
        c.execute("SELECT id FROM team_invites WHERE team_id = ? AND email = ? AND status = 'pending'",
                  (team_id, email))
        if c.fetchone():
            conn.close()
            return jsonify({"detail": "Invitation already sent"}), 400
        
        # Check if user exists - if so, add directly
        c.execute("SELECT id FROM users WHERE email = ?", (email,))
        existing_user = c.fetchone()
        
        if existing_user:
            # Add user directly as member
            member_id = str(uuid.uuid4())
            joined_at = datetime.utcnow().isoformat()
            c.execute("""INSERT INTO team_members (id, team_id, user_id, role, joined_at)
                         VALUES (?, ?, ?, 'member', ?)""",
                      (member_id, team_id, existing_user[0], joined_at))
            conn.commit()
            conn.close()
            return jsonify({"message": "Member added successfully", "added_directly": True}), 201
        
        # Create invite for non-existing user
        invite_id = str(uuid.uuid4())
        created_at = datetime.utcnow().isoformat()
        
        c.execute("""INSERT INTO team_invites (id, team_id, email, invited_by, status, created_at)
                     VALUES (?, ?, ?, ?, 'pending', ?)""",
                  (invite_id, team_id, email, user_id, created_at))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "message": "Invitation sent",
            "invite_id": invite_id,
            "added_directly": False
        }), 201
        
    except Exception as e:
        print(f"Error inviting member: {e}")
        return jsonify({"detail": "Failed to send invitation"}), 500


@app.route('/api/v1/teams/<team_id>/members/<member_id>', methods=['PUT', 'OPTIONS'])
def update_member_role(team_id, member_id):
    """Update a member's role"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        data = request.get_json()
        role = data.get('role', 'member')
        
        if role not in ['admin', 'member']:
            return jsonify({"detail": "Invalid role"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check ownership
        c.execute("SELECT owner_id FROM teams WHERE id = ?", (team_id,))
        team = c.fetchone()
        
        if not team or team[0] != user_id:
            conn.close()
            return jsonify({"detail": "Only team owner can change roles"}), 403
        
        c.execute("UPDATE team_members SET role = ? WHERE team_id = ? AND user_id = ?",
                  (role, team_id, member_id))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Role updated successfully"}), 200
        
    except Exception as e:
        print(f"Error updating member role: {e}")
        return jsonify({"detail": "Failed to update role"}), 500


@app.route('/api/v1/teams/<team_id>/members/<member_id>', methods=['DELETE', 'OPTIONS'])
def remove_member(team_id, member_id):
    """Remove a member from the team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check if owner or self-removal
        c.execute("SELECT owner_id FROM teams WHERE id = ?", (team_id,))
        team = c.fetchone()
        
        if not team:
            conn.close()
            return jsonify({"detail": "Team not found"}), 404
        
        # Owner can remove anyone, members can only remove themselves
        if team[0] != user_id and member_id != user_id:
            conn.close()
            return jsonify({"detail": "You don't have permission to remove this member"}), 403
        
        c.execute("DELETE FROM team_members WHERE team_id = ? AND user_id = ?",
                  (team_id, member_id))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Member removed successfully"}), 200
        
    except Exception as e:
        print(f"Error removing member: {e}")
        return jsonify({"detail": "Failed to remove member"}), 500


@app.route('/api/v1/teams/<team_id>/projects', methods=['POST', 'OPTIONS'])
def share_project_with_team(team_id):
    """Share a project with a team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        data = request.get_json()
        project_id = data.get('project_id')
        permission = data.get('permission', 'view')
        
        if not project_id:
            return jsonify({"detail": "Project ID is required"}), 400
        
        if permission not in ['view', 'edit']:
            return jsonify({"detail": "Invalid permission"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check project ownership
        c.execute("SELECT user_id FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        
        if not project or project[0] != user_id:
            conn.close()
            return jsonify({"detail": "You can only share your own projects"}), 403
        
        # Check team membership
        c.execute("""SELECT 1 FROM teams WHERE id = ? AND owner_id = ?
                     UNION
                     SELECT 1 FROM team_members WHERE team_id = ? AND user_id = ?""",
                  (team_id, user_id, team_id, user_id))
        
        if not c.fetchone():
            conn.close()
            return jsonify({"detail": "You must be a team member to share projects"}), 403
        
        # Check if already shared
        c.execute("SELECT id FROM project_collaborators WHERE project_id = ? AND team_id = ?",
                  (project_id, team_id))
        existing = c.fetchone()
        
        if existing:
            # Update permission
            c.execute("UPDATE project_collaborators SET permission = ? WHERE id = ?",
                      (permission, existing[0]))
        else:
            # Create new share
            collab_id = str(uuid.uuid4())
            added_at = datetime.utcnow().isoformat()
            c.execute("""INSERT INTO project_collaborators (id, project_id, team_id, permission, added_at)
                         VALUES (?, ?, ?, ?, ?)""",
                      (collab_id, project_id, team_id, permission, added_at))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Project shared successfully"}), 200
        
    except Exception as e:
        print(f"Error sharing project: {e}")
        return jsonify({"detail": "Failed to share project"}), 500


@app.route('/api/v1/teams/<team_id>/projects/<project_id>', methods=['DELETE', 'OPTIONS'])
def unshare_project(team_id, project_id):
    """Remove project sharing with a team"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Check project ownership
        c.execute("SELECT user_id FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        
        if not project or project[0] != user_id:
            conn.close()
            return jsonify({"detail": "You can only manage your own projects"}), 403
        
        c.execute("DELETE FROM project_collaborators WHERE project_id = ? AND team_id = ?",
                  (project_id, team_id))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Project unshared successfully"}), 200
        
    except Exception as e:
        print(f"Error unsharing project: {e}")
        return jsonify({"detail": "Failed to unshare project"}), 500


@app.route('/api/v1/my-invites', methods=['GET', 'OPTIONS'])
def get_my_invites():
    """Get pending team invites for current user"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get user email
        c.execute("SELECT email FROM users WHERE id = ?", (user_id,))
        user = c.fetchone()
        
        if not user:
            conn.close()
            return jsonify({"detail": "User not found"}), 404
        
        # Get pending invites
        c.execute("""
            SELECT ti.id, t.id, t.name, u.full_name, ti.created_at
            FROM team_invites ti
            JOIN teams t ON ti.team_id = t.id
            JOIN users u ON ti.invited_by = u.id
            WHERE ti.email = ? AND ti.status = 'pending'
        """, (user[0],))
        
        invites = []
        for row in c.fetchall():
            invites.append({
                "invite_id": row[0],
                "team_id": row[1],
                "team_name": row[2],
                "invited_by": row[3],
                "created_at": row[4]
            })
        
        conn.close()
        return jsonify(invites), 200
        
    except Exception as e:
        print(f"Error getting invites: {e}")
        return jsonify({"detail": "Failed to get invites"}), 500


@app.route('/api/v1/invites/<invite_id>/accept', methods=['POST', 'OPTIONS'])
def accept_invite(invite_id):
    """Accept a team invite"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get user email
        c.execute("SELECT email FROM users WHERE id = ?", (user_id,))
        user = c.fetchone()
        
        # Get invite
        c.execute("SELECT team_id, email FROM team_invites WHERE id = ? AND status = 'pending'",
                  (invite_id,))
        invite = c.fetchone()
        
        if not invite or invite[1] != user[0]:
            conn.close()
            return jsonify({"detail": "Invalid invite"}), 404
        
        # Add as member
        member_id = str(uuid.uuid4())
        joined_at = datetime.utcnow().isoformat()
        c.execute("""INSERT INTO team_members (id, team_id, user_id, role, joined_at)
                     VALUES (?, ?, ?, 'member', ?)""",
                  (member_id, invite[0], user_id, joined_at))
        
        # Update invite status
        c.execute("UPDATE team_invites SET status = 'accepted' WHERE id = ?", (invite_id,))
        
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Invite accepted"}), 200
        
    except Exception as e:
        print(f"Error accepting invite: {e}")
        return jsonify({"detail": "Failed to accept invite"}), 500


@app.route('/api/v1/invites/<invite_id>/decline', methods=['POST', 'OPTIONS'])
def decline_invite(invite_id):
    """Decline a team invite"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get user email
        c.execute("SELECT email FROM users WHERE id = ?", (user_id,))
        user = c.fetchone()
        
        # Get invite
        c.execute("SELECT email FROM team_invites WHERE id = ? AND status = 'pending'", (invite_id,))
        invite = c.fetchone()
        
        if not invite or invite[0] != user[0]:
            conn.close()
            return jsonify({"detail": "Invalid invite"}), 404
        
        c.execute("UPDATE team_invites SET status = 'declined' WHERE id = ?", (invite_id,))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Invite declined"}), 200
        
    except Exception as e:
        print(f"Error declining invite: {e}")
        return jsonify({"detail": "Failed to decline invite"}), 500


@app.route('/api/v1/shared-projects', methods=['GET', 'OPTIONS'])
def get_shared_projects():
    """Get projects shared with the user through teams"""
    if request.method == 'OPTIONS':
        return '', 200
    
    user_id, error_response, status = get_user_from_token()
    if error_response:
        return error_response, status
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get projects shared through teams user is a member of
        c.execute("""
            SELECT DISTINCT p.id, p.name, p.description, p.status, p.gwp_total, p.mci_score,
                   pc.permission, t.name as team_name, u.full_name as owner_name
            FROM project_collaborators pc
            JOIN projects p ON pc.project_id = p.id
            JOIN teams t ON pc.team_id = t.id
            JOIN users u ON p.user_id = u.id
            LEFT JOIN team_members tm ON pc.team_id = tm.team_id
            WHERE t.owner_id = ? OR tm.user_id = ?
        """, (user_id, user_id))
        
        projects = []
        for row in c.fetchall():
            projects.append({
                "id": row[0],
                "name": row[1],
                "description": row[2],
                "status": row[3],
                "gwp_total": row[4],
                "mci_score": row[5],
                "permission": row[6],
                "team_name": row[7],
                "owner_name": row[8]
            })
        
        conn.close()
        return jsonify(projects), 200
        
    except Exception as e:
        print(f"Error getting shared projects: {e}")
        return jsonify({"detail": "Failed to get shared projects"}), 500


@app.route('/api/v1/projects', methods=['GET', 'OPTIONS'])
def list_projects():
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("""SELECT id, name, description, status, product_category, 
                     target_lifespan, is_designed_for_disassembly, user_id, 
                     COALESCE(gwp_total, 0) as gwp_total, 
                     COALESCE(mci_score, 0) as mci_score,
                     created_at 
                     FROM projects WHERE user_id = ? ORDER BY created_at DESC""", (payload['user_id'],))
        projects = c.fetchall()
        conn.close()
        
        return jsonify([{
            "id": p[0],
            "name": p[1],
            "description": p[2],
            "status": p[3],
            "product_category": p[4],
            "target_lifespan": p[5],
            "is_designed_for_disassembly": bool(p[6]) if p[6] is not None else False,
            "user_id": p[7],
            "gwp_total": float(p[8]) if p[8] is not None else 0.0,
            "mci_score": float(p[9]) if p[9] is not None else 0.0,
            "created_at": p[10]
        } for p in projects]), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects', methods=['POST'])
def create_project():
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        name = data.get('name')
        if not name:
            return jsonify({"detail": "Project name required"}), 400
        
        project_id = str(uuid.uuid4())
        description = data.get('description', '')
        status = 'draft'
        product_category = data.get('product_category', '')
        target_lifespan = data.get('target_lifespan')
        is_disassembly = data.get('is_designed_for_disassembly', False)
        created_at = datetime.utcnow().isoformat()
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("""INSERT INTO projects VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                  (project_id, name, description, status, product_category, 
                   target_lifespan, 1 if is_disassembly else 0, payload['user_id'], 
                   0, 0, 0, created_at))  # gwp_total=0, mci_score=0, circular_design_score=0
        conn.commit()
        conn.close()
        
        return jsonify({
            "id": project_id,
            "name": name,
            "description": description,
            "status": status,
            "product_category": product_category,
            "target_lifespan": target_lifespan,
            "is_designed_for_disassembly": is_disassembly,
            "created_at": created_at
        }), 201
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>', methods=['GET', 'OPTIONS'])
def get_project(project_id):
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("""SELECT id, name, description, status, product_category, 
                     target_lifespan, is_designed_for_disassembly, user_id, 
                     COALESCE(gwp_total, 0) as gwp_total, created_at 
                     FROM projects WHERE id = ? AND user_id = ?""", (project_id, payload['user_id']))
        project = c.fetchone()
        conn.close()
        
        if not project:
            return jsonify({"detail": "Project not found"}), 404
        
        return jsonify({
            "id": project[0],
            "name": project[1],
            "description": project[2],
            "status": project[3],
            "product_category": project[4],
            "target_lifespan": project[5],
            "is_designed_for_disassembly": bool(project[6]) if project[6] is not None else False,
            "user_id": project[7],
            "gwp_total": float(project[8]) if project[8] is not None else 0.0,
            "created_at": project[9]
        }), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>', methods=['PUT', 'OPTIONS'])
def update_project(project_id):
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Build dynamic update query
        updates = []
        values = []
        
        if 'name' in data:
            updates.append("name = ?")
            values.append(data['name'])
        if 'description' in data:
            updates.append("description = ?")
            values.append(data['description'])
        if 'status' in data:
            updates.append("status = ?")
            values.append(data['status'])
        if 'product_category' in data:
            updates.append("product_category = ?")
            values.append(data['product_category'])
        if 'target_lifespan' in data:
            updates.append("target_lifespan = ?")
            values.append(data['target_lifespan'])
        
        if not updates:
            return jsonify({"detail": "No fields to update"}), 400
        
        values.extend([project_id, payload['user_id']])
        query = f"UPDATE projects SET {', '.join(updates)} WHERE id = ? AND user_id = ?"
        c.execute(query, values)
        conn.commit()
        
        if c.rowcount == 0:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Fetch updated project
        c.execute("""SELECT id, name, description, status, product_category, 
                     target_lifespan, is_designed_for_disassembly, user_id, 
                     COALESCE(gwp_total, 0) as gwp_total, created_at 
                     FROM projects WHERE id = ?""", (project_id,))
        project = c.fetchone()
        conn.close()
        
        return jsonify({
            "id": project[0],
            "name": project[1],
            "description": project[2],
            "status": project[3],
            "product_category": project[4],
            "target_lifespan": project[5],
            "is_designed_for_disassembly": bool(project[6]) if project[6] is not None else False,
            "gwp_total": float(project[8]) if project[8] is not None else 0.0,
            "created_at": project[9]
        }), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>', methods=['DELETE', 'OPTIONS'])
def delete_project(project_id):
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Delete materials first
        c.execute("DELETE FROM project_materials WHERE project_id = ?", (project_id,))
        
        # Delete project
        c.execute("DELETE FROM projects WHERE id = ? AND user_id = ?", (project_id, payload['user_id']))
        conn.commit()
        
        if c.rowcount == 0:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        conn.close()
        return jsonify({"message": "Project deleted successfully"}), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# JNARRDC VERIFICATION WORKFLOW
# =====================================================

# In-memory storage for verification requests (would be database in production)
VERIFICATION_REQUESTS = {}

@app.route('/api/v1/projects/<project_id>/verification', methods=['GET', 'OPTIONS'])
def get_verification_status(project_id):
    """Get verification status for a project"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        verification = VERIFICATION_REQUESTS.get(project_id)
        
        if verification:
            return jsonify({
                'verification_status': verification.get('status', 'not_submitted'),
                'verification_submitted_at': verification.get('submitted_at'),
                'verification_reviewed_at': verification.get('verified_at'),
                'verifier_notes': verification.get('verifier_notes'),
                'certificate_id': verification.get('certificate_id'),
                'request_id': verification.get('request_id'),
                'flags': verification.get('flags', [])
            }), 200
        else:
            return jsonify({
                'verification_status': 'not_submitted',
                'verification_submitted_at': None,
                'verification_reviewed_at': None,
                'verifier_notes': None,
                'certificate_id': None
            }), 200
    except Exception as e:
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/verification/submit', methods=['POST', 'OPTIONS'])
def submit_for_verification(project_id):
    """Submit project for JNARRDC verification"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        user_id = payload['user_id']
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ? AND user_id = ?", (project_id, user_id))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("SELECT COUNT(*) FROM project_materials WHERE project_id = ?", (project_id,))
        material_count = c.fetchone()[0]
        
        # Get user info
        c.execute("SELECT full_name, organization_name, email FROM users WHERE id = ?", (user_id,))
        user = c.fetchone()
        conn.close()
        
        # Check if project has materials
        if material_count == 0:
            return jsonify({"detail": "Cannot submit empty project for verification. Add materials first."}), 400
        
        # Check if already submitted
        if project_id in VERIFICATION_REQUESTS and VERIFICATION_REQUESTS[project_id]['status'] != 'rejected':
            return jsonify({"detail": "Project already submitted for verification"}), 400
        
        # Create verification request
        request_id = f"VER-{project_id[:8].upper()}-{datetime.now().strftime('%Y%m%d%H%M')}"
        
        VERIFICATION_REQUESTS[project_id] = {
            'request_id': request_id,
            'project_id': project_id,
            'project_name': project[1],
            'user_id': user_id,
            'user_name': user[0] if user else 'Unknown',
            'organization': user[1] if user else 'Unknown',
            'status': 'pending',  # pending, under_review, verified, rejected
            'submitted_at': datetime.now().isoformat(),
            'material_count': material_count,
            'gwp_total': project[8] if len(project) > 8 else 0,
            'mci_score': project[10] if len(project) > 10 else 0,
            'verified_at': None,
            'verifier_name': None,
            'verifier_notes': None,
            'certificate_id': None,
            'flags': []
        }
        
        # Add AI-generated flags (simulated)
        flags = []
        if project[8] and project[8] > 1000:  # High GWP
            flags.append({'type': 'warning', 'message': 'High GWP value - manual review recommended'})
        if material_count > 20:
            flags.append({'type': 'info', 'message': 'Large BOM - detailed verification may take longer'})
        
        VERIFICATION_REQUESTS[project_id]['flags'] = flags
        
        return jsonify({
            "message": "Project submitted for JNARRDC verification",
            "request_id": request_id,
            "status": "pending",
            "estimated_review_time": "3-5 business days"
        }), 200
        
    except Exception as e:
        print(f"Verification Submit Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/verification/certificate', methods=['GET', 'OPTIONS'])
def get_verification_certificate(project_id):
    """Get verification certificate if project is verified"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        verification = VERIFICATION_REQUESTS.get(project_id)
        
        if not verification or verification['status'] != 'verified':
            return jsonify({"detail": "Project not verified. Certificate not available."}), 404
        
        # Generate certificate data
        certificate = {
            'certificate_id': verification['certificate_id'],
            'project_name': verification['project_name'],
            'organization': verification['organization'],
            'verified_at': verification['verified_at'],
            'verifier_name': verification['verifier_name'],
            'gwp_total': verification['gwp_total'],
            'mci_score': verification['mci_score'],
            'validity': '1 year from verification date',
            'issuer': 'JNARRDC - Ministry of Mines, Government of India',
            'qr_code_data': f"https://jnarrdc.gov.in/verify/{verification['certificate_id']}"
        }
        
        return jsonify(certificate), 200
        
    except Exception as e:
        return jsonify({"detail": str(e)}), 500


# Admin endpoint to simulate verification (for demo purposes)
@app.route('/api/v1/admin/verify-project/<project_id>', methods=['POST', 'OPTIONS'])
def admin_verify_project(project_id):
    """Admin endpoint to verify a project (demo only)"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json() or {}
        action = data.get('action', 'verify')  # verify or reject
        notes = data.get('notes', '')
        
        if project_id not in VERIFICATION_REQUESTS:
            return jsonify({"detail": "No verification request found"}), 404
        
        if action == 'verify':
            cert_id = f"JNARRDC-{datetime.now().year}-{uuid.uuid4().hex[:8].upper()}"
            VERIFICATION_REQUESTS[project_id].update({
                'status': 'verified',
                'verified_at': datetime.now().isoformat(),
                'verifier_name': 'JNARRDC Verification Team',
                'verifier_notes': notes or 'LCA data verified and approved.',
                'certificate_id': cert_id
            })
            return jsonify({"message": "Project verified", "certificate_id": cert_id}), 200
        else:
            VERIFICATION_REQUESTS[project_id].update({
                'status': 'rejected',
                'verified_at': datetime.now().isoformat(),
                'verifier_name': 'JNARRDC Verification Team',
                'verifier_notes': notes or 'Verification rejected. Please review and resubmit.'
            })
            return jsonify({"message": "Project verification rejected"}), 200
            
    except Exception as e:
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/materials', methods=['GET', 'OPTIONS'])
def list_materials(project_id):
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("SELECT * FROM project_materials WHERE project_id = ?", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        return jsonify([{
            "id": m[0],
            "project_id": m[1],
            "material_name": m[2],
            "material_type": m[3],
            "quantity": m[4],
            "unit": m[5],
            "recycled_content": m[6],
            "gwp": m[7],
            "transport_distance": m[8],
            "created_at": m[9]
        } for m in materials]), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>/materials', methods=['POST'])
def add_material(project_id):
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        material_id = str(uuid.uuid4())
        
        # Calculate GWP based on material type and quantity
        gwp = calculate_gwp(
            data.get('material_type'),
            data.get('quantity', 0),
            data.get('recycled_content', 0),
            data.get('transport_distance', 0)
        )
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        c.execute("""INSERT INTO project_materials VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                  (material_id, project_id, data.get('material_name'), 
                   data.get('material_type'), data.get('quantity'), data.get('unit'),
                   data.get('recycled_content', 0), gwp, data.get('transport_distance', 0),
                   datetime.utcnow().isoformat()))
        conn.commit()
        
        # Update project total GWP and MCI score
        c.execute("SELECT SUM(gwp) FROM project_materials WHERE project_id = ?", (project_id,))
        total_gwp = c.fetchone()[0] or 0
        
        # Calculate MCI score based on all materials
        c.execute("""SELECT material_type, quantity, recycled_content 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        all_materials = c.fetchall()
        
        # Get project info for MCI calculation
        c.execute("""SELECT target_lifespan, is_designed_for_disassembly, product_category 
                     FROM projects WHERE id = ?""", (project_id,))
        project_info = c.fetchone()
        target_lifespan = project_info[0] or 10
        is_designed_for_disassembly = bool(project_info[1])
        product_category = project_info[2] or 'other'
        
        # Calculate weighted average recycled content
        total_mass = sum(m[1] or 0 for m in all_materials)
        weighted_recycled = sum((m[1] or 0) * (m[2] or 0) for m in all_materials) / total_mass if total_mass > 0 else 0
        
        # Get industry benchmark
        benchmark = INDUSTRY_BENCHMARKS.get(product_category, INDUSTRY_BENCHMARKS['other'])
        industry_avg_lifespan = benchmark['avg_lifespan']
        
        # Calculate MCI
        mci_score = calculate_mci(
            weighted_recycled,  # recycled input
            weighted_recycled * 0.8,  # assume 80% of input can be recycled at end of life
            target_lifespan,
            industry_avg_lifespan,
            is_designed_for_disassembly
        )
        
        c.execute("UPDATE projects SET status = ?, gwp_total = ?, mci_score = ? WHERE id = ?", 
                  ('calculated', total_gwp, mci_score, project_id))
        conn.commit()
        conn.close()
        
        return jsonify({
            "id": material_id,
            "project_id": project_id,
            "material_name": data.get('material_name'),
            "material_type": data.get('material_type'),
            "quantity": data.get('quantity'),
            "unit": data.get('unit'),
            "recycled_content": data.get('recycled_content', 0),
            "gwp": gwp,
            "transport_distance": data.get('transport_distance', 0)
        }), 201
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>/materials/batch', methods=['POST'])
def add_materials_batch(project_id):
    """Batch add multiple materials at once - for BOM upload"""
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        materials = data.get('materials', [])
        
        if not materials:
            return jsonify({"detail": "No materials provided"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        added_materials = []
        failed_materials = []
        
        for item in materials:
            try:
                material_id = str(uuid.uuid4())
                
                # Calculate GWP
                gwp = calculate_gwp(
                    item.get('material_type'),
                    item.get('quantity', 0),
                    item.get('recycled_content', 0),
                    item.get('transport_distance', 0)
                )
                
                c.execute("""INSERT INTO project_materials VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                          (material_id, project_id, item.get('material_name'), 
                           item.get('material_type'), item.get('quantity'), item.get('unit', 'kg'),
                           item.get('recycled_content', 0), gwp, item.get('transport_distance', 0),
                           datetime.utcnow().isoformat()))
                
                added_materials.append({
                    "id": material_id,
                    "material_name": item.get('material_name'),
                    "gwp": gwp
                })
            except Exception as e:
                failed_materials.append({
                    "material_name": item.get('material_name'),
                    "error": str(e)
                })
        
        conn.commit()
        
        # Update project total GWP and MCI score
        c.execute("SELECT SUM(gwp) FROM project_materials WHERE project_id = ?", (project_id,))
        total_gwp = c.fetchone()[0] or 0
        
        # Calculate MCI score based on all materials
        c.execute("""SELECT material_type, quantity, recycled_content 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        all_materials = c.fetchall()
        
        # Get project info for MCI calculation
        c.execute("""SELECT target_lifespan, is_designed_for_disassembly, product_category 
                     FROM projects WHERE id = ?""", (project_id,))
        project_info = c.fetchone()
        target_lifespan = project_info[0] or 10
        is_designed_for_disassembly = bool(project_info[1])
        product_category = project_info[2] or 'other'
        
        # Calculate weighted average recycled content
        total_mass = sum(m[1] or 0 for m in all_materials)
        weighted_recycled = sum((m[1] or 0) * (m[2] or 0) for m in all_materials) / total_mass if total_mass > 0 else 0
        
        # Get industry benchmark
        benchmark = INDUSTRY_BENCHMARKS.get(product_category, INDUSTRY_BENCHMARKS['other'])
        industry_avg_lifespan = benchmark['avg_lifespan']
        
        # Calculate MCI
        mci_score = calculate_mci(
            weighted_recycled,
            weighted_recycled * 0.8,
            target_lifespan,
            industry_avg_lifespan,
            is_designed_for_disassembly
        )
        
        c.execute("UPDATE projects SET status = ?, gwp_total = ?, mci_score = ? WHERE id = ?", 
                  ('calculated', total_gwp, mci_score, project_id))
        conn.commit()
        conn.close()
        
        return jsonify({
            "added": len(added_materials),
            "failed": len(failed_materials),
            "materials": added_materials,
            "errors": failed_materials,
            "total_gwp": total_gwp,
            "mci_score": mci_score
        }), 201
    except Exception as e:
        print(f"Error in batch add: {e}")
        return jsonify({"detail": str(e)}), 500

@app.route('/api/v1/projects/<project_id>/materials/<material_id>', methods=['DELETE', 'OPTIONS'])
def delete_material(project_id, material_id):
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Delete material
        c.execute("DELETE FROM project_materials WHERE id = ? AND project_id = ?", (material_id, project_id))
        conn.commit()
        
        if c.rowcount == 0:
            conn.close()
            return jsonify({"detail": "Material not found"}), 404
        
        # Recalculate project total GWP
        c.execute("SELECT SUM(gwp) FROM project_materials WHERE project_id = ?", (project_id,))
        total_gwp = c.fetchone()[0] or 0
        
        # Update project status
        c.execute("SELECT COUNT(*) FROM project_materials WHERE project_id = ?", (project_id,))
        material_count = c.fetchone()[0]
        status = 'calculated' if material_count > 0 else 'draft'
        
        c.execute("UPDATE projects SET status = ?, gwp_total = ? WHERE id = ?", 
                  (status, total_gwp, project_id))
        conn.commit()
        conn.close()
        
        return jsonify({"message": "Material deleted successfully", "new_total_gwp": total_gwp}), 200
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"detail": str(e)}), 500

def calculate_gwp(material_type, quantity, recycled_content, transport_distance):
    """Calculate GWP based on material type and parameters"""
    
    # Emission factors (kg CO2-eq per kg material) - includes critical minerals
    emission_factors = {
        # Base Metals
        'aluminium_primary': 12.5,
        'aluminium_secondary': 0.6,
        'copper_primary': 3.5,
        'copper_secondary': 0.5,
        'steel_primary': 2.1,
        'steel_secondary': 0.4,
        # Battery Minerals
        'lithium': 15.0,
        'lithium_carbonate': 15.0,
        'lithium_hydroxide': 18.0,
        'cobalt': 10.0,
        'cobalt_sulfate': 10.0,
        'nickel': 8.5,
        'nickel_class1': 12.5,
        'nickel_ferronickel': 8.5,
        'manganese': 2.8,
        'graphite': 4.2,
        # Rare Earths
        'neodymium': 35.0,
        'dysprosium': 45.0,
        'praseodymium': 32.0,
        'terbium': 50.0,
        'rare_earth_mixed': 38.0,
        # Other Critical Minerals
        'tungsten': 22.0,
        'vanadium': 28.0,
        'titanium': 8.1,
        'tantalum': 48.0,
        'indium': 142.0,
        'gallium': 185.0,
        'germanium': 165.0,
        # Precious Metals
        'platinum': 12500.0,
        'palladium': 9800.0,
        'silver': 104.0,
        'gold': 31500.0,
        # Joining/Brazing Materials
        'solder_lead_free': 25.0,
        'brazing_alloy': 85.0,  # Silver-based brazing alloy
        'flux': 3.0,
    }
    
    # Get base emission factor
    base_ef = emission_factors.get(material_type, 5.0)  # Default if not found
    
    # Calculate virgin and recycled portions
    virgin_fraction = (100 - recycled_content) / 100
    recycled_fraction = recycled_content / 100
    
    # Calculate material emissions
    virgin_emissions = quantity * base_ef * virgin_fraction
    
    # Recycled material has ~90% lower emissions
    recycled_emissions = quantity * (base_ef * 0.1) * recycled_fraction
    
    # Transport emissions (kg CO2-eq per ton-km)
    transport_emissions = (quantity / 1000) * transport_distance * 0.062
    
    # Total GWP
    total_gwp = virgin_emissions + recycled_emissions + transport_emissions
    
    return round(total_gwp, 2)


def get_scarcity_score(material_type):
    """Get Abiotic Depletion Potential (scarcity) score for critical minerals"""
    scarcity_scores = {
        # Battery Minerals
        'lithium': 85, 'lithium_carbonate': 85, 'lithium_hydroxide': 85,
        'cobalt': 92, 'cobalt_sulfate': 92,
        'nickel': 65, 'nickel_class1': 65, 'nickel_ferronickel': 55,
        'manganese': 40,
        'graphite': 70,
        # Rare Earths
        'neodymium': 95,
        'dysprosium': 98,
        'praseodymium': 90,
        'terbium': 96,
        'rare_earth_mixed': 93,
        # Other Critical
        'tungsten': 78,
        'vanadium': 72,
        'titanium': 45,
        'tantalum': 88,
        'indium': 82,
        'gallium': 80,
        'germanium': 84,
        # Precious Metals
        'platinum': 88,
        'palladium': 86,
        'silver': 60,
        'gold': 95,
        # Base metals (lower scarcity)
        'aluminium_primary': 20, 'aluminium_secondary': 20,
        'copper_primary': 35, 'copper_secondary': 35,
        'steel_primary': 15, 'steel_secondary': 15,
        # Joining materials (low scarcity)
        'solder_lead_free': 25,
        'brazing_alloy': 50,
        'flux': 5,
    }
    return scarcity_scores.get(material_type, 0)


# Industry benchmark lifespans (in years)
INDUSTRY_BENCHMARKS = {
    'batteries': {'avg_lifespan': 8, 'avg_mci': 0.35, 'name': 'Battery Pack'},
    'ev_battery': {'avg_lifespan': 10, 'avg_mci': 0.32, 'name': 'EV Battery Pack'},
    'electronics': {'avg_lifespan': 5, 'avg_mci': 0.28, 'name': 'Electronics'},
    'automotive': {'avg_lifespan': 15, 'avg_mci': 0.45, 'name': 'Automotive Component'},
    'industrial': {'avg_lifespan': 20, 'avg_mci': 0.40, 'name': 'Industrial Equipment'},
    'construction': {'avg_lifespan': 50, 'avg_mci': 0.52, 'name': 'Construction Material'},
    'packaging': {'avg_lifespan': 1, 'avg_mci': 0.65, 'name': 'Packaging'},
    'renewable_energy': {'avg_lifespan': 25, 'avg_mci': 0.48, 'name': 'Renewable Energy'},
    'magnets': {'avg_lifespan': 20, 'avg_mci': 0.25, 'name': 'Permanent Magnets'},
    'other': {'avg_lifespan': 10, 'avg_mci': 0.35, 'name': 'General Product'}
}

# Emission factors for virgin and recycled materials (kg CO2-eq per kg)
EMISSION_FACTORS = {
    # Base Metals
    'aluminium_primary': {'virgin': 12.5, 'recycled': 0.6},
    'aluminium_secondary': {'virgin': 0.6, 'recycled': 0.6},
    'copper_primary': {'virgin': 3.5, 'recycled': 0.5},
    'copper_secondary': {'virgin': 0.5, 'recycled': 0.5},
    'steel_primary': {'virgin': 2.1, 'recycled': 0.4},
    'steel_secondary': {'virgin': 0.4, 'recycled': 0.4},
    
    # Critical Minerals - Battery Metals
    'lithium_carbonate': {'virgin': 15.0, 'recycled': 1.5, 'scarcity_score': 85},
    'lithium_hydroxide': {'virgin': 18.0, 'recycled': 1.8, 'scarcity_score': 85},
    'lithium': {'virgin': 15.0, 'recycled': 1.5, 'scarcity_score': 85},
    'cobalt_sulfate': {'virgin': 10.0, 'recycled': 1.0, 'scarcity_score': 92},
    'cobalt': {'virgin': 10.0, 'recycled': 1.0, 'scarcity_score': 92},
    'nickel_class1': {'virgin': 12.5, 'recycled': 1.25, 'scarcity_score': 65},
    'nickel_ferronickel': {'virgin': 8.5, 'recycled': 0.85, 'scarcity_score': 55},
    'nickel': {'virgin': 8.5, 'recycled': 0.85, 'scarcity_score': 65},
    'manganese': {'virgin': 2.8, 'recycled': 0.3, 'scarcity_score': 40},
    'graphite': {'virgin': 4.2, 'recycled': 0.5, 'scarcity_score': 70},
    
    # Rare Earths
    'neodymium': {'virgin': 35.0, 'recycled': 5.0, 'scarcity_score': 95},
    'dysprosium': {'virgin': 45.0, 'recycled': 6.0, 'scarcity_score': 98},
    'praseodymium': {'virgin': 32.0, 'recycled': 4.5, 'scarcity_score': 90},
    'terbium': {'virgin': 50.0, 'recycled': 7.0, 'scarcity_score': 96},
    'rare_earth_mixed': {'virgin': 38.0, 'recycled': 5.5, 'scarcity_score': 93},
    
    # Other Critical Minerals
    'tungsten': {'virgin': 22.0, 'recycled': 3.0, 'scarcity_score': 78},
    'vanadium': {'virgin': 28.0, 'recycled': 4.0, 'scarcity_score': 72},
    'titanium': {'virgin': 8.1, 'recycled': 1.2, 'scarcity_score': 45},
    'platinum': {'virgin': 12500, 'recycled': 800, 'scarcity_score': 88},
    'palladium': {'virgin': 9800, 'recycled': 650, 'scarcity_score': 86},
    'indium': {'virgin': 142, 'recycled': 15, 'scarcity_score': 82},
    'gallium': {'virgin': 185, 'recycled': 20, 'scarcity_score': 80},
    'germanium': {'virgin': 165, 'recycled': 18, 'scarcity_score': 84},
    'tantalum': {'virgin': 48.0, 'recycled': 6.5, 'scarcity_score': 88},
    # Precious Metals
    'silver': {'virgin': 104, 'recycled': 8.0, 'scarcity_score': 60},
    'gold': {'virgin': 31500, 'recycled': 1200, 'scarcity_score': 95},
    # Joining/Brazing Materials
    'solder_lead_free': {'virgin': 25.0, 'recycled': 4.0, 'scarcity_score': 25},
    'brazing_alloy': {'virgin': 85.0, 'recycled': 12.0, 'scarcity_score': 50},
    'flux': {'virgin': 3.0, 'recycled': 3.0, 'scarcity_score': 5},
}

# Waste generation factors per material type (percentage of material input that becomes waste)
# Based on industry averages for metal processing
WASTE_FACTORS = {
    # Base Metals - lower waste due to established recycling infrastructure
    'aluminium': 0.03,           # 3% - well-established scrap recycling
    'aluminium_primary': 0.04,   # 4% - primary production has more losses
    'aluminium_secondary': 0.02, # 2% - secondary already optimized
    'copper': 0.03,              # 3% - high value, well recovered
    'copper_primary': 0.04,
    'copper_secondary': 0.02,
    'steel': 0.04,               # 4% - some scale/slag losses
    'steel_primary': 0.05,
    'steel_secondary': 0.03,
    'iron': 0.05,                # 5% - more slag in processing
    
    # Critical Minerals - Battery Metals - higher waste due to complex processing
    'lithium': 0.08,             # 8% - complex extraction process
    'lithium_carbonate': 0.08,
    'lithium_hydroxide': 0.08,
    'cobalt': 0.07,              # 7% - refining losses
    'cobalt_sulfate': 0.07,
    'nickel': 0.06,              # 6% - smelting losses
    'nickel_class1': 0.05,
    'nickel_ferronickel': 0.07,
    'manganese': 0.06,           # 6% - processing losses
    'graphite': 0.10,            # 10% - high waste in processing
    
    # Rare Earths - highest waste due to complex separation
    'neodymium': 0.12,           # 12% - separation losses
    'dysprosium': 0.15,          # 15% - very complex separation
    'praseodymium': 0.12,
    'terbium': 0.15,
    'rare_earth_mixed': 0.10,    # 10% - before separation
    
    # Other Critical Minerals
    'tungsten': 0.08,            # 8% - hard to process
    'vanadium': 0.09,            # 9% - extraction losses
    'titanium': 0.07,            # 7% - machining waste
    'platinum': 0.02,            # 2% - very high recovery due to value
    'palladium': 0.02,
    'indium': 0.12,              # 12% - byproduct recovery
    'gallium': 0.12,
    'germanium': 0.12,
    'tantalum': 0.10,            # 10% - complex extraction
    
    # Precious Metals - very low waste due to high value
    'silver': 0.02,              # 2%
    'gold': 0.01,                # 1% - extremely high recovery
    
    # Joining/Brazing Materials
    'solder_lead_free': 0.05,    # 5% - application losses
    'brazing_alloy': 0.06,
    'flux': 0.15,                # 15% - consumable
    
    # Default
    'default': 0.05              # 5% default for unknown materials
}


# =====================================================
# LCIA IMPACT CATEGORIES - Life Cycle Impact Assessment
# Based on ReCiPe 2016, CML 2001, and TRACI 2.1 methodologies
# =====================================================

# LCIA characterization factors per material type
# Units: Various (see individual category definitions)
LCIA_IMPACT_FACTORS = {
    # =========================================================================
    # BASE METALS
    # =========================================================================
    'aluminium_primary': {
        'gwp': 12.5,           # Global Warming Potential (kg CO2-eq/kg)
        'ap': 0.052,           # Acidification Potential (kg SO2-eq/kg)
        'ep': 0.0042,          # Eutrophication Potential (kg PO4-eq/kg)
        'odp': 2.1e-8,         # Ozone Depletion Potential (kg CFC-11-eq/kg)
        'pocp': 0.0031,        # Photochemical Ozone Creation (kg C2H4-eq/kg)
        'htp': 1.8,            # Human Toxicity Potential (kg 1,4-DCB-eq/kg)
        'faetp': 0.45,         # Freshwater Aquatic Ecotoxicity (kg 1,4-DCB-eq/kg)
        'tetp': 0.012,         # Terrestrial Ecotoxicity (kg 1,4-DCB-eq/kg)
        'adp_elements': 1.1e-5, # Abiotic Depletion - Elements (kg Sb-eq/kg)
        'adp_fossil': 165.0,   # Abiotic Depletion - Fossil (MJ/kg)
        'water_use': 85.0,     # Water Use (L/kg)
        'land_use': 0.8,       # Land Use (m2a/kg)
        # Process-level energy breakdown (kWh/kg)
        'energy_mining': 0.5,
        'energy_refining': 8.5,    # Bayer process
        'energy_smelting': 15.0,   # Hall-Héroult process
        'energy_total': 24.0,
    },
    'aluminium_secondary': {
        'gwp': 0.6,
        'ap': 0.003,
        'ep': 0.0003,
        'odp': 1.2e-9,
        'pocp': 0.0002,
        'htp': 0.12,
        'faetp': 0.03,
        'tetp': 0.001,
        'adp_elements': 5.5e-7,
        'adp_fossil': 8.5,
        'water_use': 5.0,
        'land_use': 0.02,
        'energy_mining': 0.0,
        'energy_refining': 0.0,
        'energy_smelting': 0.8,    # Remelting only
        'energy_total': 0.8,
    },
    'copper_primary': {
        'gwp': 3.5,
        'ap': 0.18,            # High due to SO2 from smelting
        'ep': 0.0025,
        'odp': 8.5e-9,
        'pocp': 0.0018,
        'htp': 45.0,           # High due to heavy metal emissions
        'faetp': 28.0,
        'tetp': 0.85,
        'adp_elements': 0.0042,
        'adp_fossil': 42.0,
        'water_use': 120.0,
        'land_use': 1.2,
        'energy_mining': 1.5,
        'energy_refining': 2.8,
        'energy_smelting': 3.2,
        'energy_total': 7.5,
    },
    'copper_secondary': {
        'gwp': 0.5,
        'ap': 0.008,
        'ep': 0.0002,
        'odp': 4.2e-10,
        'pocp': 0.0001,
        'htp': 2.5,
        'faetp': 1.2,
        'tetp': 0.04,
        'adp_elements': 2.1e-5,
        'adp_fossil': 6.0,
        'water_use': 8.0,
        'land_use': 0.03,
        'energy_mining': 0.0,
        'energy_refining': 0.5,
        'energy_smelting': 0.8,
        'energy_total': 1.3,
    },
    'steel_primary': {
        'gwp': 2.1,
        'ap': 0.008,
        'ep': 0.0012,
        'odp': 4.5e-9,
        'pocp': 0.0012,
        'htp': 0.95,
        'faetp': 0.22,
        'tetp': 0.008,
        'adp_elements': 2.8e-6,
        'adp_fossil': 28.0,
        'water_use': 45.0,
        'land_use': 0.35,
        'energy_mining': 0.8,
        'energy_refining': 1.2,
        'energy_smelting': 4.5,
        'energy_total': 6.5,
    },
    'steel_secondary': {
        'gwp': 0.4,
        'ap': 0.002,
        'ep': 0.0002,
        'odp': 2.2e-10,
        'pocp': 0.0002,
        'htp': 0.15,
        'faetp': 0.04,
        'tetp': 0.001,
        'adp_elements': 1.4e-7,
        'adp_fossil': 5.5,
        'water_use': 6.0,
        'land_use': 0.02,
        'energy_mining': 0.0,
        'energy_refining': 0.0,
        'energy_smelting': 0.6,
        'energy_total': 0.6,
    },
    
    # =========================================================================
    # CRITICAL MINERALS - BATTERY METALS
    # =========================================================================
    'lithium': {
        'gwp': 15.0,
        'ap': 0.045,
        'ep': 0.018,
        'odp': 1.8e-8,
        'pocp': 0.0028,
        'htp': 2.8,
        'faetp': 1.2,
        'tetp': 0.035,
        'adp_elements': 0.012,
        'adp_fossil': 185.0,
        'water_use': 1900.0,      # Very high - brine extraction
        'land_use': 2.5,
        'energy_mining': 2.5,
        'energy_refining': 12.0,
        'energy_smelting': 0.0,
        'energy_total': 14.5,
    },
    'lithium_carbonate': {
        'gwp': 15.0,
        'ap': 0.045,
        'ep': 0.018,
        'odp': 1.8e-8,
        'pocp': 0.0028,
        'htp': 2.8,
        'faetp': 1.2,
        'tetp': 0.035,
        'adp_elements': 0.012,
        'adp_fossil': 185.0,
        'water_use': 1900.0,
        'land_use': 2.5,
        'energy_mining': 2.5,
        'energy_refining': 12.0,
        'energy_smelting': 0.0,
        'energy_total': 14.5,
    },
    'lithium_hydroxide': {
        'gwp': 18.0,
        'ap': 0.055,
        'ep': 0.022,
        'odp': 2.2e-8,
        'pocp': 0.0035,
        'htp': 3.2,
        'faetp': 1.5,
        'tetp': 0.042,
        'adp_elements': 0.014,
        'adp_fossil': 220.0,
        'water_use': 2200.0,
        'land_use': 2.8,
        'energy_mining': 2.5,
        'energy_refining': 15.0,
        'energy_smelting': 0.0,
        'energy_total': 17.5,
    },
    'cobalt': {
        'gwp': 10.0,
        'ap': 0.28,            # High due to sulfide processing
        'ep': 0.025,
        'odp': 1.2e-8,
        'pocp': 0.0045,
        'htp': 85.0,           # Very high - toxic metal
        'faetp': 42.0,
        'tetp': 2.5,
        'adp_elements': 0.052,
        'adp_fossil': 125.0,
        'water_use': 350.0,
        'land_use': 3.5,
        'energy_mining': 3.5,
        'energy_refining': 8.5,
        'energy_smelting': 2.0,
        'energy_total': 14.0,
    },
    'cobalt_sulfate': {
        'gwp': 10.0,
        'ap': 0.28,
        'ep': 0.025,
        'odp': 1.2e-8,
        'pocp': 0.0045,
        'htp': 85.0,
        'faetp': 42.0,
        'tetp': 2.5,
        'adp_elements': 0.052,
        'adp_fossil': 125.0,
        'water_use': 350.0,
        'land_use': 3.5,
        'energy_mining': 3.5,
        'energy_refining': 8.5,
        'energy_smelting': 2.0,
        'energy_total': 14.0,
    },
    'nickel': {
        'gwp': 8.5,
        'ap': 0.15,
        'ep': 0.012,
        'odp': 8.5e-9,
        'pocp': 0.0032,
        'htp': 52.0,
        'faetp': 25.0,
        'tetp': 1.8,
        'adp_elements': 0.0085,
        'adp_fossil': 95.0,
        'water_use': 180.0,
        'land_use': 2.2,
        'energy_mining': 2.5,
        'energy_refining': 5.5,
        'energy_smelting': 4.5,
        'energy_total': 12.5,
    },
    'nickel_class1': {
        'gwp': 12.5,
        'ap': 0.22,
        'ep': 0.018,
        'odp': 1.2e-8,
        'pocp': 0.0048,
        'htp': 68.0,
        'faetp': 35.0,
        'tetp': 2.4,
        'adp_elements': 0.012,
        'adp_fossil': 145.0,
        'water_use': 250.0,
        'land_use': 2.8,
        'energy_mining': 2.5,
        'energy_refining': 8.5,
        'energy_smelting': 6.5,
        'energy_total': 17.5,
    },
    'nickel_ferronickel': {
        'gwp': 8.5,
        'ap': 0.15,
        'ep': 0.012,
        'odp': 8.5e-9,
        'pocp': 0.0032,
        'htp': 52.0,
        'faetp': 25.0,
        'tetp': 1.8,
        'adp_elements': 0.0085,
        'adp_fossil': 95.0,
        'water_use': 180.0,
        'land_use': 2.2,
        'energy_mining': 2.5,
        'energy_refining': 5.5,
        'energy_smelting': 4.5,
        'energy_total': 12.5,
    },
    'manganese': {
        'gwp': 2.8,
        'ap': 0.015,
        'ep': 0.0035,
        'odp': 3.2e-9,
        'pocp': 0.0012,
        'htp': 4.5,
        'faetp': 1.8,
        'tetp': 0.15,
        'adp_elements': 0.0012,
        'adp_fossil': 35.0,
        'water_use': 45.0,
        'land_use': 0.85,
        'energy_mining': 1.2,
        'energy_refining': 2.8,
        'energy_smelting': 1.5,
        'energy_total': 5.5,
    },
    'graphite': {
        'gwp': 4.2,
        'ap': 0.018,
        'ep': 0.0028,
        'odp': 4.5e-9,
        'pocp': 0.0015,
        'htp': 1.2,
        'faetp': 0.45,
        'tetp': 0.08,
        'adp_elements': 0.0008,
        'adp_fossil': 55.0,
        'water_use': 65.0,
        'land_use': 1.2,
        'energy_mining': 1.8,
        'energy_refining': 4.5,
        'energy_smelting': 0.0,
        'energy_total': 6.3,
    },
    
    # =========================================================================
    # RARE EARTH ELEMENTS
    # =========================================================================
    'neodymium': {
        'gwp': 35.0,
        'ap': 0.12,
        'ep': 0.085,           # High due to processing chemicals
        'odp': 3.5e-8,
        'pocp': 0.0085,
        'htp': 125.0,          # Very high - radioactive byproducts
        'faetp': 85.0,
        'tetp': 8.5,
        'adp_elements': 0.42,
        'adp_fossil': 420.0,
        'water_use': 850.0,
        'land_use': 12.0,
        'energy_mining': 8.5,
        'energy_refining': 35.0,
        'energy_smelting': 12.0,
        'energy_total': 55.5,
    },
    'dysprosium': {
        'gwp': 45.0,
        'ap': 0.15,
        'ep': 0.095,
        'odp': 4.5e-8,
        'pocp': 0.0095,
        'htp': 145.0,
        'faetp': 95.0,
        'tetp': 9.5,
        'adp_elements': 0.85,
        'adp_fossil': 550.0,
        'water_use': 950.0,
        'land_use': 14.0,
        'energy_mining': 10.0,
        'energy_refining': 45.0,
        'energy_smelting': 15.0,
        'energy_total': 70.0,
    },
    'praseodymium': {
        'gwp': 32.0,
        'ap': 0.11,
        'ep': 0.078,
        'odp': 3.2e-8,
        'pocp': 0.0078,
        'htp': 115.0,
        'faetp': 78.0,
        'tetp': 7.8,
        'adp_elements': 0.38,
        'adp_fossil': 385.0,
        'water_use': 780.0,
        'land_use': 11.0,
        'energy_mining': 7.8,
        'energy_refining': 32.0,
        'energy_smelting': 11.0,
        'energy_total': 50.8,
    },
    'terbium': {
        'gwp': 50.0,
        'ap': 0.18,
        'ep': 0.12,
        'odp': 5.0e-8,
        'pocp': 0.012,
        'htp': 165.0,
        'faetp': 105.0,
        'tetp': 10.5,
        'adp_elements': 1.2,
        'adp_fossil': 650.0,
        'water_use': 1050.0,
        'land_use': 16.0,
        'energy_mining': 12.0,
        'energy_refining': 55.0,
        'energy_smelting': 18.0,
        'energy_total': 85.0,
    },
    'rare_earth_mixed': {
        'gwp': 38.0,
        'ap': 0.13,
        'ep': 0.088,
        'odp': 3.8e-8,
        'pocp': 0.0088,
        'htp': 130.0,
        'faetp': 88.0,
        'tetp': 8.8,
        'adp_elements': 0.55,
        'adp_fossil': 450.0,
        'water_use': 880.0,
        'land_use': 13.0,
        'energy_mining': 9.0,
        'energy_refining': 38.0,
        'energy_smelting': 13.0,
        'energy_total': 60.0,
    },
    
    # =========================================================================
    # OTHER CRITICAL MINERALS
    # =========================================================================
    'tungsten': {
        'gwp': 22.0,
        'ap': 0.065,
        'ep': 0.025,
        'odp': 2.2e-8,
        'pocp': 0.0055,
        'htp': 35.0,
        'faetp': 18.0,
        'tetp': 1.5,
        'adp_elements': 0.18,
        'adp_fossil': 280.0,
        'water_use': 320.0,
        'land_use': 4.5,
        'energy_mining': 5.5,
        'energy_refining': 18.0,
        'energy_smelting': 8.5,
        'energy_total': 32.0,
    },
    'vanadium': {
        'gwp': 28.0,
        'ap': 0.085,
        'ep': 0.032,
        'odp': 2.8e-8,
        'pocp': 0.0068,
        'htp': 45.0,
        'faetp': 22.0,
        'tetp': 2.2,
        'adp_elements': 0.025,
        'adp_fossil': 350.0,
        'water_use': 280.0,
        'land_use': 3.8,
        'energy_mining': 4.5,
        'energy_refining': 22.0,
        'energy_smelting': 6.5,
        'energy_total': 33.0,
    },
    'titanium': {
        'gwp': 8.1,
        'ap': 0.035,
        'ep': 0.008,
        'odp': 8.1e-9,
        'pocp': 0.0025,
        'htp': 8.5,
        'faetp': 3.5,
        'tetp': 0.35,
        'adp_elements': 0.0015,
        'adp_fossil': 95.0,
        'water_use': 120.0,
        'land_use': 1.8,
        'energy_mining': 2.5,
        'energy_refining': 8.5,
        'energy_smelting': 12.0,
        'energy_total': 23.0,
    },
    'tantalum': {
        'gwp': 48.0,
        'ap': 0.14,
        'ep': 0.055,
        'odp': 4.8e-8,
        'pocp': 0.012,
        'htp': 95.0,
        'faetp': 55.0,
        'tetp': 5.5,
        'adp_elements': 0.65,
        'adp_fossil': 580.0,
        'water_use': 450.0,
        'land_use': 8.5,
        'energy_mining': 8.0,
        'energy_refining': 35.0,
        'energy_smelting': 15.0,
        'energy_total': 58.0,
    },
    'indium': {
        'gwp': 142.0,
        'ap': 0.42,
        'ep': 0.15,
        'odp': 1.4e-7,
        'pocp': 0.035,
        'htp': 280.0,
        'faetp': 145.0,
        'tetp': 14.5,
        'adp_elements': 2.5,
        'adp_fossil': 1850.0,
        'water_use': 1200.0,
        'land_use': 25.0,
        'energy_mining': 25.0,
        'energy_refining': 120.0,
        'energy_smelting': 35.0,
        'energy_total': 180.0,
    },
    'gallium': {
        'gwp': 185.0,
        'ap': 0.55,
        'ep': 0.18,
        'odp': 1.8e-7,
        'pocp': 0.045,
        'htp': 320.0,
        'faetp': 165.0,
        'tetp': 16.5,
        'adp_elements': 3.2,
        'adp_fossil': 2400.0,
        'water_use': 1500.0,
        'land_use': 32.0,
        'energy_mining': 32.0,
        'energy_refining': 155.0,
        'energy_smelting': 45.0,
        'energy_total': 232.0,
    },
    'germanium': {
        'gwp': 165.0,
        'ap': 0.48,
        'ep': 0.16,
        'odp': 1.6e-7,
        'pocp': 0.042,
        'htp': 295.0,
        'faetp': 155.0,
        'tetp': 15.5,
        'adp_elements': 2.8,
        'adp_fossil': 2150.0,
        'water_use': 1350.0,
        'land_use': 28.0,
        'energy_mining': 28.0,
        'energy_refining': 140.0,
        'energy_smelting': 40.0,
        'energy_total': 208.0,
    },
    
    # =========================================================================
    # PRECIOUS METALS
    # =========================================================================
    'platinum': {
        'gwp': 12500.0,
        'ap': 38.0,
        'ep': 12.0,
        'odp': 1.2e-5,
        'pocp': 3.2,
        'htp': 25000.0,
        'faetp': 12500.0,
        'tetp': 1250.0,
        'adp_elements': 185.0,
        'adp_fossil': 165000.0,
        'water_use': 285000.0,
        'land_use': 2500.0,
        'energy_mining': 2500.0,
        'energy_refining': 8500.0,
        'energy_smelting': 2500.0,
        'energy_total': 13500.0,
    },
    'palladium': {
        'gwp': 9800.0,
        'ap': 32.0,
        'ep': 9.8,
        'odp': 9.8e-6,
        'pocp': 2.8,
        'htp': 19500.0,
        'faetp': 9800.0,
        'tetp': 980.0,
        'adp_elements': 145.0,
        'adp_fossil': 128000.0,
        'water_use': 225000.0,
        'land_use': 1950.0,
        'energy_mining': 1950.0,
        'energy_refining': 6500.0,
        'energy_smelting': 1950.0,
        'energy_total': 10400.0,
    },
    'silver': {
        'gwp': 104.0,
        'ap': 0.35,
        'ep': 0.11,
        'odp': 1.0e-7,
        'pocp': 0.028,
        'htp': 210.0,
        'faetp': 105.0,
        'tetp': 10.5,
        'adp_elements': 1.8,
        'adp_fossil': 1350.0,
        'water_use': 950.0,
        'land_use': 18.0,
        'energy_mining': 18.0,
        'energy_refining': 85.0,
        'energy_smelting': 25.0,
        'energy_total': 128.0,
    },
    'gold': {
        'gwp': 31500.0,
        'ap': 95.0,
        'ep': 32.0,
        'odp': 3.2e-5,
        'pocp': 8.5,
        'htp': 65000.0,
        'faetp': 32000.0,
        'tetp': 3200.0,
        'adp_elements': 520.0,
        'adp_fossil': 420000.0,
        'water_use': 750000.0,
        'land_use': 6500.0,
        'energy_mining': 6500.0,
        'energy_refining': 22000.0,
        'energy_smelting': 6500.0,
        'energy_total': 35000.0,
    },
    
    # =========================================================================
    # JOINING/BRAZING MATERIALS
    # =========================================================================
    'solder_lead_free': {
        'gwp': 25.0,
        'ap': 0.075,
        'ep': 0.028,
        'odp': 2.5e-8,
        'pocp': 0.0065,
        'htp': 15.0,
        'faetp': 8.5,
        'tetp': 0.85,
        'adp_elements': 0.045,
        'adp_fossil': 320.0,
        'water_use': 180.0,
        'land_use': 3.2,
        'energy_mining': 3.5,
        'energy_refining': 18.0,
        'energy_smelting': 5.5,
        'energy_total': 27.0,
    },
    'brazing_alloy': {
        'gwp': 85.0,
        'ap': 0.28,
        'ep': 0.088,
        'odp': 8.5e-8,
        'pocp': 0.022,
        'htp': 165.0,
        'faetp': 85.0,
        'tetp': 8.5,
        'adp_elements': 1.2,
        'adp_fossil': 1100.0,
        'water_use': 650.0,
        'land_use': 12.0,
        'energy_mining': 12.0,
        'energy_refining': 65.0,
        'energy_smelting': 22.0,
        'energy_total': 99.0,
    },
    'flux': {
        'gwp': 3.0,
        'ap': 0.015,
        'ep': 0.005,
        'odp': 3.0e-9,
        'pocp': 0.001,
        'htp': 2.5,
        'faetp': 1.2,
        'tetp': 0.12,
        'adp_elements': 0.0005,
        'adp_fossil': 38.0,
        'water_use': 25.0,
        'land_use': 0.45,
        'energy_mining': 0.5,
        'energy_refining': 2.5,
        'energy_smelting': 0.0,
        'energy_total': 3.0,
    },
}

# LCIA Impact Category Metadata - descriptions and units
LCIA_CATEGORIES = {
    'gwp': {
        'name': 'Global Warming Potential',
        'short_name': 'GWP',
        'unit': 'kg CO2-eq',
        'description': 'Contribution to climate change through greenhouse gas emissions',
        'methodology': 'IPCC AR5 100-year',
        'category_type': 'midpoint'
    },
    'ap': {
        'name': 'Acidification Potential',
        'short_name': 'AP',
        'unit': 'kg SO2-eq',
        'description': 'Contribution to acid rain through sulfur and nitrogen oxide emissions',
        'methodology': 'CML 2001',
        'category_type': 'midpoint'
    },
    'ep': {
        'name': 'Eutrophication Potential',
        'short_name': 'EP',
        'unit': 'kg PO4-eq',
        'description': 'Contribution to nutrient enrichment in water bodies',
        'methodology': 'CML 2001',
        'category_type': 'midpoint'
    },
    'odp': {
        'name': 'Ozone Depletion Potential',
        'short_name': 'ODP',
        'unit': 'kg CFC-11-eq',
        'description': 'Contribution to stratospheric ozone layer depletion',
        'methodology': 'WMO 2014',
        'category_type': 'midpoint'
    },
    'pocp': {
        'name': 'Photochemical Ozone Creation Potential',
        'short_name': 'POCP',
        'unit': 'kg C2H4-eq',
        'description': 'Contribution to ground-level ozone (smog) formation',
        'methodology': 'ReCiPe 2016',
        'category_type': 'midpoint'
    },
    'htp': {
        'name': 'Human Toxicity Potential',
        'short_name': 'HTP',
        'unit': 'kg 1,4-DCB-eq',
        'description': 'Potential toxic effects on human health',
        'methodology': 'USEtox 2.0',
        'category_type': 'midpoint'
    },
    'faetp': {
        'name': 'Freshwater Aquatic Ecotoxicity Potential',
        'short_name': 'FAETP',
        'unit': 'kg 1,4-DCB-eq',
        'description': 'Potential toxic effects on freshwater ecosystems',
        'methodology': 'USEtox 2.0',
        'category_type': 'midpoint'
    },
    'tetp': {
        'name': 'Terrestrial Ecotoxicity Potential',
        'short_name': 'TETP',
        'unit': 'kg 1,4-DCB-eq',
        'description': 'Potential toxic effects on terrestrial ecosystems',
        'methodology': 'USEtox 2.0',
        'category_type': 'midpoint'
    },
    'adp_elements': {
        'name': 'Abiotic Depletion - Elements',
        'short_name': 'ADP-E',
        'unit': 'kg Sb-eq',
        'description': 'Depletion of non-renewable mineral resources',
        'methodology': 'CML 2001',
        'category_type': 'midpoint'
    },
    'adp_fossil': {
        'name': 'Abiotic Depletion - Fossil',
        'short_name': 'ADP-F',
        'unit': 'MJ',
        'description': 'Depletion of non-renewable fossil fuel resources',
        'methodology': 'CML 2001',
        'category_type': 'midpoint'
    },
    'water_use': {
        'name': 'Water Use',
        'short_name': 'WU',
        'unit': 'L',
        'description': 'Freshwater consumption throughout the life cycle',
        'methodology': 'AWARE',
        'category_type': 'midpoint'
    },
    'land_use': {
        'name': 'Land Use',
        'short_name': 'LU',
        'unit': 'm²a',
        'description': 'Land occupation and transformation impacts',
        'methodology': 'ReCiPe 2016',
        'category_type': 'midpoint'
    },
}

# Indian Grid Emission Factors by Region (kg CO2/kWh)
# Source: CEA CO2 Baseline Database 2023
INDIAN_GRID_FACTORS = {
    'national_average': 0.82,       # India national grid average
    'northern': 0.85,               # Northern Regional Grid
    'western': 0.78,                # Western Regional Grid  
    'southern': 0.72,               # Southern Regional Grid (more hydro)
    'eastern': 0.92,                # Eastern Regional Grid (more coal)
    'northeastern': 0.65,           # Northeastern Regional Grid (more hydro)
    'captive_coal': 1.05,           # Captive coal power plant
    'captive_gas': 0.45,            # Captive natural gas
    'renewable_solar': 0.05,        # Solar PV (lifecycle)
    'renewable_wind': 0.02,         # Wind (lifecycle)
    'renewable_hydro': 0.01,        # Hydropower (lifecycle)
}

# Transport Emission Factors (kg CO2-eq per ton-km)
TRANSPORT_EMISSION_FACTORS = {
    'road_truck': 0.062,            # Heavy goods vehicle
    'road_lcv': 0.089,              # Light commercial vehicle
    'rail_freight': 0.022,          # Rail freight (Indian Railways avg)
    'rail_electric': 0.018,         # Electric rail
    'sea_container': 0.008,         # Container ship
    'sea_bulk': 0.005,              # Bulk carrier
    'air_freight': 0.602,           # Air cargo
    'pipeline': 0.015,              # Pipeline transport
}

# Default LCIA factors for unknown materials
DEFAULT_LCIA_FACTORS = {
    'gwp': 5.0,
    'ap': 0.02,
    'ep': 0.005,
    'odp': 5.0e-9,
    'pocp': 0.002,
    'htp': 5.0,
    'faetp': 2.5,
    'tetp': 0.25,
    'adp_elements': 0.001,
    'adp_fossil': 65.0,
    'water_use': 50.0,
    'land_use': 1.0,
    'energy_mining': 1.0,
    'energy_refining': 3.0,
    'energy_smelting': 2.0,
    'energy_total': 6.0,
}


def get_lcia_factors(material_type: str) -> dict:
    """Get LCIA impact factors for a material type"""
    material_lower = material_type.lower().replace(' ', '_').replace('-', '_')
    
    # Direct match
    if material_lower in LCIA_IMPACT_FACTORS:
        return LCIA_IMPACT_FACTORS[material_lower]
    
    # Check for partial matches
    for key in LCIA_IMPACT_FACTORS:
        if key in material_lower or material_lower in key:
            return LCIA_IMPACT_FACTORS[key]
    
    # Check base material type
    base_types = ['aluminium', 'copper', 'steel', 'lithium', 'cobalt', 
                  'nickel', 'tungsten', 'titanium', 'gold', 'silver', 'platinum']
    for base in base_types:
        if base in material_lower:
            primary_key = f'{base}_primary'
            if primary_key in LCIA_IMPACT_FACTORS:
                return LCIA_IMPACT_FACTORS[primary_key]
            if base in LCIA_IMPACT_FACTORS:
                return LCIA_IMPACT_FACTORS[base]
    
    return DEFAULT_LCIA_FACTORS


def calculate_lcia_impacts(material_type: str, quantity: float, recycled_content: float = 0, 
                           transport_distance: float = 0, transport_mode: str = 'road_truck',
                           grid_region: str = 'national_average') -> dict:
    """
    Calculate comprehensive LCIA impacts for a material
    
    Args:
        material_type: Type of material
        quantity: Amount in kg
        recycled_content: Percentage of recycled content (0-100)
        transport_distance: Distance in km
        transport_mode: Mode of transport (road_truck, rail_freight, sea_container, etc.)
        grid_region: Indian grid region for electricity factor
    
    Returns:
        Dictionary with all LCIA impact values
    """
    factors = get_lcia_factors(material_type)
    
    # Calculate virgin vs recycled portions
    virgin_fraction = (100 - recycled_content) / 100
    recycled_fraction = recycled_content / 100
    
    # Recycled material has ~90% lower impacts (varies by category)
    recycled_reduction = {
        'gwp': 0.10,
        'ap': 0.15,
        'ep': 0.12,
        'odp': 0.10,
        'pocp': 0.12,
        'htp': 0.15,
        'faetp': 0.15,
        'tetp': 0.15,
        'adp_elements': 0.05,  # Resources already extracted
        'adp_fossil': 0.15,
        'water_use': 0.20,
        'land_use': 0.05,
    }
    
    impacts = {}
    for category in LCIA_CATEGORIES.keys():
        base_factor = factors.get(category, DEFAULT_LCIA_FACTORS.get(category, 0))
        reduction = recycled_reduction.get(category, 0.10)
        
        # Calculate material impacts
        virgin_impact = quantity * base_factor * virgin_fraction
        recycled_impact = quantity * base_factor * reduction * recycled_fraction
        material_impact = virgin_impact + recycled_impact
        
        impacts[category] = round(material_impact, 6)
    
    # Add transport impacts (primarily affects GWP)
    if transport_distance > 0:
        transport_factor = TRANSPORT_EMISSION_FACTORS.get(transport_mode, 0.062)
        transport_gwp = (quantity / 1000) * transport_distance * transport_factor
        impacts['gwp'] = round(impacts['gwp'] + transport_gwp, 6)
        
        # Transport also contributes to other categories (simplified)
        impacts['ap'] = round(impacts['ap'] + transport_gwp * 0.002, 6)
        impacts['pocp'] = round(impacts['pocp'] + transport_gwp * 0.0003, 6)
    
    # Add energy breakdown
    impacts['energy_breakdown'] = {
        'mining_kwh': round(factors.get('energy_mining', 0) * quantity, 2),
        'refining_kwh': round(factors.get('energy_refining', 0) * quantity, 2),
        'smelting_kwh': round(factors.get('energy_smelting', 0) * quantity, 2),
        'total_kwh': round(factors.get('energy_total', 0) * quantity, 2),
    }
    
    # Calculate electricity-related GWP based on grid region
    grid_factor = INDIAN_GRID_FACTORS.get(grid_region, INDIAN_GRID_FACTORS['national_average'])
    impacts['grid_gwp'] = round(impacts['energy_breakdown']['total_kwh'] * grid_factor, 2)
    
    return impacts


def calculate_project_lcia(materials: list, grid_region: str = 'national_average') -> dict:
    """
    Calculate comprehensive LCIA for a project with multiple materials
    
    Args:
        materials: List of material dictionaries with type, quantity, recycled_content, transport_distance
        grid_region: Indian grid region for electricity calculations
    
    Returns:
        Aggregated LCIA results with per-material and total impacts
    """
    total_impacts = {cat: 0 for cat in LCIA_CATEGORIES.keys()}
    total_impacts['grid_gwp'] = 0
    total_impacts['energy_breakdown'] = {
        'mining_kwh': 0,
        'refining_kwh': 0,
        'smelting_kwh': 0,
        'total_kwh': 0,
    }
    
    materials_impacts = []
    
    for mat in materials:
        mat_type = mat.get('type', mat.get('material_type', 'unknown'))
        quantity = float(mat.get('quantity', 0))
        recycled = float(mat.get('recycled_content', 0))
        transport = float(mat.get('transport_distance', 0))
        transport_mode = mat.get('transport_mode', 'road_truck')
        
        impacts = calculate_lcia_impacts(
            mat_type, quantity, recycled, transport, transport_mode, grid_region
        )
        
        # Add to totals
        for cat in LCIA_CATEGORIES.keys():
            total_impacts[cat] += impacts.get(cat, 0)
        
        total_impacts['grid_gwp'] += impacts.get('grid_gwp', 0)
        for key in total_impacts['energy_breakdown']:
            total_impacts['energy_breakdown'][key] += impacts['energy_breakdown'].get(key, 0)
        
        materials_impacts.append({
            'name': mat.get('name', mat_type),
            'type': mat_type,
            'quantity': quantity,
            'impacts': impacts
        })
    
    # Round totals
    for cat in LCIA_CATEGORIES.keys():
        total_impacts[cat] = round(total_impacts[cat], 4)
    total_impacts['grid_gwp'] = round(total_impacts['grid_gwp'], 2)
    for key in total_impacts['energy_breakdown']:
        total_impacts['energy_breakdown'][key] = round(total_impacts['energy_breakdown'][key], 2)
    
    return {
        'total_impacts': total_impacts,
        'materials_breakdown': materials_impacts,
        'categories_metadata': LCIA_CATEGORIES,
        'grid_region': grid_region,
        'grid_emission_factor': INDIAN_GRID_FACTORS.get(grid_region, INDIAN_GRID_FACTORS['national_average']),
    }


def get_waste_factor(material_type: str) -> float:
    """Get waste generation factor for a material type"""
    material_lower = material_type.lower().replace(' ', '_').replace('-', '_')
    
    # Direct match
    if material_lower in WASTE_FACTORS:
        return WASTE_FACTORS[material_lower]
    
    # Check for partial matches (e.g., 'aluminium_alloy_6061' matches 'aluminium')
    for key in WASTE_FACTORS:
        if key in material_lower or material_lower in key:
            return WASTE_FACTORS[key]
    
    # Check base material type
    base_types = ['aluminium', 'copper', 'steel', 'iron', 'lithium', 'cobalt', 
                  'nickel', 'tungsten', 'titanium', 'gold', 'silver', 'platinum']
    for base in base_types:
        if base in material_lower:
            return WASTE_FACTORS.get(base, WASTE_FACTORS['default'])
    
    return WASTE_FACTORS['default']


def calculate_mci(recycled_content_input, recycled_content_output, 
                  target_lifespan, industry_avg_lifespan, 
                  is_designed_for_disassembly=False):
    """
    Calculate Material Circularity Indicator (MCI) using Ellen MacArthur Foundation formula
    
    MCI = 1 - LFI × F(X)
    
    Where:
    - LFI = Linear Flow Index = (V + W) / (2M)
    - V = Virgin material input = M × (1 - Fr)
    - W = Unrecoverable waste = M × (1 - Cr) × (1 - Ef)
    - Fr = Fraction from recycled sources (recycled_content_input)
    - Cr = Fraction going to recycling (recycled_content_output)
    - Ef = Efficiency of recycling process (assume 0.9)
    - F(X) = Utility factor based on lifespan
    """
    
    # Material mass (normalized to 1)
    M = 1.0
    
    # Recycling efficiency
    Ef = 0.9
    
    # Convert percentages to fractions
    Fr = recycled_content_input / 100.0  # Fraction from recycled sources
    Cr = recycled_content_output / 100.0  # Fraction going to recycling
    
    # Boost recycling output if designed for disassembly
    if is_designed_for_disassembly:
        Cr = min(1.0, Cr * 1.2)  # 20% improvement in recyclability
    
    # Virgin material input
    V = M * (1 - Fr)
    
    # Unrecoverable waste
    W = M * (1 - Cr) * (1 - Ef)
    
    # Linear Flow Index
    LFI = (V + W) / (2 * M)
    
    # Utility factor F(X) based on lifespan comparison
    # F(X) = 0.9 / X, where X = L/Lav (actual lifespan / average lifespan)
    if industry_avg_lifespan > 0 and target_lifespan > 0:
        X = target_lifespan / industry_avg_lifespan
        F_X = 0.9 / X if X > 0 else 0.9
    else:
        F_X = 0.9
    
    # Cap F(X) at reasonable bounds
    F_X = max(0.1, min(1.0, F_X))
    
    # Calculate MCI
    MCI = 1 - (LFI * F_X)
    
    # Ensure MCI is between 0 and 1
    MCI = max(0, min(1, MCI))
    
    return round(MCI, 3)


def calculate_circular_design_score(mci_score, target_lifespan, industry_avg_lifespan,
                                    is_designed_for_disassembly, avg_recycled_content):
    """
    Calculate Circular Design Score (0-100)
    Combines: MCI (40%) + Durability (30%) + Design for Recycling (30%)
    """
    
    # MCI contribution (0-40 points)
    mci_points = mci_score * 40
    
    # Durability contribution (0-30 points)
    # Based on lifespan vs industry average
    if industry_avg_lifespan > 0:
        lifespan_ratio = target_lifespan / industry_avg_lifespan
        durability_points = min(30, lifespan_ratio * 20)  # Cap at 30
    else:
        durability_points = 15
    
    # Design for Recycling contribution (0-30 points)
    # Based on recyclability + disassembly + recycled content
    recycling_points = 0
    
    if is_designed_for_disassembly:
        recycling_points += 15
    
    # Recycled content contribution (0-15 points)
    recycling_points += (avg_recycled_content / 100) * 15
    
    # Total score
    total_score = mci_points + durability_points + recycling_points
    
    return round(min(100, max(0, total_score)), 1)


# =====================================================
# NLP PARSING ENGINE - Natural Language to BOM
# =====================================================

# Material patterns for NLP matching
NLP_MATERIAL_PATTERNS = {
    'aluminium': {
        'keywords': ['aluminium', 'aluminum'],
        'forms': ['sheet', 'wire', 'rod', 'bar', 'plate', 'tube', 'foil', 'ingot', 'casting', 'alloy'],
        'default_type': 'aluminium_primary',
        'recycled_type': 'aluminium_secondary',
        'national_baseline_recycled': 25,  # India avg recycled content %
        'gwp_factor': 12.5
    },
    'copper': {
        'keywords': ['copper', 'brass', 'bronze'],
        'forms': ['wire', 'cable', 'sheet', 'rod', 'tube', 'pipe', 'coil'],
        'default_type': 'copper_primary',
        'recycled_type': 'copper_secondary',
        'national_baseline_recycled': 35,  # India avg recycled content %
        'gwp_factor': 3.5
    },
    'steel': {
        'keywords': ['steel', 'stainless steel', 'carbon steel', 'mild steel', 'cast steel', 'tool steel', 'alloy steel', 'high-speed steel'],
        'forms': ['sheet', 'rod', 'bar', 'beam', 'wire', 'plate', 'tube', 'coil', 'shank', 'body', 'casting'],
        'default_type': 'steel_primary',
        'recycled_type': 'steel_secondary',
        'national_baseline_recycled': 40,  # India avg recycled content %
        'gwp_factor': 2.1
    },
    'iron': {
        'keywords': ['iron', 'cast iron', 'pig iron', 'wrought iron'],
        'forms': ['casting', 'ingot', 'bar', 'plate'],
        'default_type': 'steel_primary',
        'recycled_type': 'steel_secondary',
        'national_baseline_recycled': 40,
        'gwp_factor': 1.9
    },
    'lithium': {
        'keywords': ['lithium', 'lithium-ion', 'li-ion', 'lifepo4', 'lithium carbonate', 'lithium hydroxide'],
        'forms': ['carbonate', 'hydroxide', 'oxide', 'battery', 'cell'],
        'default_type': 'lithium',
        'recycled_type': 'lithium',
        'national_baseline_recycled': 5,
        'gwp_factor': 15.0
    },
    'cobalt': {
        'keywords': ['cobalt', 'cobalt sulfate', 'cobalt oxide'],
        'forms': ['sulfate', 'oxide', 'powder', 'metal'],
        'default_type': 'cobalt',
        'recycled_type': 'cobalt',
        'national_baseline_recycled': 10,
        'gwp_factor': 10.0
    },
    'nickel': {
        'keywords': ['nickel', 'nickel sulfate', 'ferronickel'],
        'forms': ['class1', 'ferronickel', 'sulfate', 'plating'],
        'default_type': 'nickel',
        'recycled_type': 'nickel',
        'national_baseline_recycled': 15,
        'gwp_factor': 8.5
    },
    # Critical Minerals - Expanded
    'manganese': {
        'keywords': ['manganese', 'mn'],
        'forms': ['dioxide', 'oxide', 'sulfate', 'electrolytic'],
        'default_type': 'manganese',
        'recycled_type': 'manganese',
        'national_baseline_recycled': 20,
        'gwp_factor': 2.8,
        'scarcity_score': 40
    },
    'graphite': {
        'keywords': ['graphite', 'carbon', 'anode'],
        'forms': ['natural', 'synthetic', 'spherical', 'flake'],
        'default_type': 'graphite',
        'recycled_type': 'graphite',
        'national_baseline_recycled': 5,
        'gwp_factor': 4.2,
        'scarcity_score': 70
    },
    # Rare Earths
    'neodymium': {
        'keywords': ['neodymium', 'nd', 'ndfeb', 'neo'],
        'forms': ['magnet', 'oxide', 'metal', 'alloy'],
        'default_type': 'neodymium',
        'recycled_type': 'neodymium',
        'national_baseline_recycled': 2,
        'gwp_factor': 35.0,
        'scarcity_score': 95
    },
    'dysprosium': {
        'keywords': ['dysprosium', 'dy'],
        'forms': ['oxide', 'metal', 'magnet'],
        'default_type': 'dysprosium',
        'recycled_type': 'dysprosium',
        'national_baseline_recycled': 1,
        'gwp_factor': 45.0,
        'scarcity_score': 98
    },
    'praseodymium': {
        'keywords': ['praseodymium', 'pr'],
        'forms': ['oxide', 'metal'],
        'default_type': 'praseodymium',
        'recycled_type': 'praseodymium',
        'national_baseline_recycled': 1,
        'gwp_factor': 32.0,
        'scarcity_score': 90
    },
    'rare_earth': {
        'keywords': ['rare earth', 'ree', 'rare-earth', 'lanthanide'],
        'forms': ['oxide', 'metal', 'mixed', 'concentrate'],
        'default_type': 'rare_earth_mixed',
        'recycled_type': 'rare_earth_mixed',
        'national_baseline_recycled': 2,
        'gwp_factor': 38.0,
        'scarcity_score': 93
    },
    # Other Critical Minerals
    'tungsten': {
        'keywords': ['tungsten', 'tungsten carbide', 'wolfram', 'carbide alloy'],
        'forms': ['carbide', 'powder', 'wire', 'electrode', 'alloy', 'bit', 'tool'],
        'default_type': 'tungsten',
        'recycled_type': 'tungsten',
        'national_baseline_recycled': 30,
        'gwp_factor': 22.0,
        'scarcity_score': 78
    },
    'vanadium': {
        'keywords': ['vanadium', 'vanadium oxide', 'vanadium pentoxide'],
        'forms': ['oxide', 'pentoxide', 'alloy', 'redox'],
        'default_type': 'vanadium',
        'recycled_type': 'vanadium',
        'national_baseline_recycled': 15,
        'gwp_factor': 28.0,
        'scarcity_score': 72
    },
    'titanium': {
        'keywords': ['titanium', 'titanium nitride', 'titanium dioxide', 'titanium oxide'],
        'forms': ['sponge', 'ingot', 'sheet', 'powder', 'dioxide', 'nitride', 'coating', 'oxide'],
        'default_type': 'titanium',
        'recycled_type': 'titanium',
        'national_baseline_recycled': 25,
        'gwp_factor': 8.1,
        'scarcity_score': 45
    },
    'tantalum': {
        'keywords': ['tantalum'],
        'forms': ['powder', 'capacitor', 'wire'],
        'default_type': 'tantalum',
        'recycled_type': 'tantalum',
        'national_baseline_recycled': 20,
        'gwp_factor': 48.0,
        'scarcity_score': 88
    },
    'indium': {
        'keywords': ['indium', 'indium tin oxide', 'ito'],
        'forms': ['oxide', 'tin oxide', 'metal'],
        'default_type': 'indium',
        'recycled_type': 'indium',
        'national_baseline_recycled': 35,
        'gwp_factor': 142.0,
        'scarcity_score': 82
    },
    'gallium': {
        'keywords': ['gallium', 'gallium arsenide', 'gaas', 'gallium nitride'],
        'forms': ['arsenide', 'nitride', 'metal'],
        'default_type': 'gallium',
        'recycled_type': 'gallium',
        'national_baseline_recycled': 25,
        'gwp_factor': 185.0,
        'scarcity_score': 80
    },
    'platinum': {
        'keywords': ['platinum', 'pt', 'pgm'],
        'forms': ['catalyst', 'wire', 'sheet', 'powder'],
        'default_type': 'platinum',
        'recycled_type': 'platinum',
        'national_baseline_recycled': 40,
        'gwp_factor': 12500.0,
        'scarcity_score': 88
    },
    'palladium': {
        'keywords': ['palladium', 'pd'],
        'forms': ['catalyst', 'alloy', 'powder'],
        'default_type': 'palladium',
        'recycled_type': 'palladium',
        'national_baseline_recycled': 45,
        'gwp_factor': 9800.0,
        'scarcity_score': 86
    },
    # Refractory and Ceramic Materials
    'magnesia': {
        'keywords': ['magnesia', 'magnesium oxide', 'mgo', 'magnesia-chrome', 'mag-chrome'],
        'forms': ['brick', 'refractory', 'lining', 'block', 'castable'],
        'default_type': 'magnesia_refractory',
        'recycled_type': 'magnesia_refractory',
        'national_baseline_recycled': 10,
        'gwp_factor': 1.8,
        'scarcity_score': 25
    },
    'zirconia': {
        'keywords': ['zirconia', 'zirconium oxide', 'zro2', 'yttria-stabilized zirconia', 'ysz'],
        'forms': ['coating', 'layer', 'ceramic', 'brick', 'powder', 'thermal barrier'],
        'default_type': 'zirconia_ceramic',
        'recycled_type': 'zirconia_ceramic',
        'national_baseline_recycled': 5,
        'gwp_factor': 8.5,
        'scarcity_score': 55
    },
    'chrome': {
        'keywords': ['chrome', 'chromium', 'chromite', 'chrome-magnesia', 'ferrochrome'],
        'forms': ['refractory', 'brick', 'ore', 'plating', 'alloy'],
        'default_type': 'chrome_refractory',
        'recycled_type': 'chrome_refractory',
        'national_baseline_recycled': 20,
        'gwp_factor': 5.2,
        'scarcity_score': 50
    },
    'refractory': {
        'keywords': ['refractory', 'firebite', 'high alumina', 'silica brick', 'fire brick'],
        'forms': ['brick', 'lining', 'castable', 'mortar', 'cement'],
        'default_type': 'refractory_generic',
        'recycled_type': 'refractory_generic',
        'national_baseline_recycled': 15,
        'gwp_factor': 1.5,
        'scarcity_score': 20
    },
    'silicon_carbide': {
        'keywords': ['silicon carbide', 'sic', 'carborundum'],
        'forms': ['brick', 'crucible', 'abrasive', 'heating element'],
        'default_type': 'silicon_carbide',
        'recycled_type': 'silicon_carbide',
        'national_baseline_recycled': 10,
        'gwp_factor': 12.0,
        'scarcity_score': 40
    },
    'alumina': {
        'keywords': ['alumina', 'aluminum oxide', 'al2o3', 'corundum'],
        'forms': ['ceramic', 'brick', 'powder', 'coating', 'abrasive'],
        'default_type': 'alumina_ceramic',
        'recycled_type': 'alumina_ceramic',
        'national_baseline_recycled': 10,
        'gwp_factor': 3.2,
        'scarcity_score': 30
    },
    # Precious Metals
    'silver': {
        'keywords': ['silver', 'ag', 'sterling silver'],
        'forms': ['wire', 'sheet', 'powder', 'paste', 'coating', 'contact'],
        'default_type': 'silver',
        'recycled_type': 'silver',
        'national_baseline_recycled': 50,
        'gwp_factor': 104.0,
        'scarcity_score': 60
    },
    'gold': {
        'keywords': ['gold', 'au'],
        'forms': ['wire', 'sheet', 'powder', 'plating', 'contact', 'bonding'],
        'default_type': 'gold',
        'recycled_type': 'gold',
        'national_baseline_recycled': 60,
        'gwp_factor': 31500.0,
        'scarcity_score': 95
    },
    # Joining/Brazing Materials
    'brazing_alloy': {
        'keywords': ['brazing alloy', 'braze', 'brazing', 'bag-1', 'silver braze', 'copper braze'],
        'forms': ['rod', 'wire', 'paste', 'ring', 'preform', 'filler'],
        'default_type': 'brazing_alloy',
        'recycled_type': 'brazing_alloy',
        'national_baseline_recycled': 30,
        'gwp_factor': 85.0,
        'scarcity_score': 50
    },
    'solder': {
        'keywords': ['solder', 'soldering', 'lead-free solder', 'sn-ag-cu', 'sac305', 'sac405'],
        'forms': ['wire', 'paste', 'bar', 'ball', 'preform'],
        'default_type': 'solder_lead_free',
        'recycled_type': 'solder_lead_free',
        'national_baseline_recycled': 35,
        'gwp_factor': 25.0,
        'scarcity_score': 25
    },
    'flux': {
        'keywords': ['flux', 'brazing flux', 'soldering flux'],
        'forms': ['powder', 'paste', 'liquid'],
        'default_type': 'flux',
        'recycled_type': 'flux',
        'national_baseline_recycled': 0,
        'gwp_factor': 3.0,
        'scarcity_score': 5
    }
}

# Product category detection
PRODUCT_CATEGORIES = {
    'mining': ['mining', 'drill', 'drilling', 'excavation', 'extraction', 'ore', 'quarry', 'underground', 'excavator', 'bucket', 'teeth', 'hard-rock', 'hardrock'],
    'metallurgy': ['metallurgy', 'smelting', 'foundry', 'casting', 'forging', 'metal processing', 'furnace', 'crucible', 'ladle', 'converter'],
    'ev_battery': ['battery', 'cell', 'ev', 'electric vehicle', 'lithium-ion', 'bms', 'cathode', 'anode'],
    'power_transmission': ['transformer', 'cable', 'wire', 'conductor', 'transmission', 'power line'],
    'construction': ['building', 'structure', 'beam', 'rebar', 'construction', 'roof'],
    'automotive': ['car', 'vehicle', 'automobile', 'automotive'],
    'electronics': ['pcb', 'circuit', 'electronic', 'component', 'chip', 'semiconductor'],
    'packaging': ['can', 'container', 'foil', 'packaging', 'wrap'],
    'appliances': ['appliance', 'refrigerator', 'ac', 'washing', 'machine'],
    'renewable_energy': ['solar', 'wind', 'turbine', 'panel', 'inverter', 'generator'],
    'magnets': ['magnet', 'permanent magnet', 'ndfeb', 'motor magnet'],
    'industrial_tools': ['tool', 'drill bit', 'cutting', 'milling', 'lathe', 'cnc'],
    'refractory': ['refractory', 'kiln', 'incinerator', 'boiler', 'heat exchanger']
}


def parse_nlp_input(description):
    """
    Parse natural language description into structured project/material data
    
    Input: "10kg copper wire, PVC coated, used in a motor for 10 years"
    Output: {
        materials: [{name, type, quantity, unit, recycled_content, ...}],
        project: {product_category, target_lifespan, ...},
        assumptions: [{field, value, reason}]
    }
    """
    import re
    
    description_lower = description.lower()
    result = {
        'materials': [],
        'project': {
            'product_category': '',
            'target_lifespan': None,
            'is_designed_for_disassembly': False
        },
        'assumptions': [],
        'tokens': []
    }
    
    # === TOKENIZATION ===
    
    # Extract quantity patterns with units (e.g., "10kg", "5 kg", "100 grams", "800kg")
    # Store as tuples of (value, unit) to convert properly
    quantity_with_unit_patterns = [
        (r'(\d+(?:\.\d+)?)\s*(?:kg|kilogram|kilograms)\b', 'kg'),
        (r'(\d+(?:\.\d+)?)\s*(?:g|gram|grams)\b', 'g'),
        (r'(\d+(?:\.\d+)?)\s*(?:t|ton|tons|tonne|tonnes)\b', 't'),
        (r'(\d+(?:\.\d+)?)\s*(?:lb|pound|pounds)\b', 'lb'),
        (r'(\d+(?:\.\d+)?)\s*(?:mm|millimeter|millimeters)\b', 'mm'),  # For thickness
    ]
    
    quantities = []
    thickness_values = []
    
    for pattern, unit in quantity_with_unit_patterns:
        matches = re.findall(pattern, description_lower)
        for match in matches:
            value = float(match)
            # Convert to kg (keep kg as-is, don't divide by 1000!)
            if unit == 'kg':
                quantities.append(value)  # Already in kg
            elif unit == 'g':
                quantities.append(value / 1000)  # grams to kg
            elif unit == 't':
                quantities.append(value * 1000)  # tons to kg
            elif unit == 'lb':
                quantities.append(value * 0.453592)  # pounds to kg
            elif unit == 'mm':
                thickness_values.append(value)  # Store thickness separately
    
    # If no quantity found, default to 1kg
    if not quantities:
        quantities = [1.0]
        result['assumptions'].append({
            'field': 'quantity',
            'value': '1 kg',
            'reason': 'No quantity specified, using default'
        })
    
    result['tokens'].append({'type': 'quantity', 'values': quantities})
    if thickness_values:
        result['tokens'].append({'type': 'thickness', 'values': thickness_values})
    
    # Extract lifespan (e.g., "10 years", "5y", "15 year")
    lifespan_patterns = [
        r'(\d+)\s*(?:years?|y|yr|yrs)',
        r'for\s+(\d+)\s*(?:years?|y)',
        r'lifespan[:\s]+(\d+)',
        r'lifetime[:\s]+(\d+)',
    ]
    
    # Patterns for months (convert to years)
    month_lifespan_patterns = [
        r'(\d+)\s*(?:months?|mo)',
        r'last[s]?\s+(\d+)\s*(?:months?|mo)',
        r'for\s+(\d+)\s*(?:months?|mo)',
        r'lifespan[:\s]+(\d+)\s*(?:months?|mo)',
    ]
    
    # Patterns for weeks (convert to years)
    week_lifespan_patterns = [
        r'(\d+)\s*(?:weeks?|wk|wks)',
        r'last[s]?\s+(\d+)\s*(?:weeks?)',
    ]
    
    # Patterns for days (convert to years)
    day_lifespan_patterns = [
        r'(\d+)\s*(?:days?|d)',
        r'last[s]?\s+(\d+)\s*(?:days?)',
    ]
    
    lifespan_found = False
    
    # Try years first
    for pattern in lifespan_patterns:
        match = re.search(pattern, description_lower)
        if match:
            lifespan_years = int(match.group(1))
            result['project']['target_lifespan'] = lifespan_years
            result['tokens'].append({'type': 'lifespan', 'value': lifespan_years, 'unit': 'years'})
            lifespan_found = True
            break
    
    # Try months if years not found
    if not lifespan_found:
        for pattern in month_lifespan_patterns:
            match = re.search(pattern, description_lower)
            if match:
                months = int(match.group(1))
                lifespan_years = round(months / 12, 2)
                result['project']['target_lifespan'] = lifespan_years
                result['tokens'].append({'type': 'lifespan', 'value': lifespan_years, 'unit': 'years', 'original': f'{months} months'})
                result['assumptions'].append({
                    'field': 'lifespan',
                    'value': f'{lifespan_years} years',
                    'reason': f'Converted from {months} months'
                })
                lifespan_found = True
                break
    
    # Try weeks if still not found
    if not lifespan_found:
        for pattern in week_lifespan_patterns:
            match = re.search(pattern, description_lower)
            if match:
                weeks = int(match.group(1))
                lifespan_years = round(weeks / 52, 2)
                result['project']['target_lifespan'] = lifespan_years
                result['tokens'].append({'type': 'lifespan', 'value': lifespan_years, 'unit': 'years', 'original': f'{weeks} weeks'})
                result['assumptions'].append({
                    'field': 'lifespan',
                    'value': f'{lifespan_years} years',
                    'reason': f'Converted from {weeks} weeks'
                })
                lifespan_found = True
                break
    
    # Try days if still not found
    if not lifespan_found:
        for pattern in day_lifespan_patterns:
            match = re.search(pattern, description_lower)
            if match:
                days = int(match.group(1))
                lifespan_years = round(days / 365, 2)
                result['project']['target_lifespan'] = lifespan_years
                result['tokens'].append({'type': 'lifespan', 'value': lifespan_years, 'unit': 'years', 'original': f'{days} days'})
                result['assumptions'].append({
                    'field': 'lifespan',
                    'value': f'{lifespan_years} years',
                    'reason': f'Converted from {days} days'
                })
                lifespan_found = True
                break
    
    # Extract recycled content (e.g., "30% recycled", "recycled content 25%")
    recycled_patterns = [
        r'(\d+)\s*%\s*recycled',
        r'recycled[:\s]+(\d+)\s*%',
        r'recycled content[:\s]+(\d+)',
    ]
    
    recycled_content = None
    for pattern in recycled_patterns:
        match = re.search(pattern, description_lower)
        if match:
            recycled_content = int(match.group(1))
            result['tokens'].append({'type': 'recycled_content', 'value': recycled_content})
            break
    
    # Check for "recycled" keyword without percentage
    if recycled_content is None and 'recycled' in description_lower:
        recycled_content = 100  # Assume fully recycled
        result['tokens'].append({'type': 'recycled_content', 'value': 100})
    
    # === MATERIAL MAPPING ===
    
    detected_materials = []
    detected_material_keys = set()  # Track already detected materials to avoid duplicates
    
    # Sort materials by keyword length (longest first) to match specific terms first
    sorted_materials = sorted(
        NLP_MATERIAL_PATTERNS.items(),
        key=lambda x: max(len(k) for k in x[1]['keywords']),
        reverse=True
    )
    
    for material_key, material_data in sorted_materials:
        if material_key in detected_material_keys:
            continue
            
        for keyword in sorted(material_data['keywords'], key=len, reverse=True):
            # Use word boundary matching to avoid false positives
            # e.g., 'co' should not match 'coated', 'coating', etc.
            pattern = r'\b' + re.escape(keyword) + r'\b'
            match_obj = re.search(pattern, description_lower)
            if match_obj:
                # === CONTEXT-AWARE FILTERING ===
                # Check if the material appears in a context that indicates it's NOT a component
                # e.g., "copper smelting" means the furnace processes copper, not made OF copper
                
                match_pos = match_obj.start()
                context_window = description_lower[max(0, match_pos-30):min(len(description_lower), match_pos+len(keyword)+30)]
                
                # Skip if material appears in process context (not as a component)
                process_indicators = [
                    f'{keyword} smelting', f'{keyword} processing', f'{keyword} refining',
                    f'{keyword} production', f'{keyword} extraction', f'{keyword} mining',
                    f'{keyword} plants', f'{keyword} plant', f'{keyword} industry',
                    f'smelting {keyword}', f'processing {keyword}', f'refining {keyword}',
                    f'produces {keyword}', f'manufacture {keyword}', f'making {keyword}'
                ]
                
                is_process_context = False
                for indicator in process_indicators:
                    if indicator in description_lower:
                        is_process_context = True
                        break
                
                if is_process_context:
                    # Skip this material - it's what the equipment processes, not made of
                    continue
                
                # Check for forms
                detected_form = None
                for form in material_data['forms']:
                    form_pattern = r'\b' + re.escape(form) + r'\b'
                    if re.search(form_pattern, description_lower):
                        detected_form = form
                        break
                
                # Determine if recycled
                is_recycled = bool(re.search(r'\b(recycled|secondary|scrap)\b', description_lower))
                material_type = material_data['recycled_type'] if is_recycled else material_data['default_type']
                
                # Use detected recycled content or national baseline
                mat_recycled = recycled_content
                if mat_recycled is None:
                    mat_recycled = material_data['national_baseline_recycled']
                    result['assumptions'].append({
                        'field': 'recycled_content',
                        'value': f"{mat_recycled}%",
                        'reason': f"Using India National Baseline for {material_key.title()}"
                    })
                
                detected_materials.append({
                    'material_key': material_key,
                    'name': f"{material_key.title()} {'(' + detected_form.title() + ')' if detected_form else ''}".strip(),
                    'type': material_type,
                    'form': detected_form,
                    'gwp_factor': material_data['gwp_factor'] if not is_recycled else material_data['gwp_factor'] * 0.05,
                    'recycled_content': mat_recycled
                })
                
                result['tokens'].append({
                    'type': 'material',
                    'material': material_key,
                    'form': detected_form,
                    'is_recycled': is_recycled,
                    'matched_keyword': keyword
                })
                
                detected_material_keys.add(material_key)
                break
    
    # If no materials detected, try to infer from context
    if not detected_materials:
        result['assumptions'].append({
            'field': 'material',
            'value': 'Unknown',
            'reason': 'Could not identify material from description. Please specify material type.'
        })
    
    # Create material entries
    for i, mat in enumerate(detected_materials):
        quantity = quantities[i] if i < len(quantities) else quantities[0]
        result['materials'].append({
            'material_name': mat['name'],
            'material_type': mat['type'],
            'quantity': quantity,
            'unit': 'kg',
            'recycled_content': mat['recycled_content'],
            'gwp_factor': mat['gwp_factor'],
            'transport_distance': 100  # Default transport distance
        })
    
    # === PRODUCT CATEGORY DETECTION ===
    
    # Sort categories by keyword length (longest first) for better matching
    for category, keywords in sorted(PRODUCT_CATEGORIES.items(), key=lambda x: max(len(k) for k in x[1]), reverse=True):
        for keyword in sorted(keywords, key=len, reverse=True):
            # Use word boundary matching
            pattern = r'\b' + re.escape(keyword) + r'\b'
            if re.search(pattern, description_lower):
                result['project']['product_category'] = category
                result['tokens'].append({'type': 'category', 'value': category, 'matched_keyword': keyword})
                break
        if result['project']['product_category']:
            break
    
    # Apply default lifespan if not specified
    if result['project']['target_lifespan'] is None:
        category = result['project']['product_category'] or 'other'
        default_lifespan = INDUSTRY_BENCHMARKS.get(category, {}).get('avg_lifespan', 10)
        result['project']['target_lifespan'] = default_lifespan
        result['assumptions'].append({
            'field': 'target_lifespan',
            'value': f"{default_lifespan} years",
            'reason': f"Using industry average for {category.replace('_', ' ').title()}"
        })
    
    # Check for disassembly keywords
    disassembly_keywords = ['disassembly', 'disassemble', 'modular', 'repairable', 'serviceable']
    for keyword in disassembly_keywords:
        if keyword in description_lower:
            result['project']['is_designed_for_disassembly'] = True
            result['tokens'].append({'type': 'disassembly', 'value': True})
            break
    
    # Detect coatings/finishes
    coating_keywords = ['pvc', 'coated', 'anodized', 'galvanized', 'painted', 'plated', 'chrome', 
                       'titanium nitride', 'tin coating', 'nickel plating', 'zinc coating', 
                       'powder coated', 'epoxy', 'enamel']
    detected_coatings = []
    for keyword in coating_keywords:
        pattern = r'\b' + re.escape(keyword) + r'\b'
        if re.search(pattern, description_lower):
            detected_coatings.append(keyword)
            result['tokens'].append({'type': 'coating', 'value': keyword})
    
    # Add coating info to result
    if detected_coatings:
        result['coatings'] = detected_coatings
    
    # === GENERATE SUGGESTED PROJECT NAME ===
    
    # Product name templates based on category and detected keywords
    PRODUCT_NAME_TEMPLATES = {
        'ev_battery': ['Battery Pack', 'EV Battery Module', 'Lithium-Ion Battery Pack', 'Battery Cell Assembly'],
        'mining': ['Mining Equipment', 'Drill Bit Assembly', 'Mining Tool', 'Extraction Equipment'],
        'metallurgy': ['Industrial Furnace', 'Smelting Furnace', 'Foundry Equipment', 'Metal Processing Unit'],
        'power_transmission': ['Power Cable', 'Transformer Unit', 'Transmission Line', 'Electrical Conductor'],
        'construction': ['Construction Component', 'Structural Element', 'Building Material', 'Rebar Assembly'],
        'automotive': ['Automotive Part', 'Vehicle Component', 'Engine Part', 'Automotive Assembly'],
        'electronics': ['Electronic Component', 'Circuit Board', 'Semiconductor Device', 'Electronic Module'],
        'packaging': ['Packaging Material', 'Container', 'Packaging Solution'],
        'appliances': ['Home Appliance', 'Appliance Component', 'Household Equipment'],
        'renewable_energy': ['Solar Panel', 'Wind Turbine Component', 'Renewable Energy System', 'Generator Unit'],
        'magnets': ['Permanent Magnet', 'Magnet Assembly', 'Motor Magnet'],
        'industrial_tools': ['Industrial Tool', 'Cutting Tool', 'Manufacturing Equipment'],
        'refractory': ['Refractory Lining', 'Furnace Lining', 'Kiln Component', 'Heat-Resistant Equipment']
    }
    
    # Try to extract a specific product name from the description
    product_keywords = [
        (r'\b(battery pack|battery module|battery cell)\b', 'Battery Pack'),
        (r'\b(drill bit|drilling tool)\b', 'Drill Bit'),
        (r'\b(solar panel|pv panel|photovoltaic)\b', 'Solar Panel'),
        (r'\b(wind turbine|turbine generator)\b', 'Wind Turbine'),
        (r'\b(transformer)\b', 'Transformer'),
        (r'\b(motor|electric motor)\b', 'Electric Motor'),
        (r'\b(cable|wire|conductor)\b', 'Cable Assembly'),
        (r'\b(circuit board|pcb)\b', 'Circuit Board'),
        (r'\b(catalytic converter)\b', 'Catalytic Converter'),
        (r'\b(frame|mounting)\b', 'Frame Assembly'),
        (r'\b(rebar|reinforcement)\b', 'Rebar'),
        (r'\b(smartphone|phone|mobile)\b', 'Smartphone'),
        (r'\b(laptop|computer)\b', 'Laptop'),
        (r'\b(generator)\b', 'Generator'),
        (r'\b(furnace|smelting furnace|blast furnace|arc furnace)\b', 'Industrial Furnace'),
        (r'\b(kiln|rotary kiln)\b', 'Industrial Kiln'),
        (r'\b(boiler)\b', 'Industrial Boiler'),
        (r'\b(heat exchanger)\b', 'Heat Exchanger'),
        (r'\b(excavator\s*bucket\s*teeth?|bucket\s*teeth?)\b', 'Excavator Bucket Teeth'),
        (r'\b(excavator|backhoe|digger)\b', 'Excavator Component'),
        (r'\b(crusher|jaw crusher|cone crusher)\b', 'Crusher'),
        (r'\b(conveyor|conveyor belt)\b', 'Conveyor System'),
        (r'\b(grinding mill|ball mill|sag mill)\b', 'Grinding Mill'),
        (r'\b(shovel|loading shovel)\b', 'Mining Shovel'),
        (r'\b(dragline)\b', 'Dragline'),
        (r'\b(haul truck|mining truck|dump truck)\b', 'Haul Truck Component'),
    ]
    
    suggested_name = None
    for pattern, name in product_keywords:
        if re.search(pattern, description_lower):
            suggested_name = name
            break
    
    # Fallback to category-based name
    if not suggested_name:
        category = result['project']['product_category']
        if category and category in PRODUCT_NAME_TEMPLATES:
            suggested_name = PRODUCT_NAME_TEMPLATES[category][0]
        else:
            # Generate from primary material
            if detected_materials:
                primary_material = detected_materials[0]['material_key'].title()
                suggested_name = f"{primary_material} Product"
            else:
                suggested_name = "LCA Project"
    
    result['suggested_name'] = suggested_name
    
    return result


def parse_nlp_with_groq(description):
    """
    Use Groq LLM to intelligently parse product descriptions into structured BOM data.
    Handles composite materials, proper weight attribution, and context understanding.
    """
    if not groq_client:
        return None  # Fall back to regex-based parsing
    
    prompt = f"""You are an expert materials engineer and LCA (Life Cycle Assessment) specialist. 
Analyze the following product description and extract structured Bill of Materials (BOM) data.

CRITICAL RULES:
1. COMPOSITE MATERIALS: If a material is described as a composite (e.g., "magnesia-chrome"), treat it as ONE material entry, not separate entries.
2. WEIGHT ATTRIBUTION: Only assign a weight to a material if the weight is directly attached to it syntactically. For example:
   - "800kg of magnesia-chrome bricks" → magnesia-chrome gets 800kg
   - "coated with a 3mm layer of zirconia" → zirconia weight is UNKNOWN (only thickness given)
3. COATINGS vs BULK: Distinguish between coating materials (thin layers) and bulk materials.
4. PROCESS vs COMPONENT: If a material is mentioned as what the equipment PROCESSES (e.g., "copper smelting plant"), do NOT include it as a component material.
5. LIFESPAN: Extract the service lifespan in years.
6. CATEGORY: Identify the product category (ev_battery, mining, metallurgy, construction, automotive, electronics, renewable_energy, etc.)

PRODUCT DESCRIPTION:
{description}

Respond ONLY with valid JSON in this exact format (no markdown, no explanation):
{{
  "suggested_name": "Product Name",
  "product_category": "category_name",
  "target_lifespan": number_or_null,
  "is_designed_for_disassembly": true_or_false,
  "materials": [
    {{
      "material_name": "Material Name",
      "material_type": "material_type_code",
      "quantity": number_or_null,
      "unit": "kg",
      "quantity_note": "explanation if quantity is estimated or unknown",
      "recycled_content": number_0_to_100,
      "is_coating": true_or_false,
      "is_composite": true_or_false,
      "composite_components": ["component1", "component2"] or null
    }}
  ],
  "coatings": ["coating1", "coating2"],
  "assumptions": [
    {{
      "field": "field_name",
      "value": "assumed_value",
      "reason": "why this was assumed"
    }}
  ]
}}"""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a precise JSON generator for LCA material extraction. Output only valid JSON, no markdown formatting."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Clean up response - remove markdown code blocks if present
        if response_text.startswith('```'):
            response_text = response_text.split('```')[1]
            if response_text.startswith('json'):
                response_text = response_text[4:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        # Parse the JSON response
        import json
        parsed_data = json.loads(response_text)
        
        # Convert to the expected format
        result = {
            'materials': [],
            'project': {
                'product_category': parsed_data.get('product_category', ''),
                'target_lifespan': parsed_data.get('target_lifespan'),
                'is_designed_for_disassembly': parsed_data.get('is_designed_for_disassembly', False)
            },
            'assumptions': parsed_data.get('assumptions', []),
            'tokens': [],
            'coatings': parsed_data.get('coatings', []),
            'suggested_name': parsed_data.get('suggested_name', 'LCA Project'),
            'parsing_method': 'groq_llm'
        }
        
        # Process materials
        for mat in parsed_data.get('materials', []):
            material_entry = {
                'material_name': mat.get('material_name', 'Unknown'),
                'material_type': mat.get('material_type', 'custom'),
                'quantity': mat.get('quantity') if mat.get('quantity') else 1.0,
                'unit': mat.get('unit', 'kg'),
                'recycled_content': mat.get('recycled_content', 0),
                'transport_distance': 100,
                'gwp_factor': get_gwp_factor(mat.get('material_type', 'custom')),
                'is_coating': mat.get('is_coating', False),
                'is_composite': mat.get('is_composite', False),
                'quantity_note': mat.get('quantity_note', '')
            }
            
            # Add to assumptions if quantity was unknown
            if mat.get('quantity') is None:
                result['assumptions'].append({
                    'field': f"{mat.get('material_name')} quantity",
                    'value': 'Unknown',
                    'reason': mat.get('quantity_note', 'Weight not specified in description')
                })
            
            result['materials'].append(material_entry)
            
            # Add token for UI display
            result['tokens'].append({
                'type': 'material',
                'material': mat.get('material_name'),
                'form': 'coating' if mat.get('is_coating') else 'bulk',
                'is_recycled': mat.get('recycled_content', 0) > 50,
                'is_composite': mat.get('is_composite', False)
            })
        
        # Add category token
        if result['project']['product_category']:
            result['tokens'].append({
                'type': 'category',
                'value': result['project']['product_category']
            })
        
        # Add lifespan token
        if result['project']['target_lifespan']:
            result['tokens'].append({
                'type': 'lifespan',
                'value': result['project']['target_lifespan']
            })
        
        return result
        
    except Exception as e:
        print(f"Groq NLP parsing error: {e}")
        return None  # Fall back to regex-based parsing


def get_gwp_factor(material_type):
    """Get GWP factor for a material type from EMISSION_FACTORS"""
    # Check direct match
    if material_type in EMISSION_FACTORS:
        return EMISSION_FACTORS[material_type]['production']
    
    # Check NLP patterns for GWP
    for key, data in NLP_MATERIAL_PATTERNS.items():
        if key in material_type.lower() or material_type.lower() in data.get('keywords', []):
            return data.get('gwp_factor', 2.0)
    
    # Default GWP factor
    return 2.0


@app.route('/api/v1/parse-nlp', methods=['POST', 'OPTIONS'])
def parse_nlp():
    """Parse natural language description into structured BOM data"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        description = data.get('description', '')
        use_ai = data.get('use_ai', True)  # Default to using Groq if available
        
        if not description:
            return jsonify({"detail": "Description is required"}), 400
        
        result = None
        
        # Try Groq-powered parsing first if enabled
        if use_ai and groq_client:
            result = parse_nlp_with_groq(description)
        
        # Fall back to regex-based parsing
        if result is None:
            result = parse_nlp_input(description)
            result['parsing_method'] = 'regex_fallback'
        
        return jsonify({
            "success": True,
            "parsed": result,
            "original_input": description
        }), 200
        
    except Exception as e:
        print(f"NLP Parse Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/parse-document', methods=['POST', 'OPTIONS'])
def parse_document():
    """Upload and parse document (PDF, DOCX, TXT) for NLP input"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        if 'file' not in request.files:
            return jsonify({"detail": "No file uploaded"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"detail": "No file selected"}), 400
        
        filename = file.filename.lower()
        extracted_text = ""
        
        # Parse based on file type
        if filename.endswith('.txt'):
            # Plain text file
            extracted_text = file.read().decode('utf-8', errors='ignore')
            
        elif filename.endswith('.pdf'):
            # PDF parsing
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
            except ImportError:
                return jsonify({"detail": "PDF parsing not available. Install PyPDF2."}), 500
            except Exception as e:
                return jsonify({"detail": f"Error parsing PDF: {str(e)}"}), 400
                
        elif filename.endswith('.docx'):
            # Word document parsing
            try:
                from docx import Document
                doc = Document(io.BytesIO(file.read()))
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
                # Also get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        extracted_text += row_text + "\n"
            except ImportError:
                return jsonify({"detail": "DOCX parsing not available. Install python-docx."}), 500
            except Exception as e:
                return jsonify({"detail": f"Error parsing DOCX: {str(e)}"}), 400
                
        elif filename.endswith('.doc'):
            return jsonify({"detail": "Legacy .doc files not supported. Please convert to .docx"}), 400
            
        elif filename.endswith('.csv'):
            # CSV parsing - convert to descriptive text
            try:
                import csv
                content = file.read().decode('utf-8', errors='ignore')
                reader = csv.reader(io.StringIO(content))
                rows = list(reader)
                if rows:
                    headers = rows[0] if rows else []
                    for row in rows[1:]:
                        if row:
                            row_desc = ", ".join(f"{headers[i] if i < len(headers) else 'col'+str(i)}: {val}" for i, val in enumerate(row) if val)
                            extracted_text += row_desc + "\n"
            except Exception as e:
                return jsonify({"detail": f"Error parsing CSV: {str(e)}"}), 400
                
        elif filename.endswith(('.xls', '.xlsx')):
            # Excel parsing
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(file.read()), data_only=True)
                for sheet in wb.worksheets:
                    extracted_text += f"Sheet: {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        if any(cell is not None for cell in row):
                            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                            extracted_text += row_text + "\n"
            except Exception as e:
                return jsonify({"detail": f"Error parsing Excel: {str(e)}"}), 400
        else:
            return jsonify({"detail": "Unsupported file type. Use PDF, DOCX, TXT, CSV, or XLSX."}), 400
        
        # Clean up extracted text
        extracted_text = extracted_text.strip()
        
        if not extracted_text:
            return jsonify({"detail": "No text could be extracted from the document"}), 400
        
        # Optionally parse with NLP if requested
        use_ai = request.form.get('use_ai', 'true').lower() == 'true'
        auto_parse = request.form.get('auto_parse', 'false').lower() == 'true'
        
        result = {
            "success": True,
            "extracted_text": extracted_text,
            "filename": file.filename,
            "char_count": len(extracted_text),
            "word_count": len(extracted_text.split())
        }
        
        # Auto-parse if requested
        if auto_parse:
            if use_ai and groq_client:
                parsed = parse_nlp_with_groq(extracted_text)
            else:
                parsed = parse_nlp_input(extracted_text)
                parsed['parsing_method'] = 'regex_fallback'
            result['parsed'] = parsed
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"Document Parse Error: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# CUSTOM DATASET UPLOAD & MANAGEMENT
# =====================================================

# In-memory storage for custom datasets (per user)
# In production, this would be stored in database
USER_CUSTOM_DATASETS = {}

@app.route('/api/v1/datasets/upload', methods=['POST', 'OPTIONS'])
def upload_dataset():
    """Upload custom material/product dataset"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        user_id = payload['user_id']
        
        data = request.get_json()
        dataset_name = data.get('name', 'Custom Dataset')
        materials = data.get('materials', [])
        
        if not materials:
            return jsonify({"detail": "No materials provided"}), 400
        
        # Validate and process materials
        processed_materials = []
        for mat in materials:
            processed_mat = {
                'name': mat.get('name', 'Unknown'),
                'type': mat.get('type', 'custom'),
                'emission_factor': float(mat.get('emission_factor', 0)),
                'recycled_content': float(mat.get('recycled_content', 0)) if mat.get('recycled_content') else None,
                'region': mat.get('region', 'Global'),
                'source': 'custom_upload',
                'scarcity_score': float(mat.get('scarcity_score', 0)) if mat.get('scarcity_score') else None
            }
            processed_materials.append(processed_mat)
        
        # Store dataset for user
        dataset_id = str(uuid.uuid4())
        if user_id not in USER_CUSTOM_DATASETS:
            USER_CUSTOM_DATASETS[user_id] = {}
        
        USER_CUSTOM_DATASETS[user_id][dataset_id] = {
            'id': dataset_id,
            'name': dataset_name,
            'materials': processed_materials,
            'created_at': datetime.utcnow().isoformat()
        }
        
        return jsonify({
            "success": True,
            "dataset_id": dataset_id,
            "name": dataset_name,
            "materials_count": len(processed_materials),
            "message": f"Successfully uploaded {len(processed_materials)} materials"
        }), 201
        
    except Exception as e:
        print(f"Dataset Upload Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/datasets', methods=['GET', 'OPTIONS'])
def list_datasets():
    """List all custom datasets for the user"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        user_id = payload['user_id']
        
        user_datasets = USER_CUSTOM_DATASETS.get(user_id, {})
        datasets = list(user_datasets.values())
        
        return jsonify(datasets), 200
        
    except Exception as e:
        print(f"List Datasets Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/datasets/<dataset_id>', methods=['GET', 'DELETE', 'OPTIONS'])
def manage_dataset(dataset_id):
    """Get or delete a specific dataset"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        user_id = payload['user_id']
        
        user_datasets = USER_CUSTOM_DATASETS.get(user_id, {})
        
        if dataset_id not in user_datasets:
            return jsonify({"detail": "Dataset not found"}), 404
        
        if request.method == 'GET':
            return jsonify(user_datasets[dataset_id]), 200
        
        elif request.method == 'DELETE':
            del USER_CUSTOM_DATASETS[user_id][dataset_id]
            return jsonify({"message": "Dataset deleted successfully"}), 200
        
    except Exception as e:
        print(f"Manage Dataset Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/materials/combined-library', methods=['GET', 'OPTIONS'])
def combined_material_library():
    """Get combined material library (system + user custom datasets)"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        user_id = payload['user_id']
        
        # System materials (default library) - Base Metals
        system_materials = [
            {"id": "al_primary", "name": "Primary Aluminium", "type": "aluminium_primary", "unit": "kg", "gwp_factor": 12.5, "source": "system", "region": "Global", "category": "base_metal"},
            {"id": "al_secondary", "name": "Secondary Aluminium (Recycled)", "type": "aluminium_secondary", "unit": "kg", "gwp_factor": 0.6, "source": "system", "region": "Global", "category": "base_metal"},
            {"id": "cu_primary", "name": "Primary Copper", "type": "copper_primary", "unit": "kg", "gwp_factor": 3.5, "source": "system", "region": "Global", "category": "base_metal"},
            {"id": "cu_secondary", "name": "Secondary Copper (Recycled)", "type": "copper_secondary", "unit": "kg", "gwp_factor": 0.5, "source": "system", "region": "Global", "category": "base_metal"},
            {"id": "steel_primary", "name": "Virgin Steel", "type": "steel_primary", "unit": "kg", "gwp_factor": 2.1, "source": "system", "region": "Global", "category": "base_metal"},
            {"id": "steel_secondary", "name": "Recycled Steel", "type": "steel_secondary", "unit": "kg", "gwp_factor": 0.4, "source": "system", "region": "Global", "category": "base_metal"},
        ]
        
        # Critical Minerals - Battery Metals
        battery_minerals = [
            {"id": "li_carbonate", "name": "Lithium Carbonate", "type": "lithium_carbonate", "unit": "kg", "gwp_factor": 15.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 85},
            {"id": "li_hydroxide", "name": "Lithium Hydroxide", "type": "lithium_hydroxide", "unit": "kg", "gwp_factor": 18.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 85},
            {"id": "cobalt_sulfate", "name": "Cobalt Sulfate", "type": "cobalt_sulfate", "unit": "kg", "gwp_factor": 10.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 92},
            {"id": "nickel_class1", "name": "Nickel Class 1", "type": "nickel_class1", "unit": "kg", "gwp_factor": 12.5, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 65},
            {"id": "nickel_ferro", "name": "Ferronickel", "type": "nickel_ferronickel", "unit": "kg", "gwp_factor": 8.5, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 55},
            {"id": "manganese", "name": "Manganese Dioxide", "type": "manganese", "unit": "kg", "gwp_factor": 2.8, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 40},
            {"id": "graphite_nat", "name": "Natural Graphite", "type": "graphite", "unit": "kg", "gwp_factor": 4.2, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 70},
            {"id": "graphite_syn", "name": "Synthetic Graphite", "type": "graphite", "unit": "kg", "gwp_factor": 6.5, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 70},
        ]
        
        # Critical Minerals - Rare Earths
        rare_earth_minerals = [
            {"id": "neodymium", "name": "Neodymium (NdFeB Magnets)", "type": "neodymium", "unit": "kg", "gwp_factor": 35.0, "source": "ecoinvent", "region": "Global", "category": "rare_earth", "scarcity_score": 95},
            {"id": "dysprosium", "name": "Dysprosium", "type": "dysprosium", "unit": "kg", "gwp_factor": 45.0, "source": "ecoinvent", "region": "Global", "category": "rare_earth", "scarcity_score": 98},
            {"id": "praseodymium", "name": "Praseodymium", "type": "praseodymium", "unit": "kg", "gwp_factor": 32.0, "source": "ecoinvent", "region": "Global", "category": "rare_earth", "scarcity_score": 90},
            {"id": "terbium", "name": "Terbium", "type": "terbium", "unit": "kg", "gwp_factor": 50.0, "source": "ecoinvent", "region": "Global", "category": "rare_earth", "scarcity_score": 96},
            {"id": "ree_mixed", "name": "Rare Earth Mixed Oxide", "type": "rare_earth_mixed", "unit": "kg", "gwp_factor": 38.0, "source": "ecoinvent", "region": "Global", "category": "rare_earth", "scarcity_score": 93},
        ]
        
        # Critical Minerals - Other Strategic
        other_critical = [
            {"id": "tungsten", "name": "Tungsten Carbide", "type": "tungsten", "unit": "kg", "gwp_factor": 22.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 78},
            {"id": "vanadium", "name": "Vanadium Pentoxide", "type": "vanadium", "unit": "kg", "gwp_factor": 28.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 72},
            {"id": "titanium", "name": "Titanium Sponge", "type": "titanium", "unit": "kg", "gwp_factor": 8.1, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 45},
            {"id": "tantalum", "name": "Tantalum Powder", "type": "tantalum", "unit": "kg", "gwp_factor": 48.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 88},
            {"id": "indium", "name": "Indium (ITO)", "type": "indium", "unit": "kg", "gwp_factor": 142.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 82},
            {"id": "gallium", "name": "Gallium Arsenide", "type": "gallium", "unit": "kg", "gwp_factor": 185.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 80},
            {"id": "germanium", "name": "Germanium", "type": "germanium", "unit": "kg", "gwp_factor": 165.0, "source": "ecoinvent", "region": "Global", "category": "critical_mineral", "scarcity_score": 84},
        ]
        
        # Precious Metals (PGMs)
        precious_metals = [
            {"id": "platinum", "name": "Platinum", "type": "platinum", "unit": "kg", "gwp_factor": 12500.0, "source": "ecoinvent", "region": "Global", "category": "precious_metal", "scarcity_score": 88},
            {"id": "palladium", "name": "Palladium", "type": "palladium", "unit": "kg", "gwp_factor": 9800.0, "source": "ecoinvent", "region": "Global", "category": "precious_metal", "scarcity_score": 86},
            {"id": "silver", "name": "Silver", "type": "silver", "unit": "kg", "gwp_factor": 104.0, "source": "ecoinvent", "region": "Global", "category": "precious_metal", "scarcity_score": 60},
            {"id": "gold", "name": "Gold", "type": "gold", "unit": "kg", "gwp_factor": 31500.0, "source": "ecoinvent", "region": "Global", "category": "precious_metal", "scarcity_score": 95},
        ]
        
        # Joining/Brazing Materials
        joining_materials = [
            {"id": "brazing_alloy", "name": "Brazing Alloy (Silver-based)", "type": "brazing_alloy", "unit": "kg", "gwp_factor": 85.0, "source": "ecoinvent", "region": "Global", "category": "joining_material", "scarcity_score": 50},
            {"id": "solder_lead_free", "name": "Lead-Free Solder (SAC305)", "type": "solder_lead_free", "unit": "kg", "gwp_factor": 25.0, "source": "ecoinvent", "region": "Global", "category": "joining_material", "scarcity_score": 25},
            {"id": "flux", "name": "Brazing/Soldering Flux", "type": "flux", "unit": "kg", "gwp_factor": 3.0, "source": "system", "region": "Global", "category": "joining_material", "scarcity_score": 5},
        ]
        
        # India-specific materials (JNARRDC baseline)
        india_materials = [
            {"id": "al_india_primary", "name": "Primary Aluminium (India)", "type": "aluminium_primary", "unit": "kg", "gwp_factor": 16.5, "source": "jnarrdc", "region": "India", "category": "base_metal"},
            {"id": "al_india_secondary", "name": "Secondary Aluminium (India)", "type": "aluminium_secondary", "unit": "kg", "gwp_factor": 0.8, "source": "jnarrdc", "region": "India", "category": "base_metal"},
            {"id": "cu_india_primary", "name": "Primary Copper (India)", "type": "copper_primary", "unit": "kg", "gwp_factor": 4.2, "source": "jnarrdc", "region": "India", "category": "base_metal"},
            {"id": "cu_india_secondary", "name": "Secondary Copper (India)", "type": "copper_secondary", "unit": "kg", "gwp_factor": 0.6, "source": "jnarrdc", "region": "India", "category": "base_metal"},
            {"id": "steel_india_primary", "name": "Virgin Steel (India)", "type": "steel_primary", "unit": "kg", "gwp_factor": 2.8, "source": "jnarrdc", "region": "India", "category": "base_metal"},
            {"id": "steel_india_secondary", "name": "Recycled Steel (India)", "type": "steel_secondary", "unit": "kg", "gwp_factor": 0.5, "source": "jnarrdc", "region": "India", "category": "base_metal"},
        ]
        
        # User custom materials
        user_datasets = USER_CUSTOM_DATASETS.get(user_id, {})
        custom_materials = []
        for dataset in user_datasets.values():
            for mat in dataset.get('materials', []):
                custom_materials.append({
                    **mat,
                    'id': f"custom_{uuid.uuid4().hex[:8]}",
                    'unit': 'kg',
                    'gwp_factor': mat.get('emission_factor', 0),
                    'dataset_name': dataset.get('name', 'Custom')
                })
        
        # Combine all materials
        all_materials = (system_materials + battery_minerals + rare_earth_minerals + 
                         other_critical + precious_metals + joining_materials + india_materials + custom_materials)
        
        return jsonify({
            "system": system_materials,
            "battery_minerals": battery_minerals,
            "rare_earth": rare_earth_minerals,
            "critical_minerals": other_critical,
            "precious_metals": precious_metals,
            "joining_materials": joining_materials,
            "india": india_materials,
            "custom": custom_materials,
            "all": all_materials
        }), 200
        
    except Exception as e:
        print(f"Combined Library Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/materials/library', methods=['GET', 'OPTIONS'])
def material_library():
    if request.method == 'OPTIONS':
        return '', 200
    
    materials = [
        # Base Metals
        {"id": "al_primary", "name": "Primary Aluminium", "type": "aluminium_primary", "unit": "kg", "gwp_factor": 12.5, "category": "base_metal"},
        {"id": "al_secondary", "name": "Secondary Aluminium (Recycled)", "type": "aluminium_secondary", "unit": "kg", "gwp_factor": 0.6, "category": "base_metal"},
        {"id": "cu_primary", "name": "Primary Copper", "type": "copper_primary", "unit": "kg", "gwp_factor": 3.5, "category": "base_metal"},
        {"id": "cu_secondary", "name": "Secondary Copper (Recycled)", "type": "copper_secondary", "unit": "kg", "gwp_factor": 0.5, "category": "base_metal"},
        {"id": "steel_primary", "name": "Virgin Steel", "type": "steel_primary", "unit": "kg", "gwp_factor": 2.1, "category": "base_metal"},
        {"id": "steel_secondary", "name": "Recycled Steel", "type": "steel_secondary", "unit": "kg", "gwp_factor": 0.4, "category": "base_metal"},
        # Battery Minerals
        {"id": "li_carbonate", "name": "Lithium Carbonate", "type": "lithium_carbonate", "unit": "kg", "gwp_factor": 15.0, "category": "critical_mineral", "scarcity_score": 85},
        {"id": "li_hydroxide", "name": "Lithium Hydroxide", "type": "lithium_hydroxide", "unit": "kg", "gwp_factor": 18.0, "category": "critical_mineral", "scarcity_score": 85},
        {"id": "cobalt_sulfate", "name": "Cobalt Sulfate", "type": "cobalt_sulfate", "unit": "kg", "gwp_factor": 10.0, "category": "critical_mineral", "scarcity_score": 92},
        {"id": "nickel_class1", "name": "Nickel Class 1", "type": "nickel_class1", "unit": "kg", "gwp_factor": 12.5, "category": "critical_mineral", "scarcity_score": 65},
        {"id": "manganese", "name": "Manganese Dioxide", "type": "manganese", "unit": "kg", "gwp_factor": 2.8, "category": "critical_mineral", "scarcity_score": 40},
        {"id": "graphite", "name": "Natural Graphite", "type": "graphite", "unit": "kg", "gwp_factor": 4.2, "category": "critical_mineral", "scarcity_score": 70},
        # Rare Earths
        {"id": "neodymium", "name": "Neodymium", "type": "neodymium", "unit": "kg", "gwp_factor": 35.0, "category": "rare_earth", "scarcity_score": 95},
        {"id": "dysprosium", "name": "Dysprosium", "type": "dysprosium", "unit": "kg", "gwp_factor": 45.0, "category": "rare_earth", "scarcity_score": 98},
        {"id": "praseodymium", "name": "Praseodymium", "type": "praseodymium", "unit": "kg", "gwp_factor": 32.0, "category": "rare_earth", "scarcity_score": 90},
        # Other Critical
        {"id": "tungsten", "name": "Tungsten", "type": "tungsten", "unit": "kg", "gwp_factor": 22.0, "category": "critical_mineral", "scarcity_score": 78},
        {"id": "vanadium", "name": "Vanadium", "type": "vanadium", "unit": "kg", "gwp_factor": 28.0, "category": "critical_mineral", "scarcity_score": 72},
        {"id": "titanium", "name": "Titanium", "type": "titanium", "unit": "kg", "gwp_factor": 8.1, "category": "critical_mineral", "scarcity_score": 45},
        {"id": "tantalum", "name": "Tantalum", "type": "tantalum", "unit": "kg", "gwp_factor": 48.0, "category": "critical_mineral", "scarcity_score": 88},
        # Precious Metals
        {"id": "silver", "name": "Silver", "type": "silver", "unit": "kg", "gwp_factor": 104.0, "category": "precious_metal", "scarcity_score": 60},
        {"id": "gold", "name": "Gold", "type": "gold", "unit": "kg", "gwp_factor": 31500.0, "category": "precious_metal", "scarcity_score": 95},
        {"id": "platinum", "name": "Platinum", "type": "platinum", "unit": "kg", "gwp_factor": 12500.0, "category": "precious_metal", "scarcity_score": 88},
        {"id": "palladium", "name": "Palladium", "type": "palladium", "unit": "kg", "gwp_factor": 9800.0, "category": "precious_metal", "scarcity_score": 86},
        # Joining/Brazing Materials
        {"id": "brazing_alloy", "name": "Brazing Alloy (Silver-based)", "type": "brazing_alloy", "unit": "kg", "gwp_factor": 85.0, "category": "joining_material", "scarcity_score": 50},
        {"id": "solder_lead_free", "name": "Lead-Free Solder (SAC305)", "type": "solder_lead_free", "unit": "kg", "gwp_factor": 25.0, "category": "joining_material", "scarcity_score": 25},
        {"id": "flux", "name": "Brazing/Soldering Flux", "type": "flux", "unit": "kg", "gwp_factor": 3.0, "category": "joining_material", "scarcity_score": 5},
    ]
    
    return jsonify(materials), 200


@app.route('/api/v1/lcia-categories', methods=['GET', 'OPTIONS'])
def get_lcia_categories():
    """Get all LCIA impact categories with metadata"""
    if request.method == 'OPTIONS':
        return '', 200
    
    return jsonify({
        'categories': LCIA_CATEGORIES,
        'grid_factors': INDIAN_GRID_FACTORS,
        'transport_factors': TRANSPORT_EMISSION_FACTORS,
    }), 200


@app.route('/api/v1/projects/<project_id>/calculate-lcia', methods=['POST', 'OPTIONS'])
def calculate_project_lcia_endpoint(project_id):
    """Calculate comprehensive LCIA for a project"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("""SELECT id, name, product_category FROM projects 
                     WHERE id = ? AND user_id = ?""", (project_id, payload['user_id']))
        project = c.fetchone()
        
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("""SELECT material_name, material_type, quantity, unit, recycled_content, 
                     transport_distance FROM project_materials WHERE project_id = ?""", (project_id,))
        materials_rows = c.fetchall()
        conn.close()
        
        if not materials_rows:
            return jsonify({"detail": "No materials found in project"}), 400
        
        # Prepare materials list
        materials = []
        for row in materials_rows:
            materials.append({
                'name': row[0],
                'type': row[1],
                'quantity': row[2] or 0,
                'unit': row[3],
                'recycled_content': row[4] or 0,
                'transport_distance': row[5] or 0,
            })
        
        # Get grid region from request or use default
        data = request.get_json() or {}
        grid_region = data.get('grid_region', 'national_average')
        
        # Calculate LCIA
        lcia_result = calculate_project_lcia(materials, grid_region)
        
        return jsonify({
            'project_id': project_id,
            'project_name': project[1],
            'product_category': project[2],
            **lcia_result
        }), 200
        
    except jwt.ExpiredSignatureError:
        return jsonify({"detail": "Token has expired"}), 401
    except jwt.InvalidTokenError:
        return jsonify({"detail": "Invalid token"}), 401
    except Exception as e:
        print(f"LCIA calculation error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/lcia/calculate-material', methods=['POST', 'OPTIONS'])
def calculate_material_lcia():
    """Calculate LCIA for a single material (for real-time calculation)"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json()
        
        material_type = data.get('material_type', 'unknown')
        quantity = float(data.get('quantity', 1))
        recycled_content = float(data.get('recycled_content', 0))
        transport_distance = float(data.get('transport_distance', 0))
        transport_mode = data.get('transport_mode', 'road_truck')
        grid_region = data.get('grid_region', 'national_average')
        
        impacts = calculate_lcia_impacts(
            material_type, quantity, recycled_content, 
            transport_distance, transport_mode, grid_region
        )
        
        return jsonify({
            'material_type': material_type,
            'quantity': quantity,
            'unit': 'kg',
            'recycled_content': recycled_content,
            'impacts': impacts,
            'categories_metadata': LCIA_CATEGORIES,
        }), 200
        
    except Exception as e:
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/industry-benchmarks', methods=['GET', 'OPTIONS'])
def get_industry_benchmarks():
    """Get industry average benchmarks for comparison"""
    if request.method == 'OPTIONS':
        return '', 200
    
    benchmarks = []
    for category, data in INDUSTRY_BENCHMARKS.items():
        benchmarks.append({
            "category": category,
            "name": data['name'],
            "avg_lifespan": data['avg_lifespan'],
            "avg_mci": data['avg_mci']
        })
    
    return jsonify(benchmarks), 200


@app.route('/api/v1/projects/<project_id>/calculate-mci', methods=['POST', 'OPTIONS'])
def calculate_project_mci(project_id):
    """Calculate MCI and Circular Design Score for a project"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("""SELECT id, name, description, status, product_category, 
                     target_lifespan, is_designed_for_disassembly, user_id, 
                     COALESCE(gwp_total, 0) as gwp_total, created_at 
                     FROM projects WHERE id = ? AND user_id = ?""", (project_id, payload['user_id']))
        project = c.fetchone()
        
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get project materials
        c.execute("""SELECT id, material_name, material_type, quantity, unit, 
                     recycled_content, gwp, transport_distance 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        
        if not materials:
            conn.close()
            return jsonify({"detail": "No materials in project. Add materials first."}), 400
        
        # Calculate averages
        total_mass = sum(m[3] for m in materials)
        avg_recycled_content = sum(m[5] * m[3] for m in materials) / total_mass if total_mass > 0 else 0
        
        # Get project settings
        product_category = project[4] or 'other'
        target_lifespan = project[5] or 10
        is_designed_for_disassembly = bool(project[6])
        
        # Get industry benchmark
        benchmark = INDUSTRY_BENCHMARKS.get(product_category, INDUSTRY_BENCHMARKS['other'])
        industry_avg_lifespan = benchmark['avg_lifespan']
        
        # Assume recycled content output = recycled content input * 0.85 (some loss)
        recycled_content_output = avg_recycled_content * 0.85
        if is_designed_for_disassembly:
            recycled_content_output = min(100, recycled_content_output * 1.15)
        
        # Calculate MCI
        mci_score = calculate_mci(
            recycled_content_input=avg_recycled_content,
            recycled_content_output=recycled_content_output,
            target_lifespan=target_lifespan,
            industry_avg_lifespan=industry_avg_lifespan,
            is_designed_for_disassembly=is_designed_for_disassembly
        )
        
        # Calculate Circular Design Score
        circular_design_score = calculate_circular_design_score(
            mci_score=mci_score,
            target_lifespan=target_lifespan,
            industry_avg_lifespan=industry_avg_lifespan,
            is_designed_for_disassembly=is_designed_for_disassembly,
            avg_recycled_content=avg_recycled_content
        )
        
        # Update project with calculated scores
        c.execute("""UPDATE projects SET mci_score = ?, circular_design_score = ?, 
                     status = 'calculated' WHERE id = ?""", 
                  (mci_score, circular_design_score, project_id))
        conn.commit()
        conn.close()
        
        return jsonify({
            "project_id": project_id,
            "mci_score": mci_score,
            "circular_design_score": circular_design_score,
            "avg_recycled_content": round(avg_recycled_content, 1),
            "recycled_content_output": round(recycled_content_output, 1),
            "target_lifespan": target_lifespan,
            "industry_avg_lifespan": industry_avg_lifespan,
            "product_category": product_category,
            "is_designed_for_disassembly": is_designed_for_disassembly,
            "total_materials": len(materials),
            "total_mass": round(total_mass, 2),
            "benchmark": {
                "name": benchmark['name'],
                "avg_mci": benchmark['avg_mci']
            }
        }), 200
    except Exception as e:
        print(f"Error calculating MCI: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/scenario', methods=['POST', 'OPTIONS'])
def calculate_scenario(project_id):
    """Calculate 'What-if' scenario with modified parameters"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        data = request.get_json()
        
        # Scenario parameters (with defaults from original)
        recycled_content_modifier = data.get('recycled_content_modifier', 0)  # % change
        lifespan_modifier = data.get('lifespan_modifier', 0)  # years change
        transport_reduction = data.get('transport_reduction', 0)  # % reduction
        design_for_disassembly = data.get('design_for_disassembly', None)  # Override
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("""SELECT id, name, product_category, target_lifespan, 
                     is_designed_for_disassembly, COALESCE(gwp_total, 0) as gwp_total
                     FROM projects WHERE id = ? AND user_id = ?""", (project_id, payload['user_id']))
        project = c.fetchone()
        
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("""SELECT material_type, quantity, recycled_content, transport_distance, gwp 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        if not materials:
            return jsonify({"detail": "No materials in project"}), 400
        
        # Original values
        original_gwp = project[5]
        original_lifespan = project[3] or 10
        original_disassembly = bool(project[4])
        
        # Scenario values
        scenario_lifespan = original_lifespan + lifespan_modifier
        scenario_disassembly = design_for_disassembly if design_for_disassembly is not None else original_disassembly
        
        # Recalculate GWP with scenario
        scenario_gwp = 0
        total_mass = 0
        avg_recycled_scenario = 0
        
        for mat in materials:
            mat_type, quantity, recycled_content, transport, original_mat_gwp = mat
            
            # Apply modifiers
            new_recycled = min(100, max(0, recycled_content + recycled_content_modifier))
            new_transport = transport * (1 - transport_reduction / 100)
            
            # Recalculate GWP
            mat_gwp = calculate_gwp(mat_type, quantity, new_recycled, new_transport)
            scenario_gwp += mat_gwp
            total_mass += quantity
            avg_recycled_scenario += new_recycled * quantity
        
        avg_recycled_scenario = avg_recycled_scenario / total_mass if total_mass > 0 else 0
        
        # Get benchmark
        product_category = project[2] or 'other'
        benchmark = INDUSTRY_BENCHMARKS.get(product_category, INDUSTRY_BENCHMARKS['other'])
        
        # Calculate scenario MCI
        recycled_output_scenario = avg_recycled_scenario * 0.85
        if scenario_disassembly:
            recycled_output_scenario = min(100, recycled_output_scenario * 1.15)
        
        scenario_mci = calculate_mci(
            recycled_content_input=avg_recycled_scenario,
            recycled_content_output=recycled_output_scenario,
            target_lifespan=scenario_lifespan,
            industry_avg_lifespan=benchmark['avg_lifespan'],
            is_designed_for_disassembly=scenario_disassembly
        )
        
        scenario_circular_score = calculate_circular_design_score(
            mci_score=scenario_mci,
            target_lifespan=scenario_lifespan,
            industry_avg_lifespan=benchmark['avg_lifespan'],
            is_designed_for_disassembly=scenario_disassembly,
            avg_recycled_content=avg_recycled_scenario
        )
        
        # Calculate improvements
        gwp_improvement = ((original_gwp - scenario_gwp) / original_gwp * 100) if original_gwp > 0 else 0
        
        return jsonify({
            "original": {
                "gwp_total": original_gwp,
                "lifespan": original_lifespan,
                "is_designed_for_disassembly": original_disassembly
            },
            "scenario": {
                "gwp_total": round(scenario_gwp, 2),
                "lifespan": scenario_lifespan,
                "is_designed_for_disassembly": scenario_disassembly,
                "avg_recycled_content": round(avg_recycled_scenario, 1),
                "mci_score": scenario_mci,
                "circular_design_score": scenario_circular_score
            },
            "improvements": {
                "gwp_reduction_percent": round(gwp_improvement, 1),
                "gwp_reduction_kg": round(original_gwp - scenario_gwp, 2)
            },
            "modifiers_applied": {
                "recycled_content_modifier": recycled_content_modifier,
                "lifespan_modifier": lifespan_modifier,
                "transport_reduction": transport_reduction,
                "design_for_disassembly": design_for_disassembly
            }
        }), 200
    except Exception as e:
        print(f"Error calculating scenario: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# ENGINE 4: DESIGN OPTIMIZATION & AI ADVISOR
# =====================================================

# Alloy specifications and max recycled content limits
ALLOY_SPECIFICATIONS = {
    'aluminium_primary': {
        'name': 'Primary Aluminium',
        'alloys': {
            'Al-1100': {'max_recycled': 0, 'impurity_limit': 0.05},
            'Al-6061': {'max_recycled': 85, 'impurity_limit': 1.0},
            'Al-6063': {'max_recycled': 80, 'impurity_limit': 0.7},
            'Al-7075': {'max_recycled': 40, 'impurity_limit': 0.5},
        },
        'general_max_recycled': 60,
        'alternative': 'aluminium_secondary',
        'gwp_reduction_with_recycled': 95
    },
    'aluminium_secondary': {
        'name': 'Secondary Aluminium',
        'general_max_recycled': 100,
        'already_recycled': True
    },
    'copper_primary': {
        'name': 'Primary Copper',
        'alloys': {
            'C11000': {'max_recycled': 30, 'impurity_limit': 0.01},  # ETP copper, very pure
            'C26000': {'max_recycled': 70, 'impurity_limit': 0.5},   # Brass
            'C28000': {'max_recycled': 75, 'impurity_limit': 0.6},
            'Berry': {'max_recycled': 90, 'impurity_limit': 1.0},    # Scrap grade
        },
        'general_max_recycled': 50,
        'alternative': 'copper_secondary',
        'gwp_reduction_with_recycled': 85
    },
    'copper_secondary': {
        'name': 'Secondary Copper',
        'general_max_recycled': 100,
        'already_recycled': True
    },
    'steel_primary': {
        'name': 'Virgin Steel',
        'alloys': {
            'Mild Steel': {'max_recycled': 95, 'impurity_limit': 0.3},
            'Stainless 304': {'max_recycled': 80, 'impurity_limit': 0.2},
            'Stainless 316': {'max_recycled': 75, 'impurity_limit': 0.15},
            'High Carbon': {'max_recycled': 60, 'impurity_limit': 0.1},
        },
        'general_max_recycled': 80,
        'alternative': 'steel_secondary',
        'gwp_reduction_with_recycled': 80
    },
    'steel_secondary': {
        'name': 'Recycled Steel',
        'general_max_recycled': 100,
        'already_recycled': True
    },
    'lithium': {
        'name': 'Lithium Carbonate',
        'general_max_recycled': 25,
        'recycling_nascent': True,
        'scarcity_warning': True
    },
    'cobalt': {
        'name': 'Cobalt Sulfate',
        'general_max_recycled': 40,
        'scarcity_warning': True
    },
    'nickel': {
        'name': 'Primary Nickel',
        'alloys': {
            'Class 1': {'max_recycled': 60, 'impurity_limit': 0.01},
            'Ferronickel': {'max_recycled': 80, 'impurity_limit': 0.3},
        },
        'general_max_recycled': 60
    }
}

# Material alternatives for lifespan extension
LIFESPAN_ALTERNATIVES = {
    'aluminium_primary': {
        'corrosion_resistant': ['Al-6061', 'Al-6063'],
        'surface_treatments': ['Anodizing', 'Powder Coating', 'Chromate Conversion'],
        'lifespan_increase': 5
    },
    'steel_primary': {
        'corrosion_resistant': ['Stainless 304', 'Stainless 316'],
        'surface_treatments': ['Galvanizing', 'Zinc Plating', 'E-coating'],
        'lifespan_increase': 10
    },
    'copper_primary': {
        'corrosion_resistant': ['Brass', 'Bronze'],
        'surface_treatments': ['Tin Plating', 'Nickel Plating'],
        'lifespan_increase': 3
    }
}


def generate_design_recommendations(project_data, materials_data):
    """
    Engine 4: AI Design Advisor
    Generates recommendations for improving circularity and reducing environmental impact
    """
    recommendations = []
    priority_score = 0
    
    total_mass = sum(m[4] for m in materials_data) if materials_data else 0
    total_gwp = sum(m[7] for m in materials_data) if materials_data else 0
    
    for mat in materials_data:
        mat_id, proj_id, mat_name, mat_type, quantity, unit, recycled_content, gwp, transport_dist, created = mat
        
        # Check for recycled content optimization
        spec = ALLOY_SPECIFICATIONS.get(mat_type, {})
        max_recycled = spec.get('general_max_recycled', 50)
        current_recycled = recycled_content or 0
        
        if current_recycled < max_recycled and not spec.get('already_recycled'):
            potential_increase = max_recycled - current_recycled
            gwp_factor = EMISSION_FACTORS.get(mat_type, {}).get('virgin', 1)
            recycled_factor = EMISSION_FACTORS.get(mat_type, {}).get('recycled', gwp_factor * 0.1)
            
            # Calculate potential GWP savings
            current_gwp = quantity * (
                (current_recycled / 100) * recycled_factor +
                ((100 - current_recycled) / 100) * gwp_factor
            )
            optimized_gwp = quantity * (
                (max_recycled / 100) * recycled_factor +
                ((100 - max_recycled) / 100) * gwp_factor
            )
            gwp_savings = current_gwp - optimized_gwp
            gwp_savings_percent = (gwp_savings / current_gwp * 100) if current_gwp > 0 else 0
            
            if potential_increase >= 10:  # Only recommend if significant increase possible
                recommendations.append({
                    'type': 'recycled_content',
                    'priority': 'high' if gwp_savings_percent > 20 else 'medium',
                    'material': mat_name,
                    'material_type': mat_type,
                    'title': f'Increase Recycled Content in {mat_name}',
                    'description': f'Based on impurity limits for {mat_type}, you can increase recycled content from {current_recycled:.0f}% to {max_recycled:.0f}% without affecting material properties.',
                    'current_value': current_recycled,
                    'recommended_value': max_recycled,
                    'impact': {
                        'gwp_savings_kg': round(gwp_savings, 2),
                        'gwp_savings_percent': round(gwp_savings_percent, 1),
                        'cost_impact': 'neutral_to_positive'  # Recycled typically cheaper
                    },
                    'confidence': 0.85
                })
                priority_score += gwp_savings_percent
        
        # Check for material substitution opportunities
        alternative = spec.get('alternative')
        if alternative and current_recycled < 30:
            recommendations.append({
                'type': 'material_substitution',
                'priority': 'medium',
                'material': mat_name,
                'material_type': mat_type,
                'title': f'Consider {ALLOY_SPECIFICATIONS.get(alternative, {}).get("name", alternative)} instead',
                'description': f'Switching to secondary/recycled material could reduce GWP by up to {spec.get("gwp_reduction_with_recycled", 80)}%.',
                'alternative_type': alternative,
                'impact': {
                    'gwp_savings_percent': spec.get('gwp_reduction_with_recycled', 80),
                    'cost_impact': 'positive'
                },
                'confidence': 0.75
            })
        
        # Transport optimization
        if transport_dist and transport_dist > 300:
            transport_gwp = transport_dist * 0.0001 * quantity  # Simplified transport emission
            recommendations.append({
                'type': 'transport_optimization',
                'priority': 'low' if transport_dist < 500 else 'medium',
                'material': mat_name,
                'title': 'Optimize Supply Chain',
                'description': f'{mat_name} is transported {transport_dist}km. Consider local sourcing to reduce transport emissions.',
                'current_distance': transport_dist,
                'recommended_distance': 100,
                'impact': {
                    'gwp_savings_kg': round(transport_gwp * 0.7, 2),
                    'gwp_savings_percent': round(transport_gwp / (gwp or 1) * 70, 1)
                },
                'confidence': 0.6
            })
        
        # Scarcity warnings for critical minerals
        if spec.get('scarcity_warning'):
            recommendations.append({
                'type': 'scarcity_alert',
                'priority': 'high',
                'material': mat_name,
                'material_type': mat_type,
                'title': f'Critical Mineral: {mat_name}',
                'description': f'{mat_name} is a critical mineral with supply constraints. Consider reduction strategies or alternative chemistries.',
                'impact': {
                    'abiotic_depletion': 'high',
                    'supply_risk': 'elevated'
                },
                'confidence': 0.9
            })
    
    # Project-level recommendations
    project_id, name, desc, status, category, lifespan, disassembly, user_id, gwp_total, mci, cds, created = project_data
    
    # Lifespan extension
    if lifespan:
        benchmark = INDUSTRY_BENCHMARKS.get(category or 'other', {})
        avg_lifespan = benchmark.get('avg_lifespan', 10)
        
        if lifespan < avg_lifespan:
            recommendations.append({
                'type': 'lifespan_extension',
                'priority': 'high',
                'title': 'Extend Product Lifespan',
                'description': f'Current target lifespan ({lifespan}y) is below industry average ({avg_lifespan}y). Consider material upgrades or surface treatments.',
                'current_lifespan': lifespan,
                'recommended_lifespan': avg_lifespan + 5,
                'impact': {
                    'mci_improvement': 15,
                    'lifetime_gwp_reduction_percent': round((1 - lifespan / (avg_lifespan + 5)) * 100, 1)
                },
                'suggestions': [
                    'Use corrosion-resistant alloys',
                    'Apply protective coatings',
                    'Design for repairability'
                ],
                'confidence': 0.8
            })
    
    # Design for disassembly
    if not disassembly:
        recommendations.append({
            'type': 'design_for_disassembly',
            'priority': 'medium',
            'title': 'Enable Design for Disassembly',
            'description': 'Designing for easy disassembly improves end-of-life material recovery and MCI score.',
            'impact': {
                'mci_improvement': 10,
                'recycled_output_improvement': 20
            },
            'suggestions': [
                'Use mechanical fasteners instead of adhesives',
                'Mark materials with recycling codes',
                'Create disassembly instructions',
                'Use modular component design'
            ],
            'confidence': 0.85
        })
    
    # Sort by priority
    priority_order = {'high': 0, 'medium': 1, 'low': 2}
    recommendations.sort(key=lambda x: priority_order.get(x.get('priority', 'low'), 3))
    
    return {
        'recommendations': recommendations[:10],  # Top 10 recommendations
        'total_recommendations': len(recommendations),
        'priority_score': round(priority_score, 1),
        'summary': {
            'high_priority': len([r for r in recommendations if r.get('priority') == 'high']),
            'medium_priority': len([r for r in recommendations if r.get('priority') == 'medium']),
            'low_priority': len([r for r in recommendations if r.get('priority') == 'low'])
        },
        'source': 'rule_based'
    }


def generate_groq_design_insights(project_data, materials_data, rule_based_recommendations):
    """
    Enhanced AI Design Advisor using Groq LLM
    Generates additional AI-powered insights alongside rule-based recommendations
    """
    if not groq_client:
        return None
    
    try:
        # Prepare context for Groq
        project_id, name, desc, status, category, lifespan, disassembly, user_id, gwp_total, mci, cds, created = project_data
        
        materials_summary = []
        for mat in materials_data:
            mat_id, proj_id, mat_name, mat_type, quantity, unit, recycled_content, gwp, transport_dist, created = mat
            materials_summary.append({
                'name': mat_name,
                'type': mat_type,
                'quantity': f"{quantity} {unit}",
                'recycled_content': f"{recycled_content or 0}%",
                'gwp': f"{gwp or 0} kg CO2e"
            })
        
        # Get top rule-based recommendations for context
        top_rules = [r['title'] for r in rule_based_recommendations[:5]]
        
        prompt = f"""Analyze this LCA project and provide strategic sustainability insights:

PROJECT: {name}
CATEGORY: {category or 'General'}
DESCRIPTION: {desc or 'No description'}
TARGET LIFESPAN: {lifespan or 'Not specified'} years
DESIGNED FOR DISASSEMBLY: {'Yes' if disassembly else 'No'}
CURRENT GWP: {gwp_total or 0} kg CO2e
CURRENT MCI: {mci or 0}

MATERIALS:
{chr(10).join([f"- {m['name']}: {m['quantity']}, {m['recycled_content']} recycled, GWP: {m['gwp']}" for m in materials_summary])}

RULE-BASED RECOMMENDATIONS ALREADY IDENTIFIED:
{chr(10).join([f"- {r}" for r in top_rules])}

Provide 3-5 ADDITIONAL strategic insights that complement the rule-based recommendations. Focus on:
1. Industry-specific best practices for {category or 'this product'}
2. Emerging technologies or materials that could improve sustainability
3. Supply chain or circular economy opportunities
4. Regulatory compliance considerations (CBAM, EPR, etc.)
5. Cost-benefit analysis of sustainability improvements

Format as JSON array:
[
  {{
    "title": "Insight title",
    "description": "Detailed description (2-3 sentences)",
    "category": "technology|supply_chain|regulatory|cost_benefit|circular_economy",
    "impact_potential": "high|medium|low",
    "implementation_timeframe": "short_term|medium_term|long_term"
  }}
]

Output ONLY the JSON array, no other text."""

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You are an expert LCA consultant specializing in sustainable materials and circular economy for the metals industry. Provide actionable, specific insights."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=1500
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Clean up response - remove markdown code blocks if present
        if response_text.startswith('```'):
            response_text = response_text.split('```')[1]
            if response_text.startswith('json'):
                response_text = response_text[4:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        import json
        insights = json.loads(response_text)
        
        return {
            'ai_insights': insights,
            'model': GROQ_MODEL,
            'source': 'groq_ai'
        }
        
    except Exception as e:
        print(f"Groq Design Advisor Error: {e}")
        return None


@app.route('/api/v1/projects/<project_id>/recommendations', methods=['GET', 'OPTIONS'])
def get_design_recommendations(project_id):
    """Get AI-powered design optimization recommendations for a project"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Not authenticated"}), 401
        
        token = auth_header.split(' ')[1]
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        
        # Check if AI enhancement is requested (default: True if Groq available)
        use_ai = request.args.get('use_ai', 'true').lower() == 'true'
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("""SELECT id, name, description, status, product_category, 
                     target_lifespan, is_designed_for_disassembly, user_id, 
                     COALESCE(gwp_total, 0), COALESCE(mci_score, 0), 
                     COALESCE(circular_design_score, 0), created_at 
                     FROM projects WHERE id = ? AND user_id = ?""", (project_id, payload['user_id']))
        project = c.fetchone()
        
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("""SELECT id, project_id, material_name, material_type, quantity, 
                     unit, recycled_content, gwp, transport_distance, created_at 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        if not materials:
            return jsonify({
                "recommendations": [],
                "total_recommendations": 0,
                "message": "Add materials to get AI recommendations"
            }), 200
        
        # Generate rule-based recommendations (always runs - reliable fallback)
        result = generate_design_recommendations(project, materials)
        result['project_id'] = project_id
        result['project_name'] = project[1]
        
        # Try to enhance with Groq AI insights if enabled
        if use_ai and groq_client:
            ai_result = generate_groq_design_insights(project, materials, result['recommendations'])
            if ai_result:
                result['ai_insights'] = ai_result.get('ai_insights', [])
                result['ai_model'] = ai_result.get('model')
                result['source'] = 'hybrid'  # Rule-based + AI enhanced
            else:
                result['ai_insights'] = []
                result['source'] = 'rule_based'
        else:
            result['ai_insights'] = []
            result['source'] = 'rule_based'
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"Recommendations Error: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# ENGINE 5: ANALYTICS & CHARTS DATA
# =====================================================

@app.route('/api/v1/projects/<project_id>/analytics', methods=['GET', 'OPTIONS'])
def get_project_analytics(project_id):
    """Get comprehensive analytics data for charts and visualizations"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials with all details
        c.execute("""SELECT id, material_name, material_type, quantity, unit, 
                     recycled_content, gwp, transport_distance 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        # Calculate analytics
        total_gwp = sum(m[6] or 0 for m in materials)
        total_mass = sum(m[3] or 0 for m in materials)
        avg_recycled = sum(m[5] or 0 for m in materials) / len(materials) if materials else 0
        
        # GWP breakdown by material
        gwp_by_material = []
        for m in materials:
            gwp_by_material.append({
                'name': m[1],
                'type': m[2],
                'gwp': round(m[6] or 0, 2),
                'percentage': round((m[6] or 0) / total_gwp * 100, 1) if total_gwp > 0 else 0
            })
        gwp_by_material.sort(key=lambda x: x['gwp'], reverse=True)
        
        # GWP by material type
        gwp_by_type = {}
        for m in materials:
            mat_type = m[2] or 'Unknown'
            gwp_by_type[mat_type] = gwp_by_type.get(mat_type, 0) + (m[6] or 0)
        
        type_breakdown = [{'name': k, 'value': round(v, 2)} for k, v in gwp_by_type.items()]
        type_breakdown.sort(key=lambda x: x['value'], reverse=True)
        
        # Recycled content analysis
        recycled_analysis = []
        for m in materials:
            recycled_analysis.append({
                'name': m[1],
                'recycled_content': m[5] or 0,
                'quantity': m[3] or 0,
                'gwp': m[6] or 0
            })
        
        # Calculate MCI (Material Circularity Indicator)
        # MCI = 1 - LFI × (1 - Rc × Ru) where LFI = Linear Flow Index
        # Simplified calculation based on recycled content and recyclability
        recyclability_scores = {
            'aluminium_primary': 0.95, 'aluminium_secondary': 0.95,
            'copper_primary': 0.90, 'copper_secondary': 0.90,
            'steel_primary': 0.85, 'steel_secondary': 0.85,
            'plastic_primary': 0.30, 'plastic_recycled': 0.40,
            'glass': 0.80, 'concrete': 0.60, 'wood': 0.50
        }
        
        total_mci_weight = 0
        weighted_mci = 0
        for m in materials:
            mat_type = m[2] or 'unknown'
            recyclability = recyclability_scores.get(mat_type, 0.5)
            recycled_input = (m[5] or 0) / 100  # Convert to fraction
            
            # Ellen MacArthur Foundation MCI formula (simplified)
            linear_flow = (1 - recycled_input) * (1 - recyclability)
            utility_factor = 1.0  # Assume standard utility
            material_mci = 1 - linear_flow * utility_factor
            
            quantity = m[3] or 0
            weighted_mci += material_mci * quantity
            total_mci_weight += quantity
        
        mci_score = round(weighted_mci / total_mci_weight, 3) if total_mci_weight > 0 else 0
        
        # MCI breakdown by material
        mci_breakdown = []
        for m in materials:
            mat_type = m[2] or 'unknown'
            recyclability = recyclability_scores.get(mat_type, 0.5)
            recycled_input = (m[5] or 0) / 100
            linear_flow = (1 - recycled_input) * (1 - recyclability)
            material_mci = 1 - linear_flow
            
            mci_breakdown.append({
                'name': m[1],
                'mci': round(material_mci, 3),
                'recycled_input': m[5] or 0,
                'recyclability': round(recyclability * 100, 0)
            })
        
        # Lifecycle stage breakdown (estimated distribution)
        lifecycle_stages = [
            {'stage': 'Raw Material Extraction', 'gwp': round(total_gwp * 0.35, 2), 'percentage': 35},
            {'stage': 'Processing & Manufacturing', 'gwp': round(total_gwp * 0.30, 2), 'percentage': 30},
            {'stage': 'Transport', 'gwp': round(total_gwp * 0.10, 2), 'percentage': 10},
            {'stage': 'Use Phase', 'gwp': round(total_gwp * 0.15, 2), 'percentage': 15},
            {'stage': 'End of Life', 'gwp': round(total_gwp * 0.10, 2), 'percentage': 10}
        ]
        
        # Process flow data for Sankey diagram
        process_flow = {
            'nodes': [
                {'id': 'raw_materials', 'name': 'Raw Materials'},
                {'id': 'recycled_input', 'name': 'Recycled Input'},
                {'id': 'manufacturing', 'name': 'Manufacturing'},
                {'id': 'use_phase', 'name': 'Use Phase'},
                {'id': 'end_of_life', 'name': 'End of Life'},
                {'id': 'recycling', 'name': 'Recycling'},
                {'id': 'waste', 'name': 'Waste'}
            ],
            'links': [
                {'source': 'raw_materials', 'target': 'manufacturing', 'value': round(total_mass * (1 - avg_recycled/100), 2)},
                {'source': 'recycled_input', 'target': 'manufacturing', 'value': round(total_mass * (avg_recycled/100), 2)},
                {'source': 'manufacturing', 'target': 'use_phase', 'value': round(total_mass * 0.95, 2)},
                {'source': 'use_phase', 'target': 'end_of_life', 'value': round(total_mass * 0.90, 2)},
                {'source': 'end_of_life', 'target': 'recycling', 'value': round(total_mass * 0.90 * 0.7, 2)},
                {'source': 'end_of_life', 'target': 'waste', 'value': round(total_mass * 0.90 * 0.3, 2)}
            ]
        }
        
        # Circular design score calculation
        disassembly_bonus = 15 if project[6] else 0  # is_designed_for_disassembly
        lifespan_bonus = min((project[5] or 0) / 2, 10)  # target_lifespan bonus
        circular_score = min(100, round(mci_score * 70 + disassembly_bonus + lifespan_bonus, 1))
        
        # Calculate scarcity metrics for critical minerals
        scarcity_analysis = []
        total_scarcity_weight = 0
        weighted_scarcity = 0
        for m in materials:
            mat_type = m[2] or 'unknown'
            scarcity = get_scarcity_score(mat_type)
            quantity = m[3] or 0
            if scarcity > 0:
                scarcity_analysis.append({
                    'name': m[1],
                    'type': mat_type,
                    'scarcity_score': scarcity,
                    'quantity': quantity,
                    'risk_level': 'Critical' if scarcity >= 85 else ('High' if scarcity >= 70 else ('Medium' if scarcity >= 50 else 'Low'))
                })
                weighted_scarcity += scarcity * quantity
                total_scarcity_weight += quantity
        
        avg_scarcity = round(weighted_scarcity / total_scarcity_weight, 1) if total_scarcity_weight > 0 else 0
        critical_mineral_count = len([s for s in scarcity_analysis if s['scarcity_score'] >= 70])
        
        return jsonify({
            'summary': {
                'total_gwp': round(total_gwp, 2),
                'total_mass': round(total_mass, 2),
                'material_count': len(materials),
                'avg_recycled_content': round(avg_recycled, 1),
                'mci_score': mci_score,
                'circular_design_score': circular_score,
                'avg_scarcity_score': avg_scarcity,
                'critical_mineral_count': critical_mineral_count
            },
            'gwp_by_material': gwp_by_material,
            'gwp_by_type': type_breakdown,
            'recycled_analysis': recycled_analysis,
            'mci_breakdown': mci_breakdown,
            'lifecycle_stages': lifecycle_stages,
            'process_flow': process_flow,
            'scarcity_analysis': scarcity_analysis
        }), 200
        
    except Exception as e:
        print(f"Analytics Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/dashboard/analytics', methods=['GET', 'OPTIONS'])
def get_dashboard_analytics():
    """Get aggregate analytics across all user projects"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Get user from token
        auth_header = request.headers.get('Authorization')
        if not auth_header or not auth_header.startswith('Bearer '):
            return jsonify({"detail": "Authorization required"}), 401
        
        token = auth_header.split(' ')[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
            user_id = payload.get('user_id')
        except:
            return jsonify({"detail": "Invalid token"}), 401
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Helper function for safe float conversion
        def safe_float(val):
            """Safely convert value to float"""
            if val is None:
                return 0.0
            try:
                return float(val)
            except (ValueError, TypeError):
                return 0.0
        
        # First, recalculate MCI for any projects that have materials but no MCI score
        c.execute("""SELECT id, target_lifespan, is_designed_for_disassembly, product_category 
                     FROM projects WHERE user_id = ? AND (mci_score IS NULL OR mci_score = 0)""", (user_id,))
        projects_to_update = c.fetchall()
        
        for proj in projects_to_update:
            proj_id, target_lifespan, is_disassembly, category = proj
            target_lifespan = safe_float(target_lifespan) or 10
            is_disassembly = bool(is_disassembly)
            category = category or 'other'
            
            # Get materials for this project
            c.execute("""SELECT material_type, quantity, recycled_content 
                         FROM project_materials WHERE project_id = ?""", (proj_id,))
            materials = c.fetchall()
            
            if materials:
                total_mass = sum(safe_float(m[1]) for m in materials)
                if total_mass > 0:
                    weighted_recycled = sum(safe_float(m[1]) * safe_float(m[2]) for m in materials) / total_mass
                    benchmark = INDUSTRY_BENCHMARKS.get(category, INDUSTRY_BENCHMARKS['other'])
                    industry_avg_lifespan = benchmark['avg_lifespan']
                    
                    mci_score = calculate_mci(
                        weighted_recycled,
                        weighted_recycled * 0.8,
                        target_lifespan,
                        industry_avg_lifespan,
                        is_disassembly
                    )
                    
                    c.execute("UPDATE projects SET mci_score = ? WHERE id = ?", (mci_score, proj_id))
        
        conn.commit()
        
        # Get all user projects (after potential updates)
        c.execute("""SELECT id, name, status, gwp_total, mci_score, circular_design_score, created_at 
                     FROM projects WHERE user_id = ? ORDER BY created_at DESC""", (user_id,))
        projects = c.fetchall()
        
        # Get all materials for user's projects
        project_ids = [p[0] for p in projects]
        if project_ids:
            placeholders = ','.join('?' * len(project_ids))
            c.execute(f"""SELECT project_id, material_type, SUM(quantity), SUM(gwp), AVG(recycled_content)
                         FROM project_materials WHERE project_id IN ({placeholders})
                         GROUP BY project_id, material_type""", project_ids)
            material_stats = c.fetchall()
        else:
            material_stats = []
        
        conn.close()
        
        # Calculate totals with proper type handling (safe_float defined above)
        total_gwp = sum(safe_float(p[3]) for p in projects)
        avg_mci = sum(safe_float(p[4]) for p in projects) / len(projects) if projects else 0
        avg_circular = sum(safe_float(p[5]) for p in projects) / len(projects) if projects else 0
        
        # Projects over time (for trend chart)
        projects_timeline = []
        for p in projects:
            projects_timeline.append({
                'id': p[0],
                'name': p[1],
                'status': p[2],
                'gwp': safe_float(p[3]),
                'mci': safe_float(p[4]),
                'created_at': p[6]
            })
        
        # Material type distribution across all projects
        type_totals = {}
        for stat in material_stats:
            mat_type = stat[1] or 'Unknown'
            type_totals[mat_type] = type_totals.get(mat_type, 0) + safe_float(stat[3])
        
        material_distribution = [{'name': k, 'value': round(v, 2)} for k, v in type_totals.items()]
        material_distribution.sort(key=lambda x: x['value'], reverse=True)
        
        return jsonify({
            'summary': {
                'total_projects': len(projects),
                'calculated_projects': len([p for p in projects if p[2] in ['calculated', 'verified']]),
                'total_gwp': round(total_gwp, 2),
                'avg_mci': round(avg_mci, 3),
                'avg_circular_score': round(avg_circular, 1)
            },
            'projects_timeline': projects_timeline,
            'material_distribution': material_distribution
        }), 200
        
    except Exception as e:
        print(f"Dashboard Analytics Error: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# ENGINE 6: GROQ AI CHAT ASSISTANT
# =====================================================

def get_ai_response(prompt, context=""):
    """Get AI response using Groq or fallback to rule-based"""
    
    if groq_client:
        try:
            system_prompt = """You are JNARRDC LCA Portal AI Assistant, an expert in Life Cycle Assessment (LCA) 
for metals and materials. You help users understand:
- Environmental impacts (GWP, carbon footprint)
- Material Circularity Indicator (MCI) and circular economy principles
- Recommendations for reducing environmental impact
- CBAM (Carbon Border Adjustment Mechanism) compliance
- Best practices for sustainable material selection

Keep responses concise, factual, and actionable. Use specific numbers when available.
Focus on practical recommendations for the Indian metals industry."""

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Context: {context}\n\nQuestion: {prompt}"}
            ]
            
            response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=messages,
                temperature=0.7,
                max_tokens=1024
            )
            
            return {
                'response': response.choices[0].message.content,
                'model': GROQ_MODEL,
                'source': 'groq_ai'
            }
        except Exception as e:
            print(f"Groq API error: {e}")
    
    # Fallback to rule-based responses
    prompt_lower = prompt.lower()
    
    if 'mci' in prompt_lower or 'circularity' in prompt_lower:
        response = """The Material Circularity Indicator (MCI) measures how circular a product is on a scale of 0 to 1:
- MCI = 0: Completely linear (virgin materials, no recycling)
- MCI = 1: Fully circular (100% recycled input and recyclable output)

To improve MCI:
1. Increase recycled content in materials
2. Design for disassembly and recyclability
3. Extend product lifespan
4. Choose highly recyclable materials (metals are excellent - aluminum: 95%, steel: 85%)"""
    
    elif 'gwp' in prompt_lower or 'carbon' in prompt_lower or 'emission' in prompt_lower:
        response = """Global Warming Potential (GWP) measures greenhouse gas emissions in kg CO₂-equivalent.

Key GWP reduction strategies:
1. Use secondary (recycled) metals - reduces GWP by 85-95%
2. Optimize transport distances
3. Increase process efficiency
4. Switch to renewable energy in manufacturing
5. Design for longer product lifespan

Primary vs Secondary metals GWP:
- Primary Aluminum: 16.5 kg CO₂-eq/kg
- Secondary Aluminum: 0.7 kg CO₂-eq/kg (95% reduction!)
- Primary Steel: 2.3 kg CO₂-eq/kg
- Secondary Steel: 0.7 kg CO₂-eq/kg"""
    
    elif 'cbam' in prompt_lower:
        response = """CBAM (Carbon Border Adjustment Mechanism) is the EU's carbon pricing system for imports.

Key points for Indian exporters:
1. Applies to: Iron, steel, aluminum, cement, fertilizers, electricity, hydrogen
2. Effective: October 2023 (reporting), January 2026 (full charges)
3. Requires: Accurate GWP data per product
4. JNARRDC LCA Portal helps: Generate CBAM-compliant reports with verified emission data

To prepare:
- Track emissions at product level
- Document recycled content percentages
- Maintain supply chain transparency"""
    
    else:
        response = """I'm the JNARRDC LCA Portal AI Assistant. I can help you with:

🌱 **Environmental Impact** - Understanding GWP and carbon footprint
📊 **Circularity Metrics** - MCI calculations and improvement strategies  
🔄 **Material Selection** - Comparing primary vs secondary metals
📋 **Compliance** - CBAM reporting requirements
💡 **Recommendations** - Actionable steps to reduce environmental impact

What would you like to know more about?"""
    
    return {
        'response': response,
        'model': 'rule-based',
        'source': 'fallback'
    }


@app.route('/api/v1/ai/chat', methods=['POST', 'OPTIONS'])
def ai_chat():
    """AI Chat endpoint using Groq or fallback"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        context = data.get('context', '')
        project_id = data.get('project_id')
        
        if not prompt:
            return jsonify({"detail": "Prompt is required"}), 400
        
        # Build context from project if provided
        if project_id:
            conn = sqlite3.connect(DATABASE)
            c = conn.cursor()
            c.execute("SELECT name, gwp_total, mci_score FROM projects WHERE id = ?", (project_id,))
            project = c.fetchone()
            if project:
                context += f"\nProject: {project[0]}, GWP: {project[1]} kg CO₂-eq, MCI: {project[2]}"
            
            c.execute("""SELECT material_name, material_type, quantity, gwp, recycled_content 
                        FROM project_materials WHERE project_id = ?""", (project_id,))
            materials = c.fetchall()
            if materials:
                mat_context = "\nMaterials: " + ", ".join([f"{m[0]} ({m[1]}): {m[3]}kg CO₂" for m in materials])
                context += mat_context
            conn.close()
        
        result = get_ai_response(prompt, context)
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"AI Chat Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/ai/transcribe', methods=['POST', 'OPTIONS'])
def ai_transcribe_voice():
    """Transcribe voice audio to text using Groq Whisper"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        if 'audio' not in request.files:
            return jsonify({"detail": "No audio file provided"}), 400
        
        audio_file = request.files['audio']
        
        if not groq_client:
            return jsonify({"detail": "Voice transcription not available - Groq not configured"}), 503
        
        # Save temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as tmp:
            audio_file.save(tmp.name)
            tmp_path = tmp.name
        
        try:
            # Use Groq's Whisper model for transcription
            with open(tmp_path, 'rb') as f:
                transcription = groq_client.audio.transcriptions.create(
                    file=(audio_file.filename or 'audio.webm', f.read()),
                    model="whisper-large-v3",
                    language="en",
                    response_format="json"
                )
            
            return jsonify({
                "text": transcription.text,
                "success": True
            }), 200
            
        finally:
            # Clean up temp file
            import os
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        
    except Exception as e:
        print(f"Voice Transcription Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/ai/analyze', methods=['POST', 'OPTIONS'])
def ai_analyze_project():
    """Deep AI analysis of a project"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.get_json()
        project_id = data.get('project_id')
        
        if not project_id:
            return jsonify({"detail": "Project ID is required"}), 400
        
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("""SELECT material_name, material_type, quantity, gwp, recycled_content, transport_distance
                    FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        # Build analysis context
        context = f"""
Project: {project[1]}
Category: {project[4]}
Target Lifespan: {project[5]} years
Designed for Disassembly: {'Yes' if project[6] else 'No'}
Total GWP: {project[8]} kg CO₂-eq
MCI Score: {project[9]}
Circular Design Score: {project[10]}

Materials:
"""
        for m in materials:
            context += f"- {m[0]} ({m[1]}): {m[2]}kg, GWP={m[3]}, Recycled={m[4]}%, Transport={m[5]}km\n"
        
        prompt = """Analyze this project and provide:
1. Key environmental hotspots
2. Top 3 improvement recommendations with estimated impact
3. Comparison to industry benchmarks
4. CBAM readiness assessment"""
        
        result = get_ai_response(prompt, context)
        result['project_id'] = project_id
        
        return jsonify(result), 200
        
    except Exception as e:
        print(f"AI Analyze Error: {e}")
        return jsonify({"detail": str(e)}), 500


# =====================================================
# ENGINE 7: CBAM EXPORT & COMPLIANCE
# =====================================================

# CBAM covered goods categories
CBAM_CATEGORIES = {
    'iron_steel': {
        'name': 'Iron and Steel',
        'cn_codes': ['7201', '7202', '7203', '7204', '7205', '7206', '7207', '7208', '7209', '7210', '7211', '7212', '7213', '7214', '7215', '7216', '7217', '7218', '7219', '7220', '7221', '7222', '7223', '7224', '7225', '7226', '7227', '7228', '7229'],
        'default_cn': '7208',
        'benchmark_ef': 1.85  # tCO2/t product
    },
    'aluminium': {
        'name': 'Aluminium',
        'cn_codes': ['7601', '7602', '7603', '7604', '7605', '7606', '7607', '7608', '7609'],
        'default_cn': '7601',
        'benchmark_ef': 1.47  # tCO2/t unwrought aluminium
    },
    'copper': {
        'name': 'Copper (not CBAM-covered but included for completeness)',
        'cn_codes': ['7401', '7402', '7403', '7404', '7405', '7406', '7407', '7408', '7409'],
        'default_cn': '7403',
        'benchmark_ef': 2.5  # tCO2/t
    }
}

def get_cbam_category(material_type):
    """Determine CBAM category from material type"""
    mat_lower = material_type.lower()
    if 'aluminium' in mat_lower or 'aluminum' in mat_lower:
        return 'aluminium'
    elif 'steel' in mat_lower or 'iron' in mat_lower:
        return 'iron_steel'
    elif 'copper' in mat_lower:
        return 'copper'
    return None


@app.route('/api/v1/projects/<project_id>/cbam-export', methods=['GET', 'OPTIONS'])
def export_cbam_report(project_id):
    """Generate CBAM-compliant export report for EU compliance"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get user info
        c.execute("SELECT full_name, organization_name, email FROM users WHERE id = ?", (project[7],))
        user = c.fetchone()
        
        # Get materials
        c.execute("""SELECT material_name, material_type, quantity, gwp, recycled_content, transport_distance
                    FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        if not materials:
            return jsonify({"detail": "No materials to export"}), 400
        
        # Build CBAM report
        report_date = datetime.now().isoformat()
        quarter = (datetime.now().month - 1) // 3 + 1
        reporting_period = f"{datetime.now().year}-Q{quarter}"
        
        # Categorize materials by CBAM category
        cbam_goods = []
        total_embedded_emissions = 0
        
        for m in materials:
            material_name, material_type, quantity, gwp, recycled_content, transport_distance = m
            cbam_cat = get_cbam_category(material_type)
            
            if cbam_cat:
                cat_info = CBAM_CATEGORIES[cbam_cat]
                # Embedded emissions in tonnes CO2
                embedded = (gwp or 0) / 1000  # Convert kg to tonnes
                total_embedded_emissions += embedded
                
                cbam_goods.append({
                    'product_description': material_name,
                    'cn_code': cat_info['default_cn'],
                    'cbam_category': cat_info['name'],
                    'quantity_kg': quantity,
                    'quantity_tonnes': quantity / 1000,
                    'embedded_emissions_tco2': round(embedded, 4),
                    'specific_embedded_emissions': round(embedded / (quantity / 1000), 4) if quantity > 0 else 0,
                    'recycled_content_percent': recycled_content or 0,
                    'benchmark_ef': cat_info['benchmark_ef'],
                    'production_country': 'IN',  # India
                    'installation_name': user[1] if user else 'Unknown',
                    'verification_status': 'pending'
                })
        
        # Calculate summary
        total_quantity_tonnes = sum(g['quantity_tonnes'] for g in cbam_goods)
        avg_specific_emissions = total_embedded_emissions / total_quantity_tonnes if total_quantity_tonnes > 0 else 0
        
        # CBAM certificate calculation (estimated)
        # Price is based on EU ETS price (approx €80-100/tCO2)
        estimated_ets_price = 90  # EUR/tCO2
        estimated_cbam_liability = total_embedded_emissions * estimated_ets_price
        
        report = {
            'report_metadata': {
                'report_id': f"CBAM-{project_id[:8].upper()}-{datetime.now().strftime('%Y%m%d')}",
                'generation_date': report_date,
                'reporting_period': reporting_period,
                'regulation_reference': 'EU Regulation 2023/956 (CBAM)',
                'report_type': 'Quarterly CBAM Declaration',
                'software_version': 'JNARRDC LCA Portal v1.0'
            },
            'declarant_information': {
                'company_name': user[1] if user else 'Not specified',
                'contact_person': user[0] if user else 'Not specified',
                'email': user[2] if user else 'Not specified',
                'country_of_origin': 'India',
                'eori_number': 'To be provided'
            },
            'project_information': {
                'project_id': project_id,
                'project_name': project[1],
                'product_category': project[4],
                'description': project[2] or 'Not specified'
            },
            'goods_declaration': cbam_goods,
            'summary': {
                'total_goods_categories': len(set(g['cbam_category'] for g in cbam_goods)),
                'total_quantity_tonnes': round(total_quantity_tonnes, 3),
                'total_embedded_emissions_tco2': round(total_embedded_emissions, 4),
                'average_specific_emissions': round(avg_specific_emissions, 4),
                'estimated_ets_price_eur': estimated_ets_price,
                'estimated_cbam_liability_eur': round(estimated_cbam_liability, 2)
            },
            'verification_requirements': {
                'accredited_verifier_required': total_embedded_emissions > 500,
                'verification_deadline': f"{datetime.now().year + 1}-05-31",
                'documentation_required': [
                    'Production process documentation',
                    'Energy consumption records',
                    'Emission factor calculations',
                    'Recycled content certificates',
                    'Transport documentation'
                ]
            },
            'compliance_notes': [
                'This report is generated for CBAM compliance purposes.',
                'Actual embedded emissions may vary based on production methods.',
                'Verification by accredited body required before final submission.',
                'Default values used where actual data not available.',
                f"India currently has no carbon price applicable for CBAM deduction."
            ]
        }
        
        return jsonify(report), 200
        
    except Exception as e:
        print(f"CBAM Export Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/cbam-export/csv', methods=['GET', 'OPTIONS'])
def export_cbam_csv(project_id):
    """Export CBAM data in CSV format"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Get the JSON report first
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        c.execute("""SELECT material_name, material_type, quantity, gwp, recycled_content
                    FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        # Build CSV content
        csv_lines = [
            "CBAM Quarterly Report - JNARRDC LCA Portal Export",
            f"Project: {project[1]}",
            f"Report Date: {datetime.now().isoformat()}",
            "",
            "CN Code,Product Description,CBAM Category,Quantity (tonnes),Embedded Emissions (tCO2),Specific Emissions (tCO2/t),Recycled Content (%),Country of Origin",
        ]
        
        for m in materials:
            material_name, material_type, quantity, gwp, recycled_content = m
            cbam_cat = get_cbam_category(material_type)
            if cbam_cat:
                cat_info = CBAM_CATEGORIES[cbam_cat]
                quantity_t = quantity / 1000
                embedded = (gwp or 0) / 1000
                specific = embedded / quantity_t if quantity_t > 0 else 0
                csv_lines.append(
                    f"{cat_info['default_cn']},{material_name},{cat_info['name']},{quantity_t:.4f},{embedded:.4f},{specific:.4f},{recycled_content or 0},IN"
                )
        
        csv_content = "\n".join(csv_lines)
        
        return csv_content, 200, {
            'Content-Type': 'text/csv',
            'Content-Disposition': f'attachment; filename="cbam_report_{project_id[:8]}.csv"'
        }
        
    except Exception as e:
        print(f"CBAM CSV Export Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/cbam-export/excel', methods=['GET', 'OPTIONS'])
def export_cbam_excel(project_id):
    """Export CBAM report as professional Excel letterhead document"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get user info
        c.execute("SELECT full_name, organization_name, email FROM users WHERE id = ?", (project[7],))
        user = c.fetchone()
        
        # Get materials
        c.execute("""SELECT material_name, material_type, quantity, gwp, recycled_content, transport_distance
                    FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        conn.close()
        
        if not materials:
            return jsonify({"detail": "No materials to export"}), 400
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "CBAM Report"
        
        # Define styles
        header_font = Font(name='Arial', size=18, bold=True, color='1F4E79')
        subheader_font = Font(name='Arial', size=12, bold=True, color='2E75B6')
        title_font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
        normal_font = Font(name='Arial', size=10)
        bold_font = Font(name='Arial', size=10, bold=True)
        small_font = Font(name='Arial', size=9, color='666666')
        
        header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
        alt_row_fill = PatternFill(start_color='F2F2F2', end_color='F2F2F2', fill_type='solid')
        light_blue_fill = PatternFill(start_color='DEEAF6', end_color='DEEAF6', fill_type='solid')
        
        thin_border = Border(
            left=Side(style='thin', color='CCCCCC'),
            right=Side(style='thin', color='CCCCCC'),
            top=Side(style='thin', color='CCCCCC'),
            bottom=Side(style='thin', color='CCCCCC')
        )
        
        # Set column widths
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 25
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 18
        ws.column_dimensions['G'].width = 15
        ws.column_dimensions['H'].width = 12
        
        row = 1
        
        # ===== LETTERHEAD =====
        ws.merge_cells('B1:G1')
        ws['B1'] = 'JNARRDC LCA PORTAL'
        ws['B1'].font = header_font
        ws['B1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30
        
        row = 2
        ws.merge_cells('B2:G2')
        ws['B2'] = 'Jawaharlal Nehru Aluminium Research Development and Design Centre'
        ws['B2'].font = Font(name='Arial', size=10, italic=True, color='666666')
        ws['B2'].alignment = Alignment(horizontal='center')
        
        row = 3
        ws.merge_cells('B3:G3')
        ws['B3'] = '━' * 60
        ws['B3'].font = Font(color='1F4E79')
        ws['B3'].alignment = Alignment(horizontal='center')
        
        # ===== REPORT TITLE =====
        row = 5
        ws.merge_cells('B5:G5')
        ws['B5'] = 'CBAM QUARTERLY DECLARATION REPORT'
        ws['B5'].font = Font(name='Arial', size=14, bold=True, color='1F4E79')
        ws['B5'].alignment = Alignment(horizontal='center')
        
        row = 6
        ws.merge_cells('B6:G6')
        ws['B6'] = 'Carbon Border Adjustment Mechanism - EU Regulation 2023/956'
        ws['B6'].font = small_font
        ws['B6'].alignment = Alignment(horizontal='center')
        
        # ===== REPORT METADATA =====
        row = 8
        quarter = (datetime.now().month - 1) // 3 + 1
        report_id = f"CBAM-{project_id[:8].upper()}-{datetime.now().strftime('%Y%m%d')}"
        
        ws['B8'] = 'Report ID:'
        ws['B8'].font = bold_font
        ws['C8'] = report_id
        ws['C8'].font = normal_font
        
        ws['E8'] = 'Report Date:'
        ws['E8'].font = bold_font
        ws['F8'] = datetime.now().strftime('%d %B %Y')
        ws['F8'].font = normal_font
        
        row = 9
        ws['B9'] = 'Reporting Period:'
        ws['B9'].font = bold_font
        ws['C9'] = f"{datetime.now().year}-Q{quarter}"
        ws['C9'].font = normal_font
        
        ws['E9'] = 'Status:'
        ws['E9'].font = bold_font
        ws['F9'] = 'Draft - Pending Verification'
        ws['F9'].font = Font(name='Arial', size=10, color='FF6600')
        
        # ===== DECLARANT INFORMATION =====
        row = 11
        ws.merge_cells('B11:G11')
        ws['B11'] = 'DECLARANT INFORMATION'
        ws['B11'].font = subheader_font
        ws['B11'].fill = light_blue_fill
        ws.row_dimensions[11].height = 22
        
        row = 12
        ws['B12'] = 'Organization:'
        ws['B12'].font = bold_font
        ws['C12'] = user[1] if user else 'Not specified'
        ws['C12'].font = normal_font
        
        ws['E12'] = 'Contact Person:'
        ws['E12'].font = bold_font
        ws['F12'] = user[0] if user else 'Not specified'
        ws['F12'].font = normal_font
        
        row = 13
        ws['B13'] = 'Email:'
        ws['B13'].font = bold_font
        ws['C13'] = user[2] if user else 'Not specified'
        ws['C13'].font = normal_font
        
        ws['E13'] = 'Country of Origin:'
        ws['E13'].font = bold_font
        ws['F13'] = 'India (IN)'
        ws['F13'].font = normal_font
        
        # ===== PROJECT INFORMATION =====
        row = 15
        ws.merge_cells('B15:G15')
        ws['B15'] = 'PROJECT INFORMATION'
        ws['B15'].font = subheader_font
        ws['B15'].fill = light_blue_fill
        ws.row_dimensions[15].height = 22
        
        row = 16
        ws['B16'] = 'Project Name:'
        ws['B16'].font = bold_font
        ws['C16'] = project[1]
        ws['C16'].font = normal_font
        
        ws['E16'] = 'Product Category:'
        ws['E16'].font = bold_font
        ws['F16'] = project[4] or 'Not specified'
        ws['F16'].font = normal_font
        
        row = 17
        ws['B17'] = 'Description:'
        ws['B17'].font = bold_font
        ws.merge_cells('C17:G17')
        ws['C17'] = project[2] or 'Not specified'
        ws['C17'].font = normal_font
        
        # ===== GOODS DECLARATION TABLE =====
        row = 19
        ws.merge_cells('B19:G19')
        ws['B19'] = 'GOODS DECLARATION'
        ws['B19'].font = subheader_font
        ws['B19'].fill = light_blue_fill
        ws.row_dimensions[19].height = 22
        
        # Table headers
        row = 21
        headers = ['#', 'Product Description', 'CN Code', 'CBAM Category', 'Qty (tonnes)', 'Emissions (tCO₂)', 'Recycled %']
        for col, header in enumerate(headers, start=2):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = title_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        ws.row_dimensions[21].height = 25
        
        # Table data
        row = 22
        total_quantity = 0
        total_emissions = 0
        
        for idx, m in enumerate(materials, 1):
            material_name, material_type, quantity, gwp, recycled_content, transport_distance = m
            cbam_cat = get_cbam_category(material_type)
            
            if cbam_cat:
                cat_info = CBAM_CATEGORIES[cbam_cat]
                quantity_t = quantity / 1000
                embedded = (gwp or 0) / 1000
                total_quantity += quantity_t
                total_emissions += embedded
                
                row_data = [
                    idx,
                    material_name,
                    cat_info['default_cn'],
                    cat_info['name'],
                    round(quantity_t, 4),
                    round(embedded, 4),
                    f"{recycled_content or 0}%"
                ]
                
                for col, value in enumerate(row_data, start=2):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = normal_font
                    cell.border = thin_border
                    cell.alignment = Alignment(horizontal='center' if col > 2 else 'left', vertical='center')
                    if idx % 2 == 0:
                        cell.fill = alt_row_fill
                
                row += 1
        
        # Total row
        ws.cell(row=row, column=2, value='TOTAL').font = bold_font
        ws.cell(row=row, column=2).border = thin_border
        ws.cell(row=row, column=3).border = thin_border
        ws.cell(row=row, column=4).border = thin_border
        ws.cell(row=row, column=5, value=round(total_quantity, 4)).font = bold_font
        ws.cell(row=row, column=5).border = thin_border
        ws.cell(row=row, column=5).alignment = Alignment(horizontal='center')
        ws.cell(row=row, column=6, value=round(total_emissions, 4)).font = bold_font
        ws.cell(row=row, column=6).border = thin_border
        ws.cell(row=row, column=6).alignment = Alignment(horizontal='center')
        ws.cell(row=row, column=7).border = thin_border
        ws.cell(row=row, column=8).border = thin_border
        
        # ===== SUMMARY =====
        row += 2
        ws.merge_cells(f'B{row}:G{row}')
        ws[f'B{row}'] = 'EMISSIONS SUMMARY & CBAM LIABILITY'
        ws[f'B{row}'].font = subheader_font
        ws[f'B{row}'].fill = light_blue_fill
        ws.row_dimensions[row].height = 22
        
        row += 1
        estimated_price = 90
        estimated_liability = total_emissions * estimated_price
        
        ws[f'B{row}'] = 'Total Embedded Emissions:'
        ws[f'B{row}'].font = bold_font
        ws[f'D{row}'] = f'{round(total_emissions, 4)} tCO₂'
        ws[f'D{row}'].font = normal_font
        
        row += 1
        ws[f'B{row}'] = 'EU ETS Reference Price:'
        ws[f'B{row}'].font = bold_font
        ws[f'D{row}'] = f'€{estimated_price}/tCO₂'
        ws[f'D{row}'].font = normal_font
        
        row += 1
        ws[f'B{row}'] = 'Estimated CBAM Liability:'
        ws[f'B{row}'].font = bold_font
        ws[f'D{row}'] = f'€{round(estimated_liability, 2)}'
        ws[f'D{row}'].font = Font(name='Arial', size=12, bold=True, color='C00000')
        
        row += 1
        ws[f'B{row}'] = 'Carbon Price Deduction (India):'
        ws[f'B{row}'].font = bold_font
        ws[f'D{row}'] = '€0.00 (No applicable carbon price)'
        ws[f'D{row}'].font = small_font
        
        # ===== FOOTER =====
        row += 3
        ws.merge_cells(f'B{row}:G{row}')
        ws[f'B{row}'] = '━' * 60
        ws[f'B{row}'].font = Font(color='CCCCCC')
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        
        row += 1
        ws.merge_cells(f'B{row}:G{row}')
        ws[f'B{row}'] = 'This report is generated for CBAM compliance purposes. Verification by accredited body required.'
        ws[f'B{row}'].font = small_font
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        
        row += 1
        ws.merge_cells(f'B{row}:G{row}')
        ws[f'B{row}'] = f'Generated by JNARRDC LCA Portal v1.0 | Data Sources: IPCC AR6, Ecoinvent 3.9 | {datetime.now().strftime("%d/%m/%Y %H:%M")}'
        ws[f'B{row}'].font = Font(name='Arial', size=8, italic=True, color='999999')
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        
        # Save to BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f"CBAM_Report_{project[1].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"CBAM Excel Export Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"detail": str(e)}), 500


# =====================================================
# BRSR EXPORT - SEBI Business Responsibility Report
# =====================================================

@app.route('/api/v1/projects/<project_id>/brsr-export', methods=['GET', 'OPTIONS'])
def brsr_export(project_id):
    """Generate BRSR (SEBI) Principle 6 Environmental Report"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        # Get project
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        # Get materials
        c.execute("""SELECT material_name, material_type, quantity, gwp, 
                     recycled_content, transport_distance 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        
        # Get user info
        c.execute("SELECT full_name, organization_name, email FROM users WHERE id = ?", (project[3],))
        user = c.fetchone()
        conn.close()
        
        # Calculate metrics
        total_mass = sum(m[2] or 0 for m in materials)
        total_gwp = sum(m[3] or 0 for m in materials)
        avg_recycled = sum(m[4] or 0 for m in materials) / len(materials) if materials else 0
        
        # Calculate virgin and recycled inputs
        recycled_input = sum((m[2] or 0) * (m[4] or 0) / 100 for m in materials)
        virgin_input = total_mass - recycled_input
        
        # Calculate waste per material using dynamic waste factors
        waste_by_material = []
        total_waste = 0
        for m in materials:
            material_name = m[0] or ''
            material_type = m[1] or ''
            quantity = m[2] or 0
            recycled_content = m[4] or 0
            
            # Get waste factor for this material type
            waste_factor = get_waste_factor(material_type)
            material_waste = quantity * waste_factor
            total_waste += material_waste
            
            waste_by_material.append({
                'material': material_name,
                'material_type': material_type,
                'quantity_kg': round(quantity, 2),
                'waste_factor_percent': round(waste_factor * 100, 1),
                'waste_generated_kg': round(material_waste, 2),
                'waste_recycled_kg': round(material_waste * (recycled_content / 100), 2),
                'waste_to_landfill_kg': round(material_waste * (1 - recycled_content / 100), 2)
            })
        
        # Calculate weighted average waste factor
        avg_waste_factor = (total_waste / total_mass) if total_mass > 0 else 0.05
        
        # BRSR Principle 6 Format
        report = {
            'report_metadata': {
                'report_id': f"BRSR-{project_id[:8].upper()}-{datetime.now().strftime('%Y%m%d')}",
                'report_type': 'BRSR Principle 6 - Environment',
                'regulation_reference': 'SEBI Circular SEBI/HO/CFD/CMD-2/P/CIR/2021/562',
                'financial_year': f"FY {datetime.now().year}-{(datetime.now().year + 1) % 100:02d}",
                'generation_date': datetime.now().isoformat(),
                'software': 'JNARRDC LCA Portal v1.0'
            },
            'entity_details': {
                'name_of_listed_entity': user[1] if user else 'Not specified',
                'cin': 'To be provided',
                'year_of_incorporation': 'To be provided',
                'registered_office_address': 'To be provided',
                'reporting_boundary': 'Standalone'
            },
            'principle_6_essential_indicators': {
                'section_a_energy_consumption': {
                    'disclosure': 'Details of total energy consumption and energy intensity',
                    'total_energy_consumption_gj': round(total_gwp * 0.0036, 2),  # Estimate
                    'energy_intensity_per_unit': round((total_gwp * 0.0036) / total_mass, 4) if total_mass > 0 else 0,
                    'renewable_energy_percent': 0,  # To be provided
                    'note': 'Energy values estimated from GWP calculations'
                },
                'section_b_water': {
                    'disclosure': 'Water withdrawal and consumption',
                    'total_water_withdrawal_kl': 'To be provided',
                    'water_intensity': 'To be provided',
                    'note': 'Water data not available from LCA module'
                },
                'section_c_emissions': {
                    'disclosure': 'Details of air emissions and GHG emissions',
                    'scope_1_emissions_mtco2e': 0,  # Direct emissions - not calculated in this LCA
                    'scope_2_emissions_mtco2e': round(total_gwp / 1000, 4),  # Indirect from materials
                    'scope_3_emissions_mtco2e': round(total_gwp / 1000 * 0.1, 4),  # Estimate for transport
                    'total_ghg_emissions_mtco2e': round(total_gwp / 1000, 4),
                    'ghg_intensity_per_rupee_turnover': 'To be calculated',
                    'methodology': 'IPCC AR6 emission factors via JNARRDC LCA Portal',
                    'breakdown_by_material': [
                        {
                            'material': m[0],
                            'gwp_kgco2eq': round(m[3] or 0, 2),
                            'percentage': round((m[3] or 0) / total_gwp * 100, 1) if total_gwp > 0 else 0
                        } for m in materials
                    ]
                },
                'section_d_waste': {
                    'disclosure': 'Details of waste generated and recycled',
                    'total_waste_generated_mt': round(total_waste / 1000, 4),
                    'waste_recycled_mt': round(sum(w['waste_recycled_kg'] for w in waste_by_material) / 1000, 4),
                    'waste_to_landfill_mt': round(sum(w['waste_to_landfill_kg'] for w in waste_by_material) / 1000, 4),
                    'waste_intensity': round(avg_waste_factor, 4),
                    'avg_waste_factor_percent': round(avg_waste_factor * 100, 1),
                    'breakdown_by_material': waste_by_material,
                    'note': 'Waste estimated using material-specific waste factors based on industry averages'
                },
                'section_e_circularity': {
                    'disclosure': 'Details related to circularity',
                    'total_raw_material_consumed_mt': round(total_mass / 1000, 4),
                    'recycled_input_material_mt': round(recycled_input / 1000, 4),
                    'virgin_input_material_mt': round(virgin_input / 1000, 4),
                    'recycled_content_percentage': round(avg_recycled, 1),
                    'products_reclaimed_at_eol_percent': round(avg_recycled * 0.8, 1),  # Estimate
                    'mci_score': project[10] if len(project) > 10 else 0,
                    'circular_design_score': project[11] if len(project) > 11 else 0,
                    'designed_for_disassembly': bool(project[6]) if len(project) > 6 else False
                }
            },
            'principle_6_leadership_indicators': {
                'water_discharge_quality': 'Not applicable to LCA scope',
                'biodiversity_impact': 'Not assessed',
                'emission_management': {
                    'nox_sox_control': 'To be provided',
                    'pm_control': 'To be provided'
                },
                'environmental_compliance': {
                    'show_cause_notices': 0,
                    'penalties_paid': 0
                }
            },
            'material_breakdown': [
                {
                    'material_name': m[0],
                    'material_type': m[1],
                    'quantity_kg': m[2] or 0,
                    'recycled_content_percent': m[4] or 0,
                    'gwp_kgco2eq': round(m[3] or 0, 2),
                    'virgin_fraction_kg': round((m[2] or 0) * (100 - (m[4] or 0)) / 100, 2),
                    'recycled_fraction_kg': round((m[2] or 0) * (m[4] or 0) / 100, 2)
                } for m in materials
            ],
            'summary_metrics': {
                'total_materials_kg': round(total_mass, 2),
                'total_gwp_kgco2eq': round(total_gwp, 2),
                'average_recycled_content': round(avg_recycled, 1),
                'circular_readiness': 'High' if avg_recycled >= 50 else ('Medium' if avg_recycled >= 25 else 'Low')
            },
            'compliance_notes': [
                'This report follows SEBI BRSR format for Principle 6 (Environment)',
                'GWP calculations use IPCC AR6 and Ecoinvent 3.9 emission factors',
                'MCI calculated using Ellen MacArthur Foundation methodology',
                'Some indicators require additional data from facility operations',
                'Report generated for disclosure purposes - verify with auditor'
            ]
        }
        
        return jsonify(report), 200
        
    except Exception as e:
        print(f"BRSR Export Error: {e}")
        return jsonify({"detail": str(e)}), 500


@app.route('/api/v1/projects/<project_id>/brsr-export/excel', methods=['GET', 'OPTIONS'])
def brsr_export_excel(project_id):
    """Generate BRSR Excel Report with SEBI Format"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        conn = sqlite3.connect(DATABASE)
        c = conn.cursor()
        
        c.execute("SELECT * FROM projects WHERE id = ?", (project_id,))
        project = c.fetchone()
        if not project:
            conn.close()
            return jsonify({"detail": "Project not found"}), 404
        
        c.execute("""SELECT material_name, material_type, quantity, gwp, 
                     recycled_content, transport_distance 
                     FROM project_materials WHERE project_id = ?""", (project_id,))
        materials = c.fetchall()
        
        c.execute("SELECT full_name, organization_name, email FROM users WHERE id = ?", (project[3],))
        user = c.fetchone()
        conn.close()
        
        # Calculate metrics
        total_mass = sum(m[2] or 0 for m in materials)
        total_gwp = sum(m[3] or 0 for m in materials)
        avg_recycled = sum(m[4] or 0 for m in materials) / len(materials) if materials else 0
        recycled_input = sum((m[2] or 0) * (m[4] or 0) / 100 for m in materials)
        virgin_input = total_mass - recycled_input
        
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "BRSR Principle 6"
        
        # Styles
        header_font = Font(name='Arial', size=16, bold=True, color='1F4E79')
        subheader_font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
        bold_font = Font(name='Arial', size=10, bold=True)
        normal_font = Font(name='Arial', size=10)
        small_font = Font(name='Arial', size=9, italic=True, color='666666')
        green_fill = PatternFill(start_color='2E7D32', end_color='2E7D32', fill_type='solid')
        light_green_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
        
        # Column widths
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 25
        ws.column_dimensions['D'].width = 20
        ws.column_dimensions['E'].width = 20
        
        row = 1
        
        # Header
        ws.merge_cells('B1:E1')
        ws['B1'] = 'BUSINESS RESPONSIBILITY AND SUSTAINABILITY REPORT'
        ws['B1'].font = header_font
        ws['B1'].alignment = Alignment(horizontal='center')
        ws.row_dimensions[1].height = 30
        
        row = 2
        ws.merge_cells('B2:E2')
        ws['B2'] = 'PRINCIPLE 6: ENVIRONMENT - JNARRDC LCA Portal'
        ws['B2'].font = Font(name='Arial', size=12, italic=True, color='666666')
        ws['B2'].alignment = Alignment(horizontal='center')
        
        row = 4
        ws.merge_cells('B4:E4')
        ws['B4'] = 'SECTION A: ENERGY CONSUMPTION'
        ws['B4'].font = subheader_font
        ws['B4'].fill = green_fill
        ws.row_dimensions[4].height = 22
        
        row = 5
        ws['B5'] = 'Parameter'
        ws['C5'] = 'Unit'
        ws['D5'] = 'Current FY'
        ws['E5'] = 'Previous FY'
        for cell in ['B5', 'C5', 'D5', 'E5']:
            ws[cell].font = bold_font
            ws[cell].fill = light_green_fill
        
        row = 6
        ws['B6'] = 'Total energy consumption'
        ws['C6'] = 'GJ'
        ws['D6'] = round(total_gwp * 0.0036, 2)
        ws['E6'] = '-'
        
        row = 7
        ws['B7'] = 'Energy intensity per unit output'
        ws['C7'] = 'GJ/MT'
        ws['D7'] = round((total_gwp * 0.0036) / (total_mass/1000), 4) if total_mass > 0 else 0
        ws['E7'] = '-'
        
        # Section C: Emissions
        row = 9
        ws.merge_cells('B9:E9')
        ws['B9'] = 'SECTION C: GHG EMISSIONS'
        ws['B9'].font = subheader_font
        ws['B9'].fill = green_fill
        ws.row_dimensions[9].height = 22
        
        row = 10
        ws['B10'] = 'Parameter'
        ws['C10'] = 'Unit'
        ws['D10'] = 'Current FY'
        ws['E10'] = 'Previous FY'
        for cell in ['B10', 'C10', 'D10', 'E10']:
            ws[cell].font = bold_font
            ws[cell].fill = light_green_fill
        
        row = 11
        ws['B11'] = 'Scope 1 (Direct emissions)'
        ws['C11'] = 'MTCO2e'
        ws['D11'] = 0
        ws['E11'] = '-'
        
        row = 12
        ws['B12'] = 'Scope 2 (Indirect - Materials)'
        ws['C12'] = 'MTCO2e'
        ws['D12'] = round(total_gwp / 1000, 4)
        ws['E12'] = '-'
        
        row = 13
        ws['B13'] = 'Scope 3 (Transport)'
        ws['C13'] = 'MTCO2e'
        ws['D13'] = round(total_gwp / 1000 * 0.1, 4)
        ws['E13'] = '-'
        
        row = 14
        ws['B14'] = 'Total GHG Emissions'
        ws['C14'] = 'MTCO2e'
        ws['D14'] = round(total_gwp / 1000, 4)
        ws['D14'].font = bold_font
        ws['E14'] = '-'
        
        # Section E: Circularity
        row = 16
        ws.merge_cells('B16:E16')
        ws['B16'] = 'SECTION E: CIRCULARITY'
        ws['B16'].font = subheader_font
        ws['B16'].fill = green_fill
        ws.row_dimensions[16].height = 22
        
        row = 17
        ws['B17'] = 'Parameter'
        ws['C17'] = 'Unit'
        ws['D17'] = 'Current FY'
        ws['E17'] = 'Previous FY'
        for cell in ['B17', 'C17', 'D17', 'E17']:
            ws[cell].font = bold_font
            ws[cell].fill = light_green_fill
        
        row = 18
        ws['B18'] = 'Total raw material consumed'
        ws['C18'] = 'MT'
        ws['D18'] = round(total_mass / 1000, 4)
        ws['E18'] = '-'
        
        row = 19
        ws['B19'] = 'Recycled input material'
        ws['C19'] = 'MT'
        ws['D19'] = round(recycled_input / 1000, 4)
        ws['E19'] = '-'
        
        row = 20
        ws['B20'] = 'Virgin input material'
        ws['C20'] = 'MT'
        ws['D20'] = round(virgin_input / 1000, 4)
        ws['E20'] = '-'
        
        row = 21
        ws['B21'] = 'Recycled content percentage'
        ws['C21'] = '%'
        ws['D21'] = round(avg_recycled, 1)
        ws['E21'] = '-'
        
        row = 22
        ws['B22'] = 'Material Circularity Index (MCI)'
        ws['C22'] = 'Score (0-1)'
        ws['D22'] = project[10] if len(project) > 10 else 0
        ws['E22'] = '-'
        
        # Material breakdown
        row = 24
        ws.merge_cells('B24:E24')
        ws['B24'] = 'MATERIAL BREAKDOWN'
        ws['B24'].font = subheader_font
        ws['B24'].fill = green_fill
        ws.row_dimensions[24].height = 22
        
        row = 25
        ws['B25'] = 'Material'
        ws['C25'] = 'Quantity (kg)'
        ws['D25'] = 'Recycled %'
        ws['E25'] = 'GWP (kgCO2eq)'
        for cell in ['B25', 'C25', 'D25', 'E25']:
            ws[cell].font = bold_font
            ws[cell].fill = light_green_fill
        
        row = 26
        for m in materials:
            ws[f'B{row}'] = m[0]
            ws[f'C{row}'] = m[2] or 0
            ws[f'D{row}'] = f"{m[4] or 0}%"
            ws[f'E{row}'] = round(m[3] or 0, 2)
            row += 1
        
        # Footer
        row += 2
        ws.merge_cells(f'B{row}:E{row}')
        ws[f'B{row}'] = f'Report generated by JNARRDC LCA Portal | {datetime.now().strftime("%d/%m/%Y")} | SEBI BRSR Format'
        ws[f'B{row}'].font = small_font
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        
        # Save
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f"BRSR_Report_{project[1].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"BRSR Excel Export Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"detail": str(e)}), 500


if __name__ == '__main__':
    # Initialize databases
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id TEXT PRIMARY KEY, 
                  email TEXT UNIQUE NOT NULL,
                  password TEXT NOT NULL,
                  full_name TEXT,
                  organization_name TEXT,
                  created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS projects
                 (id TEXT PRIMARY KEY,
                  name TEXT NOT NULL,
                  description TEXT,
                  status TEXT,
                  product_category TEXT,
                  target_lifespan INTEGER,
                  is_designed_for_disassembly INTEGER,
                  user_id TEXT NOT NULL,
                  gwp_total REAL DEFAULT 0,
                  mci_score REAL DEFAULT 0,
                  circular_design_score REAL DEFAULT 0,
                  created_at TEXT,
                  FOREIGN KEY(user_id) REFERENCES users(id))''')
    
    # Migration: Add new columns if they don't exist
    try:
        c.execute("SELECT gwp_total FROM projects LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating projects table: adding gwp_total column")
        c.execute("ALTER TABLE projects ADD COLUMN gwp_total REAL DEFAULT 0")
        conn.commit()
    
    try:
        c.execute("SELECT mci_score FROM projects LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating projects table: adding mci_score column")
        c.execute("ALTER TABLE projects ADD COLUMN mci_score REAL DEFAULT 0")
        conn.commit()
    
    try:
        c.execute("SELECT circular_design_score FROM projects LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating projects table: adding circular_design_score column")
        c.execute("ALTER TABLE projects ADD COLUMN circular_design_score REAL DEFAULT 0")
        conn.commit()
    
    # Migration: Add tier and tier_expires_at columns to users table
    try:
        c.execute("SELECT tier FROM users LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating users table: adding tier column")
        c.execute("ALTER TABLE users ADD COLUMN tier TEXT DEFAULT 'free'")
        conn.commit()
    
    try:
        c.execute("SELECT tier_expires_at FROM users LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating users table: adding tier_expires_at column")
        c.execute("ALTER TABLE users ADD COLUMN tier_expires_at TEXT")
        conn.commit()
    
    try:
        c.execute("SELECT project_limit FROM users LIMIT 1")
    except sqlite3.OperationalError:
        print("⚠️  Migrating users table: adding project_limit column")
        c.execute("ALTER TABLE users ADD COLUMN project_limit INTEGER DEFAULT 3")
        conn.commit()
        
    c.execute('''CREATE TABLE IF NOT EXISTS project_materials
                 (id TEXT PRIMARY KEY,
                  project_id TEXT NOT NULL,
                  material_name TEXT NOT NULL,
                  material_type TEXT NOT NULL,
                  quantity REAL NOT NULL,
                  unit TEXT,
                  recycled_content REAL DEFAULT 0,
                  gwp REAL,
                  transport_distance REAL DEFAULT 0,
                  created_at TEXT,
                  FOREIGN KEY(project_id) REFERENCES projects(id))''')
    
    # Team Management Tables
    c.execute('''CREATE TABLE IF NOT EXISTS teams
                 (id TEXT PRIMARY KEY,
                  name TEXT NOT NULL,
                  description TEXT,
                  owner_id TEXT NOT NULL,
                  created_at TEXT,
                  FOREIGN KEY(owner_id) REFERENCES users(id))''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS team_members
                 (id TEXT PRIMARY KEY,
                  team_id TEXT NOT NULL,
                  user_id TEXT NOT NULL,
                  role TEXT DEFAULT 'member',
                  joined_at TEXT,
                  FOREIGN KEY(team_id) REFERENCES teams(id),
                  FOREIGN KEY(user_id) REFERENCES users(id),
                  UNIQUE(team_id, user_id))''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS team_invites
                 (id TEXT PRIMARY KEY,
                  team_id TEXT NOT NULL,
                  email TEXT NOT NULL,
                  invited_by TEXT NOT NULL,
                  status TEXT DEFAULT 'pending',
                  created_at TEXT,
                  FOREIGN KEY(team_id) REFERENCES teams(id),
                  FOREIGN KEY(invited_by) REFERENCES users(id))''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS project_collaborators
                 (id TEXT PRIMARY KEY,
                  project_id TEXT NOT NULL,
                  team_id TEXT,
                  user_id TEXT,
                  permission TEXT DEFAULT 'view',
                  added_at TEXT,
                  FOREIGN KEY(project_id) REFERENCES projects(id),
                  FOREIGN KEY(team_id) REFERENCES teams(id),
                  FOREIGN KEY(user_id) REFERENCES users(id))''')
    
    conn.commit()
    conn.close()
    
    print("✅ Database initialized")
    print("🚀 Starting backend on http://127.0.0.1:5000")
    app.run(host='0.0.0.0', port=5000, debug=True)
